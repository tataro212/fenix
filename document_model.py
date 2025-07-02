"""
Structured Document Model for PDF Translation Pipeline

This module defines a comprehensive object-oriented representation of documents
that preserves structural integrity throughout the translation process.

The model addresses the core architectural flaw of treating documents as simple
strings by providing structured content blocks that maintain semantic meaning,
hierarchy, and relationships.
"""

from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional, Union, Tuple
from abc import ABC, abstractmethod
from enum import Enum
import hashlib
import logging
import networkx as nx
import numpy as np
from datetime import datetime

logger = logging.getLogger(__name__)


class ContentType(Enum):
    """Enumeration of content block types"""
    HEADING = "heading"
    PARAGRAPH = "paragraph"
    FOOTNOTE = "footnote"
    TABLE = "table"
    LIST_ITEM = "list_item"
    MATHEMATICAL_FORMULA = "mathematical_formula"
    FIGURE_CAPTION = "figure_caption"
    QUOTE = "quote"
    CODE_BLOCK = "code_block"


@dataclass
class ContentBlock(ABC):
    """
    Abstract base class for all document content blocks.
    
    This class provides the foundation for structured document representation,
    ensuring that all content maintains its semantic meaning and metadata
    throughout the translation pipeline.
    """
    content: str
    page_num: int = 0
    position: int = 0  # Position within the page
    bbox: tuple = field(default_factory=lambda: (0, 0, 0, 0))  # Bounding box (x0, y0, x1, y1)
    font_info: Dict[str, Any] = field(default_factory=dict)
    block_id: str = field(default="")
    
    def __post_init__(self):
        """Generate unique block ID if not provided"""
        if not self.block_id:
            content_hash = hashlib.md5(self.content.encode()).hexdigest()[:8]
            self.block_id = f"{self.get_content_type().value}_{self.page_num}_{content_hash}"
    
    @abstractmethod
    def get_content_type(self) -> ContentType:
        """Return the content type for this block"""
        pass
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert block to dictionary representation"""
        return {
            'type': self.get_content_type().value,
            'content': self.content,
            'page_num': self.page_num,
            'position': self.position,
            'bbox': self.bbox,
            'font_info': self.font_info,
            'block_id': self.block_id
        }


@dataclass
class Heading(ContentBlock):
    """
    Represents a document heading with hierarchical level.
    
    Headings are crucial for TOC generation and document structure.
    The level determines the hierarchy (1 = main heading, 2 = subheading, etc.)
    """
    level: int = 1  # 1 for #, 2 for ##, 3 for ###, etc.
    
    def get_content_type(self) -> ContentType:
        return ContentType.HEADING
    
    def to_dict(self) -> Dict[str, Any]:
        result = super().to_dict()
        result['level'] = self.level
        return result
    
    def get_markdown_prefix(self) -> str:
        """Get the markdown prefix for this heading level"""
        return '#' * self.level


@dataclass
class Paragraph(ContentBlock):
    """
    Represents a standard text paragraph.
    
    Paragraphs are the most common content blocks and preserve
    the natural text flow and formatting.
    """
    
    def get_content_type(self) -> ContentType:
        return ContentType.PARAGRAPH


@dataclass
class Footnote(ContentBlock):
    """
    Represents a footnote with reference information.
    
    Footnotes are collected separately and rendered at the end
    of the document to maintain proper academic formatting.
    """
    reference_id: str = ""  # e.g., "1", "a", "*"
    reference_text: str = ""  # The text that references this footnote
    
    def get_content_type(self) -> ContentType:
        return ContentType.FOOTNOTE
    
    def to_dict(self) -> Dict[str, Any]:
        result = super().to_dict()
        result['reference_id'] = self.reference_id
        result['reference_text'] = self.reference_text
        return result


@dataclass
class Table(ContentBlock):
    """
    Represents a table with structured data.
    
    Tables maintain their markdown representation for now,
    but can be extended to support more sophisticated table handling.
    """
    rows: int = 0
    columns: int = 0
    headers: List[str] = field(default_factory=list)
    
    def get_content_type(self) -> ContentType:
        return ContentType.TABLE
    
    def to_dict(self) -> Dict[str, Any]:
        result = super().to_dict()
        result.update({
            'rows': self.rows,
            'columns': self.columns,
            'headers': self.headers
        })
        return result


@dataclass
class ListItem(ContentBlock):
    """
    Represents a list item (ordered or unordered).
    """
    list_level: int = 1  # Nesting level
    is_ordered: bool = False  # True for numbered lists, False for bullet points
    item_number: Optional[int] = None  # For ordered lists
    
    def get_content_type(self) -> ContentType:
        return ContentType.LIST_ITEM
    
    def to_dict(self) -> Dict[str, Any]:
        result = super().to_dict()
        result.update({
            'list_level': self.list_level,
            'is_ordered': self.is_ordered,
            'item_number': self.item_number
        })
        return result


@dataclass
class MathematicalFormula(ContentBlock):
    """
    Represents mathematical formulas and equations.
    """
    formula_type: str = "inline"  # "inline" or "block"
    latex_representation: str = ""
    
    def get_content_type(self) -> ContentType:
        return ContentType.MATHEMATICAL_FORMULA
    
    def to_dict(self) -> Dict[str, Any]:
        result = super().to_dict()
        result.update({
            'formula_type': self.formula_type,
            'latex_representation': self.latex_representation
        })
        return result


@dataclass
class Page:
    """
    Represents a single page containing multiple content blocks.
    
    Pages maintain the spatial and logical organization of content,
    preserving the document's original structure.
    """
    page_number: int
    content_blocks: List[ContentBlock] = field(default_factory=list)
    page_metadata: Dict[str, Any] = field(default_factory=dict)
    
    def add_block(self, block: ContentBlock) -> None:
        """Add a content block to this page"""
        block.page_num = self.page_number
        if not block.position:
            block.position = len(self.content_blocks)
        self.content_blocks.append(block)
    
    def get_headings(self) -> List[Heading]:
        """Get all headings on this page"""
        return [block for block in self.content_blocks if isinstance(block, Heading)]
    
    def get_footnotes(self) -> List[Footnote]:
        """Get all footnotes on this page"""
        return [block for block in self.content_blocks if isinstance(block, Footnote)]
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert page to dictionary representation"""
        return {
            'page_number': self.page_number,
            'content_blocks': [block.to_dict() for block in self.content_blocks],
            'page_metadata': self.page_metadata
        }


@dataclass
class Document:
    """
    Represents the complete document with all pages and metadata.
    
    This is the top-level container that maintains the entire document
    structure and provides methods for document-wide operations.
    """
    title: str = ""
    pages: List[Page] = field(default_factory=list)
    document_metadata: Dict[str, Any] = field(default_factory=dict)
    toc_entries: List[Dict[str, Any]] = field(default_factory=list)
    
    def add_page(self, page: Page) -> None:
        """Add a page to the document"""
        self.pages.append(page)
    
    def get_all_headings(self) -> List[Heading]:
        """Get all headings from all pages for TOC generation"""
        headings = []
        for page in self.pages:
            headings.extend(page.get_headings())
        return headings
    
    def get_toc_entries(self) -> List[Dict[str, Any]]:
        """Get all TOC entries, preserving their structure"""
        return self.toc_entries
    
    def update_toc_entries(self, entries: List[Dict[str, Any]]) -> None:
        """Update TOC entries, preserving their structure"""
        self.toc_entries = entries
    
    def translate_toc_entries(self, translation_map: Dict[str, str]) -> None:
        """Translate TOC entries while preserving their structure"""
        for entry in self.toc_entries:
            if entry['text'] in translation_map:
                entry['text'] = translation_map[entry['text']]
    
    def get_all_content_blocks(self) -> List[ContentBlock]:
        """Get all content blocks from all pages in order"""
        blocks = []
        for page in self.pages:
            blocks.extend(page.content_blocks)
        return blocks
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert document to dictionary representation"""
        return {
            'title': self.title,
            'pages': [page.to_dict() for page in self.pages],
            'document_metadata': self.document_metadata,
            'toc_entries': self.toc_entries,
            'total_pages': len(self.pages),
            'total_blocks': len(self.get_all_content_blocks()),
            'total_headings': len(self.get_all_headings()),
            'total_toc_entries': len(self.toc_entries)
        }
    
    def get_statistics(self) -> Dict[str, int]:
        """Get document statistics"""
        stats = {
            'total_pages': len(self.pages),
            'total_blocks': 0,
            'headings': 0,
            'paragraphs': 0,
            'footnotes': 0,
            'tables': 0,
            'list_items': 0,
            'mathematical_formulas': 0
        }
        
        for page in self.pages:
            stats['total_blocks'] += len(page.content_blocks)
            for block in page.content_blocks:
                content_type = block.get_content_type().value
                if content_type in stats:
                    stats[content_type] += 1
        
        return stats


# Utility functions for document manipulation

def create_document_from_dict(data: Dict[str, Any]) -> Document:
    """Create a Document object from dictionary representation"""
    doc = Document(
        title=data.get('title', ''),
        document_metadata=data.get('document_metadata', {})
    )

    for page_data in data.get('pages', []):
        page = Page(
            page_number=page_data['page_number'],
            page_metadata=page_data.get('page_metadata', {})
        )

        for block_data in page_data.get('content_blocks', []):
            block = create_content_block_from_dict(block_data)
            if block:
                page.add_block(block)

        doc.add_page(page)

    return doc


def create_content_block_from_dict(data: Dict[str, Any]) -> Optional[ContentBlock]:
    """Create a ContentBlock object from dictionary representation"""
    block_type = data.get('type', '')

    # Common fields
    common_fields = {
        'content': data.get('content', ''),
        'page_num': data.get('page_num', 0),
        'position': data.get('position', 0),
        'bbox': tuple(data.get('bbox', (0, 0, 0, 0))),
        'font_info': data.get('font_info', {}),
        'block_id': data.get('block_id', '')
    }

    # Create specific block types
    if block_type == ContentType.HEADING.value:
        return Heading(level=data.get('level', 1), **common_fields)
    elif block_type == ContentType.PARAGRAPH.value:
        return Paragraph(**common_fields)
    elif block_type == ContentType.FOOTNOTE.value:
        return Footnote(
            reference_id=data.get('reference_id', ''),
            reference_text=data.get('reference_text', ''),
            **common_fields
        )
    elif block_type == ContentType.TABLE.value:
        return Table(
            rows=data.get('rows', 0),
            columns=data.get('columns', 0),
            headers=data.get('headers', []),
            **common_fields
        )
    elif block_type == ContentType.LIST_ITEM.value:
        return ListItem(
            list_level=data.get('list_level', 1),
            is_ordered=data.get('is_ordered', False),
            item_number=data.get('item_number'),
            **common_fields
        )
    elif block_type == ContentType.MATHEMATICAL_FORMULA.value:
        return MathematicalFormula(
            formula_type=data.get('formula_type', 'inline'),
            latex_representation=data.get('latex_representation', ''),
            **common_fields
        )
    else:
        logger.warning(f"Unknown content block type: {block_type}")
        return None


def merge_documents(doc1: Document, doc2: Document) -> Document:
    """Merge two documents into one"""
    merged = Document(
        title=f"{doc1.title} + {doc2.title}",
        document_metadata={**doc1.document_metadata, **doc2.document_metadata}
    )

    # Add all pages from both documents
    for page in doc1.pages:
        merged.add_page(page)

    # Adjust page numbers for second document
    page_offset = len(doc1.pages)
    for page in doc2.pages:
        new_page = Page(
            page_number=page.page_number + page_offset,
            page_metadata=page.page_metadata
        )
        for block in page.content_blocks:
            new_block = create_content_block_from_dict(block.to_dict())
            if new_block:
                new_block.page_num = new_page.page_number
                new_page.add_block(new_block)
        merged.add_page(new_page)

    return merged


class DocumentNode:
    def __init__(self, node_id: int, bbox: Tuple[int, int, int, int], class_label: str, confidence: float, text: Optional[str] = None, semantic_role: Optional[str] = None, extra: Optional[Dict[str, Any]] = None):
        self.node_id = node_id
        self.bbox = bbox  # (x1, y1, x2, y2)
        self.class_label = class_label
        self.confidence = confidence
        self.text = text
        self.semantic_role = semantic_role
        self.extra = extra or {}
        self.state = 'unlinked'  # or 'linked', 'validated', etc.

    def __repr__(self):
        return f"<DocumentNode id={self.node_id} class={self.class_label} conf={self.confidence:.2f} role={self.semantic_role} state={self.state}>"


class DocumentGraph:
    def __init__(self):
        self.graph = nx.DiGraph()
        self.node_counter = 0

    def add_node(self, bbox: Tuple[int, int, int, int], class_label: str, confidence: float, text: Optional[str] = None, semantic_role: Optional[str] = None, extra: Optional[Dict[str, Any]] = None) -> int:
        node = DocumentNode(
            node_id=self.node_counter,
            bbox=bbox,
            class_label=class_label,
            confidence=confidence,
            text=text,
            semantic_role=semantic_role,
            extra=extra
        )
        self.graph.add_node(self.node_counter, data=node)
        self.node_counter += 1
        return node.node_id

    def add_edge(self, from_id: int, to_id: int, relationship: str, confidence_weight: float = 1.0, extra: Optional[Dict[str, Any]] = None):
        self.graph.add_edge(from_id, to_id, relationship=relationship, confidence_weight=confidence_weight, extra=extra or {})

    def get_node(self, node_id: int) -> DocumentNode:
        return self.graph.nodes[node_id]['data']

    def get_neighbors(self, node_id: int, relationship: Optional[str] = None) -> List[int]:
        neighbors = []
        for succ in self.graph.successors(node_id):
            edge_data = self.graph.get_edge_data(node_id, succ)
            if relationship is None or edge_data.get('relationship') == relationship:
                neighbors.append(succ)
        return neighbors

    def nodes(self) -> List[DocumentNode]:
        return [self.graph.nodes[n]['data'] for n in self.graph.nodes]

    def edges(self) -> List[Tuple[int, int, Dict[str, Any]]]:
        return [(u, v, d) for u, v, d in self.graph.edges(data=True)]

    def __repr__(self):
        return f"<DocumentGraph nodes={len(self.graph.nodes)} edges={len(self.graph.edges)}>"


def add_yolo_detections_to_graph(graph, detections, page_num=None):
    """
    Add YOLO detection dicts to a DocumentGraph as nodes.
    Args:
        graph: DocumentGraph instance
        detections: List of dicts with keys 'label', 'confidence', 'bounding_box', optionally 'class_id', 'text', etc.
        page_num: Optional page number to add as extra info
    Returns:
        List of node IDs added to the graph
    """
    node_ids = []
    for det in detections:
        extra = {}
        if 'class_id' in det:
            extra['class_id'] = det['class_id']
        if page_num is not None:
            extra['page_num'] = page_num
        node_id = graph.add_node(
            bbox=tuple(det['bounding_box']),
            class_label=det['label'],
            confidence=det['confidence'],
            text=det.get('text'),
            semantic_role=None,
            extra=extra
        )
        node_ids.append(node_id)
    return node_ids


def add_ocr_text_regions_to_graph(graph, text_regions, page_num=None):
    """
    Add OCR text regions to a DocumentGraph as nodes.
    Args:
        graph: DocumentGraph instance
        text_regions: List of dicts with keys 'text', 'confidence', 'bbox' (left, top, width, height)
        page_num: Optional page number to add as extra info
    Returns:
        List of node IDs added to the graph
    """
    node_ids = []
    for region in text_regions:
        # Convert (left, top, width, height) to (x1, y1, x2, y2)
        left, top, width, height = region['bbox']
        bbox = (left, top, left + width, top + height)
        
        extra = {}
        if page_num is not None:
            extra['page_num'] = page_num
        extra['ocr_confidence'] = region['confidence']
        
        node_id = graph.add_node(
            bbox=bbox,
            class_label='text',  # OCR regions are always text
            confidence=region['confidence'] / 100.0,  # Convert to 0-1 scale
            text=region['text'],
            semantic_role=None,
            extra=extra
        )
        node_ids.append(node_id)
    return node_ids


class RegionTextAssociationMatrix:
    """
    Computes and manages association scores between YOLO regions and OCR text blocks.
    This is the foundation for the confidence-driven region-text mapping.
    """
    
    def __init__(self, yolo_nodes, ocr_nodes):
        """
        Initialize the association matrix.
        
        Args:
            yolo_nodes: List of DocumentNode objects from YOLO detections
            ocr_nodes: List of DocumentNode objects from OCR text regions
        """
        self.yolo_nodes = yolo_nodes
        self.ocr_nodes = ocr_nodes
        self.matrix = np.zeros((len(yolo_nodes), len(ocr_nodes)))
        self.compute_association_scores()
    
    def compute_association_scores(self):
        """Compute association scores for all region-text pairs"""
        for i, yolo_node in enumerate(self.yolo_nodes):
            for j, ocr_node in enumerate(self.ocr_nodes):
                score = self._compute_single_association_score(yolo_node, ocr_node)
                self.matrix[i, j] = score
    
    def _compute_single_association_score(self, yolo_node, ocr_node):
        """
        Compute association score between a YOLO region and OCR text block.
        
        Returns:
            float: Association score between 0 and 1
        """
        # 1. Spatial overlap (IoU)
        iou_score = self._compute_iou(yolo_node.bbox, ocr_node.bbox)
        
        # 2. Class compatibility score
        class_score = self._compute_class_compatibility(yolo_node.class_label, ocr_node.class_label)
        
        # 3. Confidence product
        confidence_score = yolo_node.confidence * ocr_node.confidence
        
        # 4. Text content relevance (if text is available)
        content_score = self._compute_content_relevance(yolo_node, ocr_node)
        
        # Combine scores with weights
        final_score = (
            0.4 * iou_score +
            0.3 * class_score +
            0.2 * confidence_score +
            0.1 * content_score
        )
        
        return min(1.0, max(0.0, final_score))
    
    def _compute_iou(self, bbox1, bbox2):
        """Compute Intersection over Union between two bounding boxes"""
        x1_1, y1_1, x2_1, y2_1 = bbox1
        x1_2, y1_2, x2_2, y2_2 = bbox2
        
        # Compute intersection
        x1_i = max(x1_1, x1_2)
        y1_i = max(y1_1, y1_2)
        x2_i = min(x2_1, x2_2)
        y2_i = min(y2_1, y2_2)
        
        if x2_i <= x1_i or y2_i <= y1_i:
            return 0.0
        
        intersection = (x2_i - x1_i) * (y2_i - y1_i)
        
        # Compute union
        area1 = (x2_1 - x1_1) * (y2_1 - y1_1)
        area2 = (x2_2 - x1_2) * (y2_2 - y1_2)
        union = area1 + area2 - intersection
        
        return intersection / union if union > 0 else 0.0
    
    def _compute_class_compatibility(self, yolo_class, ocr_class):
        """Compute compatibility score between YOLO and OCR classes"""
        # OCR nodes are always 'text', so we need to check if YOLO class expects text
        text_compatible_classes = ['text', 'title', 'paragraph', 'list', 'caption', 'quote', 'footnote']
        
        if yolo_class in text_compatible_classes:
            return 1.0
        elif yolo_class in ['figure', 'table', 'equation']:
            return 0.1  # Low compatibility for visual elements
        else:
            return 0.5  # Medium compatibility for other classes
    
    def _compute_content_relevance(self, yolo_node, ocr_node):
        """Compute content relevance score based on text content"""
        if not ocr_node.text or not yolo_node.text:
            return 0.5  # Neutral score if no text available
        
        # Simple text similarity (can be enhanced with more sophisticated NLP)
        yolo_text = yolo_node.text.lower()
        ocr_text = ocr_node.text.lower()
        
        # Check if OCR text contains YOLO text or vice versa
        if yolo_text in ocr_text or ocr_text in yolo_text:
            return 1.0
        elif any(word in ocr_text for word in yolo_text.split()):
            return 0.7
        else:
            return 0.3
    
    def get_best_associations(self, threshold=0.5):
        """
        Get the best region-text associations above a threshold.
        
        Args:
            threshold: Minimum association score
            
        Returns:
            List of tuples: (yolo_node_id, ocr_node_id, score)
        """
        associations = []
        for i in range(len(self.yolo_nodes)):
            for j in range(len(self.ocr_nodes)):
                score = self.matrix[i, j]
                if score >= threshold:
                    associations.append((i, j, score))
        
        # Sort by score (highest first)
        associations.sort(key=lambda x: x[2], reverse=True)
        return associations
    
    def get_association_score(self, yolo_node_id, ocr_node_id):
        """Get association score for specific node pair"""
        return self.matrix[yolo_node_id, ocr_node_id]
    
    def print_matrix_summary(self):
        """Print a summary of the association matrix"""
        print(f"Association Matrix Summary:")
        print(f"  YOLO regions: {len(self.yolo_nodes)}")
        print(f"  OCR text blocks: {len(self.ocr_nodes)}")
        print(f"  Total associations: {len(self.yolo_nodes) * len(self.ocr_nodes)}")
        
        # Count strong associations
        strong_associations = sum(1 for score in self.matrix.flatten() if score >= 0.7)
        medium_associations = sum(1 for score in self.matrix.flatten() if 0.3 <= score < 0.7)
        
        print(f"  Strong associations (≥0.7): {strong_associations}")
        print(f"  Medium associations (0.3-0.7): {medium_associations}")
        
        if len(self.matrix) > 0:
            print(f"  Average association score: {np.mean(self.matrix):.3f}")
            print(f"  Max association score: {np.max(self.matrix):.3f}")


def build_association_matrix_from_graph(graph):
    """
    Build association matrix from a DocumentGraph containing both YOLO and OCR nodes.
    
    Args:
        graph: DocumentGraph with mixed YOLO and OCR nodes
        
    Returns:
        RegionTextAssociationMatrix
    """
    yolo_nodes = [node for node in graph.nodes() if node.extra.get('class_id') is not None]
    ocr_nodes = [node for node in graph.nodes() if node.extra.get('ocr_confidence') is not None]
    
    return RegionTextAssociationMatrix(yolo_nodes, ocr_nodes)


class DocumentGraphPopulator:
    """
    Populates a DocumentGraph with relationships and assigns text to regions
    using the association matrix and spatial/semantic heuristics.
    """
    
    def __init__(self, graph, association_matrix):
        """
        Initialize the graph populator.
        
        Args:
            graph: DocumentGraph instance
            association_matrix: RegionTextAssociationMatrix instance
        """
        self.graph = graph
        self.association_matrix = association_matrix
        self.yolo_nodes = association_matrix.yolo_nodes
        self.ocr_nodes = association_matrix.ocr_nodes
        
    def populate_graph(self, association_threshold=0.5):
        """
        Main method to populate the graph with relationships and text assignments.
        
        Args:
            association_threshold: Minimum association score for text assignment
        """
        # Step 1: Assign text to YOLO regions using association matrix
        self._assign_text_to_regions(association_threshold)
        
        # Step 2: Establish spatial relationships
        self._establish_spatial_relationships()
        
        # Step 3: Establish semantic relationships
        self._establish_semantic_relationships()
        
        # Step 4: Infer reading order
        self._infer_reading_order()
        
        # Step 5: Update node states
        self._update_node_states()
    
    def _assign_text_to_regions(self, threshold):
        """Assign OCR text to YOLO regions based on association scores"""
        best_associations = self.association_matrix.get_best_associations(threshold)
        
        # Track which OCR nodes have been assigned
        assigned_ocr = set()
        
        for yolo_idx, ocr_idx, score in best_associations:
            yolo_node = self.yolo_nodes[yolo_idx]
            ocr_node = self.ocr_nodes[ocr_idx]
            
            # Skip if OCR node already assigned to a better match
            if ocr_idx in assigned_ocr:
                continue
            
            # Assign text to YOLO region
            if not yolo_node.text:
                yolo_node.text = ocr_node.text
                yolo_node.extra['assigned_ocr_id'] = ocr_idx
                yolo_node.extra['association_score'] = score
                assigned_ocr.add(ocr_idx)
                
                logger.debug(f"Assigned text '{ocr_node.text}' to {yolo_node.class_label} (score: {score:.3f})")
    
    def _establish_spatial_relationships(self):
        """Establish spatial relationships between nodes (above, below, left, right)"""
        nodes = self.graph.nodes()
        
        for i, node1 in enumerate(nodes):
            for j, node2 in enumerate(nodes):
                if i == j:
                    continue
                
                relationship = self._compute_spatial_relationship(node1, node2)
                if relationship:
                    confidence = self._compute_spatial_confidence(node1, node2, relationship)
                    self.graph.add_edge(
                        node1.node_id, 
                        node2.node_id, 
                        relationship=relationship,
                        confidence_weight=confidence
                    )
    
    def _compute_spatial_relationship(self, node1, node2):
        """Compute spatial relationship between two nodes"""
        x1_1, y1_1, x2_1, y2_1 = node1.bbox
        x1_2, y1_2, x2_2, y2_2 = node2.bbox
        
        # Calculate centers
        center1_x = (x1_1 + x2_1) / 2
        center1_y = (y1_1 + y2_1) / 2
        center2_x = (x1_2 + x2_2) / 2
        center2_y = (y1_2 + y2_2) / 2
        
        # Vertical relationships
        if center2_y > center1_y + 20:  # Below (with some tolerance)
            return 'is_below'
        elif center1_y > center2_y + 20:  # Above
            return 'is_above'
        
        # Horizontal relationships
        elif center2_x > center1_x + 20:  # Right
            return 'is_right_of'
        elif center1_x > center2_x + 20:  # Left
            return 'is_left_of'
        
        return None
    
    def _compute_spatial_confidence(self, node1, node2, relationship):
        """Compute confidence for spatial relationship"""
        # Base confidence on distance and overlap
        x1_1, y1_1, x2_1, y2_1 = node1.bbox
        x1_2, y1_2, x2_2, y2_2 = node2.bbox
        
        # Calculate distance between centers
        center1_x = (x1_1 + x2_1) / 2
        center1_y = (y1_1 + y2_1) / 2
        center2_x = (x1_2 + x2_2) / 2
        center2_y = (y1_2 + y2_2) / 2
        
        distance = ((center2_x - center1_x) ** 2 + (center2_y - center1_y) ** 2) ** 0.5
        
        # Higher confidence for closer elements
        max_distance = 500  # Adjust based on typical document layout
        confidence = max(0.1, 1.0 - (distance / max_distance))
        
        return confidence
    
    def _establish_semantic_relationships(self):
        """Establish semantic relationships based on class labels and content"""
        nodes = self.graph.nodes()
        
        for i, node1 in enumerate(nodes):
            for j, node2 in enumerate(nodes):
                if i == j:
                    continue
                
                relationship = self._compute_semantic_relationship(node1, node2)
                if relationship:
                    confidence = self._compute_semantic_confidence(node1, node2, relationship)
                    self.graph.add_edge(
                        node1.node_id,
                        node2.node_id,
                        relationship=relationship,
                        confidence_weight=confidence
                    )
    
    def _compute_semantic_relationship(self, node1, node2):
        """Compute semantic relationship between two nodes"""
        class1 = node1.class_label
        class2 = node2.class_label
        
        # Title -> Paragraph relationships
        if class1 == 'title' and class2 in ['paragraph', 'text']:
            return 'has_content'
        elif class2 == 'title' and class1 in ['paragraph', 'text']:
            return 'belongs_to'
        
        # Figure -> Caption relationships
        if class1 == 'figure' and class2 == 'caption':
            return 'has_caption'
        elif class2 == 'figure' and class1 == 'caption':
            return 'captions'
        
        # List relationships
        if class1 == 'list' and class2 in ['text', 'paragraph']:
            return 'contains_item'
        elif class2 == 'list' and class1 in ['text', 'paragraph']:
            return 'is_list_item'
        
        return None
    
    def _compute_semantic_confidence(self, node1, node2, relationship):
        """Compute confidence for semantic relationship"""
        # Base confidence on class compatibility and spatial proximity
        base_confidence = 0.8
        
        # Boost confidence if nodes are spatially close
        spatial_relationship = self._compute_spatial_relationship(node1, node2)
        if spatial_relationship in ['is_below', 'is_right_of']:
            base_confidence += 0.1
        
        return min(1.0, base_confidence)
    
    def _infer_reading_order(self):
        """Infer reading order by spatial positioning (no edges created to avoid cycles)"""
        # Reading order is computed on-demand in get_reading_order()
        # This method is kept for future enhancements but doesn't create edges
        pass
    
    def _update_node_states(self):
        """Update node states based on assignments and relationships"""
        for node in self.graph.nodes():
            if node.text:
                node.state = 'linked'
            else:
                node.state = 'unlinked'
    
    def get_reading_order(self):
        """Get nodes in reading order"""
        # Use spatial ordering as primary method (more reliable for document layout)
        nodes = self.graph.nodes()
        sorted_nodes = sorted(nodes, key=lambda n: (n.bbox[1], n.bbox[0]))
        
        # Try topological sort as fallback for simple cases
        try:
            # Create a simplified graph with only reading order edges
            reading_graph = nx.DiGraph()
            for node in nodes:
                reading_graph.add_node(node.node_id)
            
            # Add only reading order edges
            for _, _, edge_data in self.graph.edges():
                if edge_data.get('relationship') == 'precedes_in_reading':
                    # Add edge if it doesn't create a cycle
                    reading_graph.add_edge(edge_data.get('from'), edge_data.get('to'))
            
            if nx.is_directed_acyclic_graph(reading_graph):
                reading_order = list(nx.topological_sort(reading_graph))
                return [self.graph.get_node(node_id) for node_id in reading_order]
        except:
            pass
        
        # Fall back to spatial ordering
        return sorted_nodes
    
    def print_graph_summary(self):
        """Print a summary of the populated graph"""
        nodes = self.graph.nodes()
        edges = self.graph.edges()
        
        print(f"Populated Graph Summary:")
        print(f"  Total nodes: {len(nodes)}")
        print(f"  Total edges: {len(edges)}")
        
        # Count nodes by state
        linked_nodes = sum(1 for n in nodes if n.state == 'linked')
        unlinked_nodes = sum(1 for n in nodes if n.state == 'unlinked')
        
        print(f"  Linked nodes: {linked_nodes}")
        print(f"  Unlinked nodes: {unlinked_nodes}")
        
        # Count edges by relationship type
        relationship_counts = {}
        for _, _, edge_data in edges:
            rel_type = edge_data.get('relationship', 'unknown')
            relationship_counts[rel_type] = relationship_counts.get(rel_type, 0) + 1
        
        print(f"  Edge types:")
        for rel_type, count in relationship_counts.items():
            print(f"    {rel_type}: {count}")


def populate_document_graph(graph, association_matrix, threshold=0.5):
    """
    Convenience function to populate a document graph.
    
    Args:
        graph: DocumentGraph instance
        association_matrix: RegionTextAssociationMatrix instance
        threshold: Association threshold for text assignment
        
    Returns:
        DocumentGraphPopulator instance
    """
    populator = DocumentGraphPopulator(graph, association_matrix)
    populator.populate_graph(threshold)
    return populator


class DocumentGraphRefiner:
    """
    Implements feedback loop and refinement for the document graph.
    Uses cross-validation, confidence propagation, and specialized solvers
    to improve the accuracy and robustness of the graph structure.
    """
    
    def __init__(self, graph, association_matrix):
        """
        Initialize the graph refiner.
        
        Args:
            graph: DocumentGraph instance
            association_matrix: RegionTextAssociationMatrix instance
        """
        self.graph = graph
        self.association_matrix = association_matrix
        self.yolo_nodes = association_matrix.yolo_nodes
        self.ocr_nodes = association_matrix.ocr_nodes
        self.refinement_history = []
        self.max_iterations = 5
        
    def refine_graph(self):
        """
        Main refinement loop that iteratively improves the graph until stable.
        
        Returns:
            bool: True if refinements were made, False if graph is stable
        """
        logger.info("🔄 Starting graph refinement process...")
        
        for iteration in range(self.max_iterations):
            logger.info(f"  Iteration {iteration + 1}/{self.max_iterations}")
            
            refinements_made = False
            
            # Step 1: Cross-validate visual and textual cues
            if self._cross_validate_visual_textual():
                refinements_made = True
            
            # Step 2: Propagate confidence from high-certainty nodes
            if self._propagate_confidence():
                refinements_made = True
            
            # Step 3: Apply specialized solvers
            if self._apply_specialized_solvers():
                refinements_made = True
            
            # Step 4: Validate and correct semantic roles
            if self._validate_semantic_roles():
                refinements_made = True
            
            # Step 5: Resolve conflicts and ambiguities
            if self._resolve_conflicts():
                refinements_made = True
            
            # Record iteration results
            self.refinement_history.append({
                'iteration': iteration + 1,
                'refinements_made': refinements_made,
                'total_nodes': len(self.graph.nodes()),
                'linked_nodes': sum(1 for n in self.graph.nodes() if n.state == 'linked')
            })
            
            # If no refinements made, graph is stable
            if not refinements_made:
                logger.info(f"✅ Graph stabilized after {iteration + 1} iterations")
                return True
        
        logger.warning(f"⚠️ Graph refinement stopped after {self.max_iterations} iterations")
        return False
    
    def _cross_validate_visual_textual(self):
        """Cross-validate visual and textual cues to correct misclassifications"""
        refinements_made = False
        
        for yolo_node in self.yolo_nodes:
            if not yolo_node.text:
                continue
            
            # Check if text content suggests a different class
            suggested_class = self._infer_class_from_text(yolo_node.text)
            if suggested_class and suggested_class != yolo_node.class_label:
                # Only change if confidence is low or text strongly suggests different class
                if (yolo_node.confidence < 0.7 or 
                    self._text_class_confidence(yolo_node.text, suggested_class) > 0.8):
                    
                    old_class = yolo_node.class_label
                    yolo_node.class_label = suggested_class
                    yolo_node.extra['corrected_class'] = old_class
                    yolo_node.extra['correction_confidence'] = self._text_class_confidence(yolo_node.text, suggested_class)
                    
                    logger.debug(f"Corrected class: {old_class} -> {suggested_class} for text: '{yolo_node.text[:50]}...'")
                    refinements_made = True
        
        return refinements_made
    
    def _infer_class_from_text(self, text):
        """Infer document element class from text content"""
        text_lower = text.lower().strip()
        
        # Title indicators
        title_indicators = ['chapter', 'section', 'introduction', 'conclusion', 'abstract', 'summary']
        if any(indicator in text_lower for indicator in title_indicators):
            return 'title'
        
        # Figure caption indicators
        caption_indicators = ['figure', 'fig.', 'image', 'photo', 'diagram', 'chart']
        if any(indicator in text_lower for indicator in caption_indicators):
            return 'caption'
        
        # List indicators
        list_indicators = ['•', '1.', '2.', '3.', 'a)', 'b)', 'c)', '- ']
        if any(text.startswith(indicator) for indicator in list_indicators):
            return 'list'
        
        # Table indicators
        table_indicators = ['table', 'tab.', 'row', 'column', 'header']
        if any(indicator in text_lower for indicator in table_indicators):
            return 'table'
        
        # Quote indicators
        quote_indicators = ['"', '"', ''', ''', 'quoted', 'said']
        if any(indicator in text for indicator in quote_indicators):
            return 'quote'
        
        return None
    
    def _text_class_confidence(self, text, suggested_class):
        """Compute confidence for text-based class suggestion"""
        text_lower = text.lower().strip()
        
        if suggested_class == 'title':
            # Short, capitalized text suggests title
            if len(text) < 100 and text[0].isupper() and not text.endswith('.'):
                return 0.9
            return 0.6
        
        elif suggested_class == 'caption':
            # Contains figure/table references
            if any(word in text_lower for word in ['figure', 'fig', 'table', 'tab']):
                return 0.8
            return 0.5
        
        elif suggested_class == 'list':
            # Starts with list markers
            if any(text.startswith(marker) for marker in ['•', '1.', '2.', '- ']):
                return 0.9
            return 0.4
        
        return 0.3
    
    def _propagate_confidence(self):
        """Propagate confidence from high-certainty nodes to neighbors"""
        refinements_made = False
        
        # Find high-confidence anchor nodes
        anchor_nodes = [node for node in self.graph.nodes() 
                       if node.confidence > 0.9 and node.state == 'linked']
        
        for anchor in anchor_nodes:
            # Find neighboring nodes
            neighbors = self.graph.get_neighbors(anchor.node_id)
            
            for neighbor_id in neighbors:
                neighbor = self.graph.get_node(neighbor_id)
                
                # Boost confidence of neighbors based on relationship strength
                edge_data = self.graph.graph.get_edge_data(anchor.node_id, neighbor_id)
                if edge_data and edge_data.get('confidence_weight', 0) > 0.7:
                    # Boost neighbor confidence
                    confidence_boost = anchor.confidence * edge_data['confidence_weight'] * 0.1
                    neighbor.confidence = min(1.0, neighbor.confidence + confidence_boost)
                    
                    # If neighbor was unlinked but now has high confidence, try to link it
                    if (neighbor.state == 'unlinked' and 
                        neighbor.confidence > 0.6 and 
                        not neighbor.text):
                        
                        # Try to find matching OCR text
                        best_match = self._find_best_ocr_match(neighbor)
                        if best_match:
                            neighbor.text = best_match['text']
                            neighbor.state = 'linked'
                            neighbor.extra['propagated_from'] = anchor.node_id
                            refinements_made = True
        
        return refinements_made
    
    def _find_best_ocr_match(self, yolo_node):
        """Find the best OCR text match for a YOLO node"""
        best_match = None
        best_score = 0
        
        for ocr_node in self.ocr_nodes:
            if ocr_node.text and not ocr_node.extra.get('assigned_to'):
                # Compute spatial overlap
                iou = self._compute_iou(yolo_node.bbox, ocr_node.bbox)
                if iou > 0.3:  # Minimum overlap threshold
                    score = iou * ocr_node.confidence
                    if score > best_score:
                        best_score = score
                        best_match = {
                            'text': ocr_node.text,
                            'score': score,
                            'ocr_node_id': ocr_node.node_id
                        }
        
        return best_match if best_score > 0.5 else None
    
    def _apply_specialized_solvers(self):
        """Apply specialized solvers for complex regions (tables, figures, etc.)"""
        refinements_made = False
        
        for node in self.graph.nodes():
            if node.class_label == 'table':
                if self._solve_table_structure(node):
                    refinements_made = True
            
            elif node.class_label == 'figure':
                if self._solve_figure_caption(node):
                    refinements_made = True
            
            elif node.class_label == 'list':
                if self._solve_list_structure(node):
                    refinements_made = True
        
        return refinements_made
    
    def _solve_table_structure(self, table_node):
        """Specialized solver for table structure"""
        # This would implement table-specific logic
        # For now, just mark as processed
        table_node.extra['table_solver_applied'] = True
        return False  # No refinements made in this simple implementation
    
    def _solve_figure_caption(self, figure_node):
        """Specialized solver for figure-caption relationships"""
        # Look for caption nodes near the figure
        figure_bbox = figure_node.bbox
        caption_candidates = []
        
        for node in self.graph.nodes():
            if (node.class_label == 'caption' and 
                self._is_near(figure_bbox, node.bbox)):
                caption_candidates.append(node)
        
        if caption_candidates:
            # Find the closest caption
            closest_caption = min(caption_candidates, 
                                key=lambda c: self._distance(figure_bbox, c.bbox))
            
            # Create relationship
            self.graph.add_edge(
                figure_node.node_id,
                closest_caption.node_id,
                relationship='has_caption',
                confidence_weight=0.9
            )
            
            return True
        
        return False
    
    def _solve_list_structure(self, list_node):
        """Specialized solver for list structure"""
        # Look for list items near the list node
        list_bbox = list_node.bbox
        item_candidates = []
        
        for node in self.graph.nodes():
            if (node.class_label in ['text', 'paragraph'] and 
                self._is_near(list_bbox, node.bbox) and
                self._looks_like_list_item(node.text)):
                item_candidates.append(node)
        
        if item_candidates:
            # Create relationships to list items
            for item in item_candidates:
                self.graph.add_edge(
                    list_node.node_id,
                    item.node_id,
                    relationship='contains_item',
                    confidence_weight=0.8
                )
            
            return True
        
        return False
    
    def _looks_like_list_item(self, text):
        """Check if text looks like a list item"""
        if not text:
            return False
        
        text_stripped = text.strip()
        list_patterns = [r'^\s*[•\-\*]\s+', r'^\s*\d+\.\s+', r'^\s*[a-zA-Z]\.\s+']
        
        import re
        return any(re.match(pattern, text_stripped) for pattern in list_patterns)
    
    def _validate_semantic_roles(self):
        """Validate and correct semantic roles based on context"""
        refinements_made = False
        
        for node in self.graph.nodes():
            if node.semantic_role is None:
                # Infer semantic role based on class and context
                inferred_role = self._infer_semantic_role(node)
                if inferred_role:
                    node.semantic_role = inferred_role
                    refinements_made = True
        
        return refinements_made
    
    def _infer_semantic_role(self, node):
        """Infer semantic role for a node based on class and context"""
        class_label = node.class_label
        
        if class_label == 'title':
            # Determine heading level based on text length and content
            if node.text:
                if len(node.text) < 50:
                    return 'h1'
                elif len(node.text) < 100:
                    return 'h2'
                else:
                    return 'h3'
        
        elif class_label == 'paragraph':
            return 'p'
        
        elif class_label == 'list':
            return 'ul'  # or 'ol' based on content
        
        elif class_label == 'table':
            return 'table'
        
        elif class_label == 'figure':
            return 'figure'
        
        elif class_label == 'caption':
            return 'figcaption'
        
        return None
    
    def _resolve_conflicts(self):
        """Resolve conflicts and ambiguities in the graph"""
        refinements_made = False
        
        # Resolve overlapping regions
        nodes = self.graph.nodes()
        for i, node1 in enumerate(nodes):
            for j, node2 in enumerate(nodes[i+1:], i+1):
                if self._has_significant_overlap(node1, node2):
                    if self._resolve_overlap(node1, node2):
                        refinements_made = True
        
        return refinements_made
    
    def _has_significant_overlap(self, node1, node2):
        """Check if two nodes have significant spatial overlap"""
        iou = self._compute_iou(node1.bbox, node2.bbox)
        return iou > 0.5  # Significant overlap threshold
    
    def _resolve_overlap(self, node1, node2):
        """Resolve overlap between two nodes"""
        # Keep the node with higher confidence
        if node1.confidence > node2.confidence:
            # Remove node2 or merge its content into node1
            if node2.text and not node1.text:
                node1.text = node2.text
                node1.state = 'linked'
            return True
        else:
            # Remove node1 or merge its content into node2
            if node1.text and not node2.text:
                node2.text = node1.text
                node2.state = 'linked'
            return True
    
    def _compute_iou(self, bbox1, bbox2):
        """Compute Intersection over Union between two bounding boxes"""
        x1_1, y1_1, x2_1, y2_1 = bbox1
        x1_2, y1_2, x2_2, y2_2 = bbox2
        
        # Compute intersection
        x1_i = max(x1_1, x1_2)
        y1_i = max(y1_1, y1_2)
        x2_i = min(x2_1, x2_2)
        y2_i = min(y2_1, y2_2)
        
        if x2_i <= x1_i or y2_i <= y1_i:
            return 0.0
        
        intersection = (x2_i - x1_i) * (y2_i - y1_i)
        
        # Compute union
        area1 = (x2_1 - x1_1) * (y2_1 - y1_1)
        area2 = (x2_2 - x1_2) * (y2_2 - y1_2)
        union = area1 + area2 - intersection
        
        return intersection / union if union > 0 else 0.0
    
    def _is_near(self, bbox1, bbox2, threshold=50):
        """Check if two bounding boxes are near each other"""
        x1_1, y1_1, x2_1, y2_1 = bbox1
        x1_2, y1_2, x2_2, y2_2 = bbox2
        
        center1_x = (x1_1 + x2_1) / 2
        center1_y = (y1_1 + y2_1) / 2
        center2_x = (x1_2 + x2_2) / 2
        center2_y = (y1_2 + y2_2) / 2
        
        distance = ((center2_x - center1_x) ** 2 + (center2_y - center1_y) ** 2) ** 0.5
        return distance < threshold
    
    def _distance(self, bbox1, bbox2):
        """Compute distance between centers of two bounding boxes"""
        x1_1, y1_1, x2_1, y2_1 = bbox1
        x1_2, y1_2, x2_2, y2_2 = bbox2
        
        center1_x = (x1_1 + x2_1) / 2
        center1_y = (y1_1 + y2_1) / 2
        center2_x = (x1_2 + x2_2) / 2
        center2_y = (y1_2 + y2_2) / 2
        
        return ((center2_x - center1_x) ** 2 + (center2_y - center1_y) ** 2) ** 0.5
    
    def get_refinement_summary(self):
        """Get a summary of the refinement process"""
        return {
            'total_iterations': len(self.refinement_history),
            'refinements_made': any(r['refinements_made'] for r in self.refinement_history),
            'final_nodes': len(self.graph.nodes()),
            'final_linked_nodes': sum(1 for n in self.graph.nodes() if n.state == 'linked'),
            'history': self.refinement_history
        }
    
    def print_refinement_summary(self):
        """Print a summary of the refinement process"""
        summary = self.get_refinement_summary()
        
        print(f"Graph Refinement Summary:")
        print(f"  Total iterations: {summary['total_iterations']}")
        print(f"  Refinements made: {summary['refinements_made']}")
        print(f"  Final nodes: {summary['final_nodes']}")
        print(f"  Final linked nodes: {summary['final_linked_nodes']}")
        
        if summary['history']:
            print(f"  Iteration details:")
            for record in summary['history']:
                print(f"    Iteration {record['iteration']}: {record['refinements_made']} refinements, {record['linked_nodes']}/{record['total_nodes']} linked")


def refine_document_graph(graph, association_matrix):
    """
    Convenience function to refine a document graph.
    
    Args:
        graph: DocumentGraph instance
        association_matrix: RegionTextAssociationMatrix instance
        
    Returns:
        DocumentGraphRefiner instance
    """
    refiner = DocumentGraphRefiner(graph, association_matrix)
    refiner.refine_graph()
    return refiner


class DocumentReconstructor:
    """
    Reconstructs documents from the refined graph with proper structure preservation.
    Generates output in various formats (HTML, Markdown, Word) with translated content.
    Enhanced with translation confidence scoring and format preservation.
    """
    
    def __init__(self, graph, translation_service=None):
        """
        Initialize the document reconstructor.
        
        Args:
            graph: Refined DocumentGraph instance
            translation_service: Enhanced translation service for content translation
        """
        self.graph = graph
        self.translation_service = translation_service
        self.reading_order = None
        self.structured_content = []
        self.translation_confidence_scores = {}
        self.original_translated_mapping = {}
        self.format_preservation_map = {}
        
    def reconstruct_document(self, output_format='html', target_language='en', preserve_layout=True):
        """
        Reconstruct the document in the specified format with enhanced translation.
        
        Args:
            output_format: 'html', 'markdown', 'word', or 'json'
            target_language: Target language for translation
            preserve_layout: Whether to preserve original layout structure
            
        Returns:
            str or dict: Reconstructed document content
        """
        logger.info(f"🔨 Reconstructing document in {output_format} format...")
        
        # Get reading order from graph
        self.reading_order = self._get_reading_order()
        
        # Extract and structure content
        self.structured_content = self._extract_structured_content()
        
        # Translate content if needed with enhanced features
        if self.translation_service and target_language != 'en':
            self._translate_content_enhanced(target_language)
        
        # Generate output in requested format
        if output_format == 'html':
            return self._generate_html(preserve_layout)
        elif output_format == 'markdown':
            return self._generate_markdown(preserve_layout)
        elif output_format == 'word':
            return self._generate_word_document(preserve_layout)
        elif output_format == 'json':
            return self._generate_json_structure()
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    def _translate_content_enhanced(self, target_language):
        """Enhanced translation with confidence scoring and format preservation"""
        logger.info(f"🌐 Translating content to {target_language} with enhanced features...")
        
        for i, block in enumerate(self.structured_content):
            if block['text']:
                try:
                    # Preserve format markers before translation
                    preserved_text, format_map = self._preserve_format_markers(block['text'])
                    
                    # Get context for better translation
                    context = self._get_translation_context(i)
                    
                    # Use enhanced translation service
                    if hasattr(self.translation_service, 'translate_text_enhanced'):
                        translated_text = self.translation_service.translate_text_enhanced(
                            preserved_text,
                            target_language=target_language,
                            prev_context=context.get('prev', ''),
                            next_context=context.get('next', ''),
                            item_type=block['type']
                        )
                    else:
                        # Fallback to standard translation
                        translated_text = self.translation_service.translate_text(
                            preserved_text, 
                            target_language=target_language
                        )
                    
                    # Restore format markers
                    translated_text = self._restore_format_markers(translated_text, format_map)
                    
                    # Calculate translation confidence
                    confidence_score = self._calculate_translation_confidence(
                        block['text'], translated_text, block['type']
                    )
                    
                    # Store translation data
                    block['translated_text'] = translated_text
                    block['original_text'] = block['text']
                    block['text'] = translated_text
                    block['translation_confidence'] = confidence_score
                    block['format_preserved'] = True
                    
                    # Update mappings
                    self.translation_confidence_scores[block['node_id']] = confidence_score
                    self.original_translated_mapping[block['original_text']] = {
                        'translated': translated_text,
                        'confidence': confidence_score,
                        'format_map': format_map
                    }
                    
                except Exception as e:
                    logger.warning(f"Translation failed for text: {block['text'][:50]}... Error: {e}")
                    block['translated_text'] = block['text']  # Keep original
                    block['translation_confidence'] = 0.0
                    block['format_preserved'] = False
    
    def _preserve_format_markers(self, text):
        """Preserve formatting markers by replacing them with unique placeholders"""
        format_map = {}
        preserved_text = text
        
        # Define format markers to preserve
        format_patterns = {
            'bold': r'\*\*(.*?)\*\*',
            'italic': r'\*(.*?)\*',
            'code': r'`(.*?)`',
            'link': r'\[([^\]]+)\]\(([^)]+)\)',
            'list_item': r'^[\s]*[-*+]\s+',
            'numbered_list': r'^[\s]*\d+\.\s+',
            'heading': r'^#{1,6}\s+',
            'blockquote': r'^>\s+',
            'horizontal_rule': r'^---$',
            'table_separator': r'^\|[\s\-\|:]+\|$'
        }
        
        import re
        for format_type, pattern in format_patterns.items():
            matches = re.finditer(pattern, preserved_text, re.MULTILINE)
            for j, match in enumerate(matches):
                marker = match.group()
                placeholder = f"__FORMAT_{format_type}_{j}__"
                preserved_text = preserved_text.replace(marker, placeholder)
                format_map[placeholder] = {
                    'marker': marker,
                    'type': format_type,
                    'position': match.start()
                }
        
        return preserved_text, format_map
    
    def _restore_format_markers(self, text, format_map):
        """Restore formatting markers from placeholders"""
        restored_text = text
        
        # Sort placeholders by their original position to maintain order
        sorted_placeholders = sorted(
            format_map.items(),
            key=lambda x: x[1]['position']
        )
        
        for placeholder, info in sorted_placeholders:
            restored_text = restored_text.replace(placeholder, info['marker'])
            
        return restored_text
    
    def _get_translation_context(self, block_index):
        """Get context from surrounding blocks for better translation"""
        context = {'prev': '', 'next': ''}
        
        if block_index > 0:
            prev_block = self.structured_content[block_index - 1]
            context['prev'] = prev_block.get('text', '')[:200]  # Last 200 chars
        
        if block_index < len(self.structured_content) - 1:
            next_block = self.structured_content[block_index + 1]
            context['next'] = next_block.get('text', '')[:200]  # First 200 chars
        
        return context
    
    def _calculate_translation_confidence(self, original_text, translated_text, block_type):
        """Calculate confidence score for translation quality"""
        try:
            # Base confidence on text length similarity
            length_ratio = len(translated_text) / max(len(original_text), 1)
            
            # Adjust confidence based on block type
            type_confidence = {
                'title': 0.9,      # Titles should be concise
                'heading': 0.85,   # Headings should be concise
                'paragraph': 0.8,  # Paragraphs can vary more
                'list': 0.85,      # Lists should maintain structure
                'figure': 0.95,    # Figure captions should be accurate
                'table': 0.9,      # Table content should be precise
                'footnote': 0.8    # Footnotes can vary
            }.get(block_type, 0.8)
            
            # Check for format preservation
            format_preserved = self._check_format_preservation(original_text, translated_text)
            format_confidence = 0.9 if format_preserved else 0.6
            
            # Calculate final confidence
            confidence = (length_ratio * 0.3 + type_confidence * 0.4 + format_confidence * 0.3)
            
            return min(max(confidence, 0.0), 1.0)
            
        except Exception as e:
            logger.warning(f"Error calculating translation confidence: {e}")
            return 0.5
    
    def _check_format_preservation(self, original_text, translated_text):
        """Check if formatting has been preserved in translation"""
        try:
            import re
            
            # Check for common formatting patterns
            format_indicators = [
                r'\*\*.*?\*\*',  # Bold
                r'\*.*?\*',      # Italic
                r'`.*?`',        # Code
                r'\[.*?\]\(.*?\)', # Links
                r'^[\s]*[-*+]\s+', # List items
                r'^[\s]*\d+\.\s+', # Numbered lists
                r'^#{1,6}\s+',   # Headings
                r'^>\s+',        # Blockquotes
            ]
            
            original_formats = 0
            translated_formats = 0
            
            for pattern in format_indicators:
                original_formats += len(re.findall(pattern, original_text, re.MULTILINE))
                translated_formats += len(re.findall(pattern, translated_text, re.MULTILINE))
            
            # If no formatting in original, consider it preserved
            if original_formats == 0:
                return True
            
            # Calculate preservation ratio
            preservation_ratio = translated_formats / max(original_formats, 1)
            return preservation_ratio >= 0.8  # 80% preservation threshold
            
        except Exception as e:
            logger.warning(f"Error checking format preservation: {e}")
            return False
    
    def _get_reading_order(self):
        """Get the reading order of nodes from the graph"""
        # Use the graph's reading order if available
        if hasattr(self.graph, 'get_reading_order'):
            return self.graph.get_reading_order()
        
        # Fallback: sort by spatial position (top to bottom, left to right)
        nodes = list(self.graph.nodes())
        nodes.sort(key=lambda n: (n.bbox[1], n.bbox[0]))  # Sort by y, then x
        return nodes
    
    def _extract_structured_content(self):
        """Extract structured content from the graph"""
        structured_content = []
        
        for node in self.reading_order:
            if not node.text:
                continue
            
            content_block = {
                'type': node.class_label,
                'semantic_role': node.semantic_role,
                'text': node.text,
                'confidence': node.confidence,
                'bbox': node.bbox,
                'node_id': node.node_id,
                'extra': node.extra.copy() if node.extra else {},
                'children': []
            }
            
            # Find child nodes (nodes that belong to this node)
            children = self._get_child_nodes(node.node_id)
            for child in children:
                child_content = {
                    'type': child.class_label,
                    'semantic_role': child.semantic_role,
                    'text': child.text,
                    'confidence': child.confidence,
                    'bbox': child.bbox,
                    'node_id': child.node_id
                }
                content_block['children'].append(child_content)
            
            structured_content.append(content_block)
        
        return structured_content
    
    def _get_child_nodes(self, parent_id):
        """Get child nodes of a parent node"""
        children = []
        for edge in self.graph.graph.edges(data=True):
            if (edge[0] == parent_id and 
                edge[2].get('relationship') in ['contains', 'has_content', 'contains_item']):
                child_id = edge[1]
                child_node = self.graph.get_node(child_id)
                if child_node:
                    children.append(child_node)
        return children
    
    def _generate_html(self, preserve_layout=True):
        """Generate HTML output with optional layout preservation"""
        html_parts = []
        
        # HTML header
        html_parts.append("""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Translated Document</title>
    <style>
        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }
        .document-block { margin-bottom: 1em; }
        .title { font-size: 1.5em; font-weight: bold; margin-bottom: 0.5em; }
        .paragraph { text-align: justify; }
        .list { margin-left: 20px; }
        .list-item { margin-bottom: 0.5em; }
        .figure { text-align: center; margin: 1em 0; }
        .caption { font-style: italic; font-size: 0.9em; }
        .table { border-collapse: collapse; width: 100%; margin: 1em 0; }
        .table th, .table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        .table th { background-color: #f2f2f2; }
        .quote { font-style: italic; border-left: 3px solid #ccc; padding-left: 10px; }
        .footnote { font-size: 0.8em; color: #666; }
    </style>
</head>
<body>""")
        
        # Generate content blocks
        for block in self.structured_content:
            html_block = self._generate_html_block(block, preserve_layout)
            html_parts.append(html_block)
        
        # HTML footer
        html_parts.append("</body>\n</html>")
        
        return '\n'.join(html_parts)
    
    def _generate_html_block(self, block, preserve_layout):
        """Generate HTML for a single content block"""
        block_type = block['type']
        text = block['text']
        semantic_role = block.get('semantic_role', '')
        
        # Add layout preservation if requested
        style_attr = ""
        if preserve_layout and block['bbox']:
            x1, y1, x2, y2 = block['bbox']
            style_attr = f' style="position: relative; left: {x1}px; top: {y1}px; width: {x2-x1}px;"'
        
        if block_type == 'title':
            heading_level = semantic_role if semantic_role.startswith('h') else 'h1'
            return f'<{heading_level} class="title document-block"{style_attr}>{text}</{heading_level}>'
        
        elif block_type == 'paragraph':
            return f'<p class="paragraph document-block"{style_attr}>{text}</p>'
        
        elif block_type == 'list':
            list_type = 'ol' if semantic_role == 'ol' else 'ul'
            items_html = []
            for child in block['children']:
                items_html.append(f'<li class="list-item">{child["text"]}</li>')
            items_content = '\n'.join(items_html)
            return f'<{list_type} class="list document-block"{style_attr}>\n{items_content}\n</{list_type}>'
        
        elif block_type == 'figure':
            return f'<figure class="figure document-block"{style_attr}>\n<div class="figure-content">[Figure Placeholder]</div>\n<figcaption class="caption">{text}</figcaption>\n</figure>'
        
        elif block_type == 'table':
            return f'<table class="table document-block"{style_attr}>\n<tr><td>{text}</td></tr>\n</table>'
        
        elif block_type == 'quote':
            return f'<blockquote class="quote document-block"{style_attr}>{text}</blockquote>'
        
        elif block_type == 'footnote':
            return f'<div class="footnote document-block"{style_attr}>{text}</div>'
        
        else:
            return f'<div class="document-block"{style_attr}>{text}</div>'
    
    def _generate_markdown(self, preserve_layout=True):
        """Generate Markdown output"""
        markdown_parts = []
        
        for block in self.structured_content:
            markdown_block = self._generate_markdown_block(block, preserve_layout)
            markdown_parts.append(markdown_block)
        
        return '\n\n'.join(markdown_parts)
    
    def _generate_markdown_block(self, block, preserve_layout):
        """Generate Markdown for a single content block"""
        block_type = block['type']
        text = block['text']
        semantic_role = block.get('semantic_role', '')
        
        # Add layout comment if requested
        layout_comment = ""
        if preserve_layout and block['bbox']:
            x1, y1, x2, y2 = block['bbox']
            layout_comment = f"<!-- Layout: ({x1},{y1},{x2},{y2}) -->\n"
        
        if block_type == 'title':
            heading_level = 1
            if semantic_role == 'h2':
                heading_level = 2
            elif semantic_role == 'h3':
                heading_level = 3
            return f"{layout_comment}{'#' * heading_level} {text}"
        
        elif block_type == 'paragraph':
            return f"{layout_comment}{text}"
        
        elif block_type == 'list':
            list_marker = "1." if semantic_role == 'ol' else "-"
            items_md = []
            for child in block['children']:
                items_md.append(f"{list_marker} {child['text']}")
            return f"{layout_comment}{chr(10).join(items_md)}"
        
        elif block_type == 'figure':
            return f"{layout_comment}![Figure](figure_placeholder.png)\n\n*{text}*"
        
        elif block_type == 'table':
            return f"{layout_comment}| {text} |\n| --- |"
        
        elif block_type == 'quote':
            return f"{layout_comment}> {text}"
        
        elif block_type == 'footnote':
            return f"{layout_comment}^[{text}]"
        
        else:
            return f"{layout_comment}{text}"
    
    def _generate_word_document(self, preserve_layout=True):
        """Generate Word document using python-docx"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not available for Word document generation")
            return None
        
        doc = Document()
        
        for block in self.structured_content:
            self._add_word_block(doc, block, preserve_layout)
        
        return doc
    
    def _add_word_block(self, doc, block, preserve_layout):
        """Add a content block to the Word document"""
        block_type = block['type']
        text = block['text']
        semantic_role = block.get('semantic_role', '')
        
        if block_type == 'title':
            heading_level = 1
            if semantic_role == 'h2':
                heading_level = 2
            elif semantic_role == 'h3':
                heading_level = 3
            
            heading = doc.add_heading(text, level=heading_level)
            if preserve_layout and block['bbox']:
                # Add layout information as comment
                x1, y1, x2, y2 = block['bbox']
                heading._element.append(f"<!-- Layout: ({x1},{y1},{x2},{y2}) -->")
        
        elif block_type == 'paragraph':
            p = doc.add_paragraph(text)
            if preserve_layout and block['bbox']:
                x1, y1, x2, y2 = block['bbox']
                p._element.append(f"<!-- Layout: ({x1},{y1},{x2},{y2}) -->")
        
        elif block_type == 'list':
            if semantic_role == 'ol':
                # Numbered list
                for child in block['children']:
                    doc.add_paragraph(child['text'], style='List Number')
            else:
                # Bullet list
                for child in block['children']:
                    doc.add_paragraph(child['text'], style='List Bullet')
        
        elif block_type == 'figure':
            # Add figure placeholder
            p = doc.add_paragraph("[Figure Placeholder]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add caption
            caption = doc.add_paragraph(text)
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.style = 'Caption'
        
        elif block_type == 'table':
            # Simple table
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = text
        
        elif block_type == 'quote':
            p = doc.add_paragraph(text)
            p.style = 'Quote'
        
        else:
            doc.add_paragraph(text)
    
    def _generate_json_structure(self):
        """Generate JSON structure of the document"""
        return {
            'metadata': {
                'total_blocks': len(self.structured_content),
                'reading_order': [block['node_id'] for block in self.structured_content],
                'generated_at': str(datetime.now())
            },
            'content': self.structured_content
        }
    
    def get_reconstruction_summary(self):
        """Get a summary of the reconstruction process with translation metrics"""
        summary = {
            'total_blocks': len(self.structured_content),
            'block_types': {},
            'translation_applied': any('translated_text' in block for block in self.structured_content),
            'layout_preserved': any('bbox' in block for block in self.structured_content),
            'translation_metrics': {
                'total_translated': 0,
                'average_confidence': 0.0,
                'high_confidence_blocks': 0,  # confidence >= 0.8
                'medium_confidence_blocks': 0,  # 0.6 <= confidence < 0.8
                'low_confidence_blocks': 0,  # confidence < 0.6
                'format_preserved_blocks': 0,
                'format_lost_blocks': 0
            }
        }
        
        # Count block types and translation metrics
        total_confidence = 0.0
        translated_count = 0
        
        for block in self.structured_content:
            block_type = block['type']
            summary['block_types'][block_type] = summary['block_types'].get(block_type, 0) + 1
            
            # Translation metrics
            if 'translation_confidence' in block:
                confidence = block['translation_confidence']
                total_confidence += confidence
                translated_count += 1
                
                if confidence >= 0.8:
                    summary['translation_metrics']['high_confidence_blocks'] += 1
                elif confidence >= 0.6:
                    summary['translation_metrics']['medium_confidence_blocks'] += 1
                else:
                    summary['translation_metrics']['low_confidence_blocks'] += 1
                
                # Format preservation
                if block.get('format_preserved', False):
                    summary['translation_metrics']['format_preserved_blocks'] += 1
                else:
                    summary['translation_metrics']['format_lost_blocks'] += 1
        
        # Calculate average confidence
        if translated_count > 0:
            summary['translation_metrics']['total_translated'] = translated_count
            summary['translation_metrics']['average_confidence'] = total_confidence / translated_count
        
        return summary
    
    def print_reconstruction_summary(self):
        """Print a comprehensive summary of the reconstruction process"""
        summary = self.get_reconstruction_summary()
        
        print(f"Document Reconstruction Summary:")
        print(f"  Total content blocks: {summary['total_blocks']}")
        print(f"  Translation applied: {summary['translation_applied']}")
        print(f"  Layout preserved: {summary['layout_preserved']}")
        
        # Translation metrics
        if summary['translation_metrics']['total_translated'] > 0:
            print(f"\nTranslation Quality Metrics:")
            print(f"  Total translated blocks: {summary['translation_metrics']['total_translated']}")
            print(f"  Average confidence: {summary['translation_metrics']['average_confidence']:.3f}")
            print(f"  High confidence (≥0.8): {summary['translation_metrics']['high_confidence_blocks']}")
            print(f"  Medium confidence (0.6-0.8): {summary['translation_metrics']['medium_confidence_blocks']}")
            print(f"  Low confidence (<0.6): {summary['translation_metrics']['low_confidence_blocks']}")
            print(f"  Format preserved: {summary['translation_metrics']['format_preserved_blocks']}")
            print(f"  Format lost: {summary['translation_metrics']['format_lost_blocks']}")
        
        print(f"\nBlock type distribution:")
        for block_type, count in summary['block_types'].items():
            print(f"    {block_type}: {count}")
        
        # Show confidence distribution by block type
        if summary['translation_metrics']['total_translated'] > 0:
            print(f"\nConfidence by block type:")
            type_confidence = {}
            type_count = {}
            
            for block in self.structured_content:
                if 'translation_confidence' in block:
                    block_type = block['type']
                    confidence = block['translation_confidence']
                    
                    if block_type not in type_confidence:
                        type_confidence[block_type] = 0.0
                        type_count[block_type] = 0
                    
                    type_confidence[block_type] += confidence
                    type_count[block_type] += 1
            
            for block_type, total_conf in type_confidence.items():
                avg_conf = total_conf / type_count[block_type]
                print(f"    {block_type}: {avg_conf:.3f} ({type_count[block_type]} blocks)")


def reconstruct_document_from_graph(graph, translation_service=None, output_format='html', target_language='en', preserve_layout=True):
    """
    Convenience function to reconstruct a document from a graph.
    
    Args:
        graph: Refined DocumentGraph instance
        translation_service: Translation service for content translation
        output_format: 'html', 'markdown', 'word', or 'json'
        target_language: Target language for translation
        preserve_layout: Whether to preserve original layout structure
        
    Returns:
        str or dict: Reconstructed document content
    """
    reconstructor = DocumentReconstructor(graph, translation_service)
    return reconstructor.reconstruct_document(output_format, target_language, preserve_layout)
