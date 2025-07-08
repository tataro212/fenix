"""
Document Generator Module for Ultimate PDF Translator

Handles Word document creation and PDF conversion with proper formatting and structure
"""

import os
import logging
import re
import time
import json
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn # Ensure qn is imported

from config_manager import config_manager
from utils import sanitize_for_xml, sanitize_filepath

logger = logging.getLogger(__name__)

# Import structured document model
try:
    from structured_document_model import (
        Document as StructuredDocument, ContentType, Heading, Paragraph, ImagePlaceholder,
        Table, CodeBlock, ListItem, Footnote, Equation, Caption, Metadata
    )
    STRUCTURED_MODEL_AVAILABLE = True
except ImportError:
    STRUCTURED_MODEL_AVAILABLE = False
    logger.warning("Structured document model not available")

# Import Digital Twin model for enhanced document generation
try:
    from digital_twin_model import (
        DocumentModel, PageModel,
        TextBlock, ImageBlock,
        TableBlock, TOCEntry, BlockType, StructuralRole
    )
    DIGITAL_TWIN_MODEL_AVAILABLE = True
except ImportError:
    DIGITAL_TWIN_MODEL_AVAILABLE = False
    logger.warning("Digital Twin model not available")

# Global bookmark counter
# bookmark_id_counter = 0 # Comment out or remove if self.bookmark_id replaces its ToC usage

class WordDocumentGenerator:
    """Generates Word documents with proper structure and formatting"""
    
    def __init__(self, translation_service=None, pdf_parser=None):
        """Initialize the document generator with translation service and PDF parser"""
        self.translation_service = translation_service
        self.pdf_parser = pdf_parser
        self.logger = logging.getLogger(__name__)
        self.word_settings = config_manager.word_output_settings
        self.toc_entries = []
        self.bookmark_id = 0
    
    def _add_heading_block(self, doc, block_item):
        text_content = None
        level = 1

        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, Heading):
            text_content = block_item.content
            level = block_item.level
        elif isinstance(block_item, dict):
            text_content = block_item.get('content') or block_item.get('text')
            level = block_item.get('level')
            type_str = block_item.get('type', '').lower()
            if level is None and type_str.startswith('h') and type_str[1:].isdigit():
                level = int(type_str[1:])
            if level is None: level = 1 # Default if not found
        else:
            logger.warning(f"Skipping unknown heading block type: {type(block_item)}")
            return

        if not text_content:
            logger.warning("Skipping empty heading block.")
            return

        safe_content = sanitize_for_xml(text_content)
        p = doc.add_paragraph(style=f'Heading {level}')
        
        self.bookmark_id += 1
        bookmark_name = f"_Toc_Bookmark_{self.bookmark_id}"
        run = p.add_run(safe_content)
        
        tag_start = OxmlElement('w:bookmarkStart')
        tag_start.set(qn('w:id'), str(self.bookmark_id))
        tag_start.set(qn('w:name'), bookmark_name)
        tag_end = OxmlElement('w:bookmarkEnd')
        tag_end.set(qn('w:id'), str(self.bookmark_id))

        if run._r is not None and p._p is not None:
            # CORRECTED LINE: Use addprevious on the run element itself
            run._r.addprevious(tag_start)
            run._r.addnext(tag_end)
        else:
            logger.warning(f"Could not precisely position bookmark for heading: {safe_content}. Run or paragraph element was None.")
            # Fallback: append to paragraph, less ideal but avoids error
            p._p.append(tag_start)
            p._p.append(tag_end)
        
        self.toc_entries.append({
            'text': text_content,
            'level': level,
            'bookmark': bookmark_name
        })
        logger.debug(f"Added Heading '{text_content}' (L{level}) with bookmark '{bookmark_name}'.")

    def _add_paragraph_block(self, doc, block_item):
        text_content = None
        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, Paragraph):
            text_content = block_item.content
        elif isinstance(block_item, dict):
            text_content = block_item.get('content') or block_item.get('text')
        else:
            logger.warning(f"Skipping unknown paragraph block type: {type(block_item)}")
            return

        if text_content:
            safe_content = sanitize_for_xml(text_content)
            doc.add_paragraph(safe_content)
        else:
            logger.debug("Skipping empty paragraph block.")

    def _add_image_placeholder_block(self, doc, block_item, image_folder_path):
        image_filename = None
        caption_text = None

        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, ImagePlaceholder):
            image_filename = block_item.image_path
            caption_text = block_item.caption
        elif isinstance(block_item, dict):
            image_filename = block_item.get('image_path') or block_item.get('filepath') # common dict keys
            caption_text = block_item.get('caption') or block_item.get('text') # some dicts might use 'text' for caption
        else:
            logger.warning(f"Skipping unknown image block type: {type(block_item)}")
            return

        if not image_filename:
            logger.warning("Skipping image block with no image_path/filename.")
            return

        actual_image_path = image_filename
        if image_folder_path and not os.path.isabs(image_filename):
            actual_image_path = os.path.join(image_folder_path, image_filename)

        if os.path.exists(actual_image_path):
            try:
                doc.add_picture(actual_image_path, width=Inches(self.word_settings.get('image_width_inches', 6.0)))
                if caption_text:
                    safe_caption = sanitize_for_xml(caption_text)
                    caption_style_name = self.word_settings.get('caption_style', 'Caption')
                    try:
                        p_caption = doc.add_paragraph(style=caption_style_name)
                        p_caption.add_run(safe_caption)
                    except KeyError: # Style not found
                        logger.warning(f"Caption style '{caption_style_name}' not found. Adding caption as normal paragraph.")
                        p_caption = doc.add_paragraph()
                        p_caption.add_run(safe_caption)
                    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logger.debug(f"Added image: {actual_image_path}")
            except Exception as e:
                logger.error(f"Error adding image {actual_image_path}: {e}")
        else:
            logger.warning(f"Image not found at path: {actual_image_path}. Skipping image.")

    def _add_list_item_block(self, doc, block_item):
        text_content = None
        level = 0 # Placeholder for list level if available in dict/model
        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, ListItem):
            text_content = block_item.content
            level = getattr(block_item, 'level', 1) # Assuming ListItem might have a level
        elif isinstance(block_item, dict):
            text_content = block_item.get('content') or block_item.get('text')
            level = block_item.get('level', 1)
        else:
            logger.warning(f"Skipping unknown list item block type: {type(block_item)}")
            return

        if text_content:
            safe_content = sanitize_for_xml(text_content)
            # Basic list item handling, python-docx requires more for actual bullets/numbering
            # This adds it as a paragraph with indentation.
            p = doc.add_paragraph()
            p.add_run(safe_content)
            p.paragraph_format.left_indent = Inches(0.25 * level)
            logger.debug(f"Added list item (as indented paragraph): {safe_content}")
        else:
            logger.debug("Skipping empty list item block.")
            
    def _add_table_block(self, doc, block_item):
        """
        Fully implement table reconstruction as required by Sub-task 2.4.
        
        Creates a properly formatted table in the Word document using python-docx,
        with "Table Grid" style for professional formatting.
        """
        try:
            # Import TableModel for type checking
            from models import TableModel
            
            table_data = None
            header_row = None
            caption = None
            
            # Handle TableModel objects
            if isinstance(block_item, TableModel):
                table_data = block_item.content  # List[List[str]]
                header_row = block_item.header_row
                caption = block_item.caption
                self.logger.info(f"Processing TableModel with {len(table_data)} rows")
            
            # Handle dictionary format (for compatibility)
            elif isinstance(block_item, dict):
                # Extract table data from various possible formats
                if 'content' in block_item and isinstance(block_item['content'], list):
                    table_data = block_item['content']
                elif 'rows' in block_item:
                    table_data = block_item['rows']
                elif 'translated_rows' in block_item:
                    table_data = block_item['translated_rows']
                else:
                    # Fallback: try to create single-cell table with text content
                    text_content = block_item.get('content') or block_item.get('text', '')
                    if text_content:
                        table_data = [[text_content]]
                
                header_row = block_item.get('header_row')
                caption = block_item.get('caption')
            
            # Handle structured model Table objects
            elif STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, Table):
                if hasattr(block_item, 'rows'):
                    table_data = block_item.rows
                    caption = getattr(block_item, 'caption', None)
            
            # Validate table data
            if not table_data or not isinstance(table_data, list) or len(table_data) == 0:
                self.logger.warning("No valid table data found, creating placeholder")
                doc.add_paragraph("[Table: No data available]")
                return
            
            # Ensure all rows have the same number of columns
            max_cols = max(len(row) if isinstance(row, list) else 1 for row in table_data)
            normalized_data = []
            
            for row in table_data:
                if isinstance(row, list):
                    # Pad row to max_cols length
                    normalized_row = row + [''] * (max_cols - len(row))
                    normalized_data.append(normalized_row[:max_cols])  # Truncate if too long
                else:
                    # Single value, convert to list
                    normalized_data.append([str(row)] + [''] * (max_cols - 1))
            
            # Create the Word table
            word_table = doc.add_table(rows=len(normalized_data), cols=max_cols)
            
            # Apply "Table Grid" style for professional formatting
            try:
                word_table.style = 'Table Grid'
            except KeyError:
                # Fallback if Table Grid style not available
                self.logger.warning("Table Grid style not available, using default table style")
            
            # Populate table cells with data
            for row_idx, row_data in enumerate(normalized_data):
                table_row = word_table.rows[row_idx]
                
                for col_idx, cell_data in enumerate(row_data):
                    cell = table_row.cells[col_idx]
                    
                    # Clean and sanitize cell content
                    cell_text = sanitize_for_xml(str(cell_data).strip())
                    cell.text = cell_text
                    
                    # Apply header formatting to first row if it's identified as header
                    if row_idx == 0 and (header_row or len(normalized_data) > 1):
                        # Make header row bold
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            
            # Add caption if available
            if caption:
                caption_p = doc.add_paragraph()
                caption_run = caption_p.add_run(f"Table: {sanitize_for_xml(caption)}")
                caption_run.font.italic = True
                caption_run.font.size = Pt(10)
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self.logger.info(f"✅ Successfully created table with {len(normalized_data)} rows and {max_cols} columns")
            
        except Exception as e:
            self.logger.error(f"❌ Error creating table block: {e}")
            # Fallback to simple paragraph
            doc.add_paragraph(f"[Table processing error: {str(e)[:100]}]")


    def _add_code_block(self, doc, block_item):
        # Placeholder - actual code block formatting can be complex
        text_content = None
        if STRUCTURED_MODEL_AVAILABLE and isinstance(block_item, CodeBlock):
            text_content = block_item.content
        elif isinstance(block_item, dict):
            text_content = block_item.get('content') or block_item.get('text')
        
        if text_content:
            logger.debug(f"Adding code block (as preformatted text): {text_content[:100]}")
            # Simple approach: add as a paragraph with a "code" style if available, or just monospace font
            # For real code blocks, you might want to preserve whitespace carefully.
            p = doc.add_paragraph()
            run = p.add_run(sanitize_for_xml(text_content))
            try:
                # Attempt to apply a 'Code' style if it exists in the template
                p.style = 'Code' 
            except KeyError:
                # Fallback to basic monospace font
                font = run.font
                font.name = 'Courier New' # Or another common monospace font
                font.size = Pt(10)
        else:
            logger.warning(f"Skipping empty code block: {type(block_item)}")

    # --- New method for inserting ToC (Step 3) ---
    def _insert_toc(self, doc, force_skip_toc=False):
        """Insert table of contents at the beginning of the document based on collected toc_entries."""
        if force_skip_toc:
            logger.info("TOC generation explicitly skipped - document has no original TOC")
            return
            
        if not self.toc_entries:
            logger.info("No TOC entries found, skipping TOC generation")
            return
            
        logger.info("Generating Table of Contents...")
        
        # Create a new document for the TOC
        toc_doc = Document()
        
        # Configure fonts for TOC document too
        self._configure_document_fonts_for_unicode(toc_doc)
        
        # Add TOC title
        title = toc_doc.add_paragraph("Table of Contents")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style = 'Title'
        
        # Add a decorative line
        line = toc_doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.add_run("_" * 40)
        
        # Sort TOC entries by level and text
        sorted_entries = sorted(self.toc_entries, key=lambda x: (x['level'], x['text']))
        
        # Add TOC entries with hyperlinks and page numbers
        for entry in sorted_entries:
            level = entry['level']
            text_for_display = sanitize_for_xml(entry['text'])
            bookmark = entry['bookmark']
            
            # Create TOC entry paragraph
            p = toc_doc.add_paragraph()
            p.style = 'Normal'
            
            # Set indentation based on level
            indent_size = Pt(20 * (level - 1))
            p.paragraph_format.left_indent = indent_size
            
            # Add heading text with hyperlink
            try:
                hyperlink_run = self._add_hyperlink(p, text_for_display, bookmark)
                if level == 1:
                    hyperlink_run.bold = True
                    hyperlink_run.font.size = Pt(12)
                elif level == 2:
                    hyperlink_run.font.size = Pt(11)
                else:
                    hyperlink_run.font.size = Pt(10)
                    hyperlink_run.italic = True
            except Exception as e:
                logger.debug(f"Could not create hyperlink for {text_for_display}: {e}")
                # Fallback to regular text
                text_run = p.add_run(text_for_display)
                if level == 1:
                    text_run.bold = True
                    text_run.font.size = Pt(12)
                elif level == 2:
                    text_run.font.size = Pt(11)
                else:
                    text_run.font.size = Pt(10)
                    text_run.italic = True

            # Add dots and page number
            dots_needed = max(3, 60 - len(text_for_display))
            dots_run = p.add_run(" " + "." * dots_needed + " ")
            dots_run.font.color.rgb = RGBColor(128, 128, 128)

            # Add page reference field
            fldSimple = OxmlElement('w:fldSimple')
            fldSimple.set(qn('w:instr'), f'PAGEREF {bookmark} \\h')
            
            run_in_field = OxmlElement('w:r')
            text_in_field = OxmlElement('w:t')
            text_in_field.text = "..."  # Placeholder that will be replaced with actual page number
            run_in_field.append(text_in_field)
            fldSimple.append(run_in_field)
            
            p._p.append(fldSimple)
        
        # Add spacing after TOC
        toc_doc.add_paragraph()
        
        # Insert TOC at the beginning of the document (Directive I requirement)
        # Insert elements in reverse order to maintain proper sequence
        toc_elements = list(toc_doc.element.body)
        for i, element in enumerate(reversed(toc_elements)):
            doc.element.body.insert(0, element)
        
        logger.info(f"Successfully generated TOC with {len(sorted_entries)} entries")

    def create_word_document_from_structured_document(self, structured_document_or_list, output_filepath,
                                                     image_folder_path=None, cover_page_data=None):
        """
        UNIFIED Document Generation Method - handles both StructuredDocument objects and list[dict] structures.
        
        This is now the sole, authoritative method for generating final documents as per Directive I.
        Implements strict two-pass reconstruction:
        - Pass 1: Content and Bookmark Generation
        - Pass 2: TOC Insertion
        """
        
        # Handle both StructuredDocument objects and list[dict] data structures
        if isinstance(structured_document_or_list, list):
            # Handle list[dict] structure from pipeline
            return self._create_from_list_structure(structured_document_or_list, output_filepath, image_folder_path)
        
        # Handle StructuredDocument objects (original functionality)
        if not STRUCTURED_MODEL_AVAILABLE:
            raise Exception("Structured document model not available")

        if not isinstance(structured_document_or_list, StructuredDocument):
            raise ValueError(f"Expected StructuredDocument or list[dict], got {type(structured_document_or_list)}")

        structured_document = structured_document_or_list

        # --- PASS 1 INITIALIZATION: Initialize TOC system ---
        self.toc_entries = []
        self.bookmark_id = 0

        # Normalize paths to handle mixed separators
        output_filepath = os.path.normpath(output_filepath)
        if image_folder_path:
            image_folder_path = os.path.normpath(image_folder_path)

        logger.info(f"--- Creating Word Document from Structured Document: {structured_document.title} ---")
        logger.info(f"📊 Processing {len(structured_document.content_blocks)} content blocks")

        # --- PASS 1: Content and Bookmark Generation ---
        doc = Document()

        # Add document title as first heading
        if structured_document.title:
            # Sanitize title before adding as heading
            safe_title = sanitize_for_xml(structured_document.title)
            title_heading = doc.add_heading(safe_title, level=0)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add cover page if provided
        if cover_page_data:
            self._add_cover_page(doc, cover_page_data, image_folder_path)
            doc.add_page_break()

        # Main content processing loop - populates toc_entries and adds bookmarks
        for block in structured_document.content_blocks:
            self._add_content_block(doc, block, image_folder_path)

        # --- PASS 2: TOC Insertion (conditional based on complexity analysis) ---
        # Check if we should skip TOC generation based on document complexity
        skip_toc = getattr(self, '_skip_toc_generation', False)
        self._insert_toc(doc, force_skip_toc=skip_toc)

        # Save document
        try:
            # Use centralized sanitization from utils
            sanitized_filepath = sanitize_filepath(output_filepath)

            # Ensure the output directory exists before saving
            output_dir = os.path.dirname(sanitized_filepath)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
                logger.info(f"Created output directory: {output_dir}")

            doc.save(sanitized_filepath)
            logger.info(f"Word document saved successfully: {sanitized_filepath}")
            # Return the actual file path used for saving
            return sanitized_filepath
        except Exception as e:
            logger.error(f"Error saving Word document: {e}")
            logger.error(f"Attempted path: {output_filepath}")
            return None

        # After saving the document, export the final output order
        final_output_blocks = []
        for page in structured_document_or_list:
            for block in page.get_all_blocks():
                text = getattr(block, 'translated_text', None) or getattr(block, 'original_text', None)
                final_output_blocks.append({
                    'block_id': getattr(block, 'block_id', None),
                    'block_type': getattr(block, 'block_type', None),
                    'text': text
                })
                # ... existing code for writing block to document ...
        # ... existing code ...
        # After saving the document, export the final output order
        output_dir = os.path.dirname(output_filepath)
        output_name = os.path.splitext(os.path.basename(output_filepath))[0]
        debug_path = os.path.join(output_dir, f"{output_name}_final_output_order.json")
        with open(debug_path, 'w', encoding='utf-8') as f:
            json.dump(final_output_blocks, f, ensure_ascii=False, indent=2)
        self.logger.info(f"🪪 Final output block order exported to: {debug_path}")

    def _create_from_list_structure(self, structured_content: List[Dict[str, Any]], output_path: str, images_dir: str = None):
        """
        Create Word document from list[dict] structure using unified two-pass approach.
        This replaces the deprecated create_word_document_with_structure method.
        """
        self.logger.info(f"--- Creating Word Document from List Structure: {os.path.basename(output_path)} ---")
        
        try:
            # --- PASS 1 INITIALIZATION: Initialize TOC system ---
            self.toc_entries = []
            self.bookmark_id = 0
            
            # --- PASS 1: Content and Bookmark Generation ---
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'  # Use a font that supports a wide range of characters
            font.size = Pt(11)

            all_text_sections = structured_content
            self.logger.info(f"Processing {len(all_text_sections)} content sections for document generation.")

            if not all_text_sections:
                self.logger.warning("No content sections found. Creating empty document.")
                doc.save(output_path)
                return True

            # Sort by reading order: top-to-bottom, then left-to-right
            all_text_sections.sort(key=lambda x: (x.get('bbox', [0, 0, 0, 0])[1], x.get('bbox', [0, 0, 0, 0])[0]))

            # Process each section and populate TOC entries if headings are found
            for section in all_text_sections:
                text = section.get('text', '')
                label = section.get('label', 'paragraph')

                if label in ['title', 'heading', 'header'] or (label == 'text' and len(text) < 100 and text.endswith(':')):
                    # Treat as heading - this will populate toc_entries
                    block_item = {
                        'content': text,
                        'type': 'heading',
                        'level': 1 if label == 'title' else 2
                    }
                    self._add_heading_block(doc, block_item)
                elif label in ('header', 'footer'):
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.italic = True
                    run.font.size = Pt(9)
                else:
                    doc.add_paragraph(text)

            # --- PASS 2: TOC Insertion (unconditional) ---
            self._insert_toc(doc)

            doc.save(output_path)
            self.logger.info(f"Word document saved successfully: {output_path}")
            
            return True

        except Exception as e:
            self.logger.error(f"Failed to save Word document: {e}", exc_info=True)
            return False

    def _add_table_of_contents_from_document(self, doc, structured_document):
        """Add table of contents from structured document headings"""
        try:
            # Add TOC title with enhanced styling and sanitization
            safe_toc_title = sanitize_for_xml(self.word_settings['toc_title'])
            toc_title = doc.add_heading(safe_toc_title, level=1)
            toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add decorative line under title
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run("─" * 50)
            title_run.font.color.rgb = RGBColor(128, 128, 128)

            # Extract headings from structured document
            headings = structured_document.get_blocks_by_type(ContentType.HEADING)

            if not headings:
                logger.warning("No headings found for table of contents")
                return

            # Add TOC entries
            for i, heading_block in enumerate(headings):
                if isinstance(heading_block, Heading):
                    self._add_toc_entry_from_heading_block(doc, heading_block, i + 1)

            # Add spacing after TOC
            doc.add_paragraph()

            logger.info(f"Table of contents added with {len(headings)} entries")

        except Exception as e:
            logger.warning(f"Could not add table of contents: {e}")

    def _add_toc_entry_from_heading_block(self, doc, heading_block, entry_number):
        """Add TOC entry from a Heading content block with improved page estimation"""
        try:
            text = heading_block.content
            level = heading_block.level
            original_page_num = heading_block.page_num

            # Improved page estimation for final document
            # Account for TOC pages, cover page, etc.
            estimated_page = self._estimate_final_page_number(original_page_num, entry_number, level)

            # Sanitize text for XML compatibility
            safe_text = sanitize_for_xml(text)

            # Create TOC paragraph
            toc_paragraph = doc.add_paragraph()
            toc_paragraph.style = 'Normal'

            # Set indentation based on level
            base_indent = 0.2
            level_indent = (level - 1) * self.word_settings['list_indent_per_level_inches']
            toc_paragraph.paragraph_format.left_indent = Inches(base_indent + level_indent)

            # Add entry number for main headings
            if level == 1:
                number_run = toc_paragraph.add_run(f"{entry_number}. ")
                number_run.bold = True
                number_run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue

            # Add heading text with hyperlink
            try:
                bookmark_name = f"heading_{entry_number}_{level}"
                hyperlink_run = self._add_hyperlink(toc_paragraph, safe_text, bookmark_name)
                if level == 1:
                    hyperlink_run.bold = True
                    hyperlink_run.font.size = Pt(12)
                elif level == 2:
                    hyperlink_run.font.size = Pt(11)
                else:
                    hyperlink_run.font.size = Pt(10)
                    hyperlink_run.italic = True
            except Exception as e:
                logger.debug(f"Could not create hyperlink for {safe_text}: {e}")
                # Fallback to regular text
                text_run = toc_paragraph.add_run(safe_text)
                if level == 1:
                    text_run.bold = True
                    text_run.font.size = Pt(12)
                elif level == 2:
                    text_run.font.size = Pt(11)
                else:
                    text_run.font.size = Pt(10)
                    text_run.italic = True

            # Add dots and page number
            dots_needed = max(3, 60 - len(text) - len(str(estimated_page)))
            dots_run = toc_paragraph.add_run(" " + "." * dots_needed + " ")
            dots_run.font.color.rgb = RGBColor(128, 128, 128)

            page_run = toc_paragraph.add_run(str(estimated_page))
            page_run.bold = True if level == 1 else False

        except Exception as e:
            logger.warning(f"Could not add TOC entry: {e}")

    def _estimate_final_page_number(self, original_page_num, entry_number, level):
        """Estimate the final page number in the Word document accounting for TOC and cover pages"""
        # Base estimation: original page + offset for cover page and TOC
        toc_pages = 2  # Estimate 2 pages for TOC
        cover_pages = 1 if self.word_settings.get('add_cover_page', False) else 0

        # Calculate offset
        page_offset = cover_pages + toc_pages

        # For the first few headings, use a more conservative estimation
        if entry_number <= 3:
            estimated_page = page_offset + entry_number
        else:
            # Use original page number with offset
            estimated_page = max(original_page_num + page_offset, page_offset + entry_number)

        return estimated_page

    def _generate_content_with_page_tracking(self, doc, structured_document, image_folder_path):
        """
        Proposition 2: Pass 1 - Generate content and track heading page numbers.

        Renders all content blocks except TOC and creates bookmarks for headings.
        Returns a mapping of heading block_id to actual page number.
        """
        logger.info("🔄 Pass 1: Generating content with page tracking...")

        heading_page_map = {}
        heading_counter = 0

        # Reserve space for TOC (we'll insert it later)
        toc_placeholder = doc.add_paragraph("[TABLE OF CONTENTS PLACEHOLDER]")
        toc_placeholder.style = 'Normal'
        doc.add_page_break()

        # Process all content blocks and track heading positions
        for block in structured_document.content_blocks:
            if isinstance(block, Heading):
                # Add heading with bookmark
                self._add_heading_block(doc, block)
                heading_counter += 1
                
                # Track page number for this heading
                current_page = len(doc.paragraphs) // 50  # Approximate page number
                heading_page_map[block.block_id] = current_page
            else:
                self._add_content_block(doc, block, image_folder_path)

        return heading_page_map

    def _generate_toc_with_page_numbers(self, doc, structured_document, heading_page_map):
        """
        Proposition 2: Pass 2 - Generate TOC with accurate page numbers.
        """
        logger.info("🔄 Pass 2: Generating TOC with page numbers...")

        # Get TOC entries from document
        toc_entries = structured_document.get_toc_entries()
        if not toc_entries:
            logger.warning("No TOC entries found in document")
            return

        # Add TOC title
        toc_title = doc.add_heading("Table of Contents", level=1)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add TOC entries with page numbers
        for entry in toc_entries:
            # Find corresponding heading block
            heading_blocks = [b for b in structured_document.get_all_content_blocks() 
                             if isinstance(b, Heading) and b.content == entry['text']]
            
            if heading_blocks:
                heading = heading_blocks[0]
                page_num = heading_page_map.get(heading.block_id, 1)
                
                # Add TOC entry with page number
                p = doc.add_paragraph()
                p.style = f'TOC {entry["level"]}'
                
                # Add entry text
                run = p.add_run(entry['text'])
                
                # Add dots and page number
                p.add_run(' ' + '.' * (50 - len(entry['text'])) + ' ')
                p.add_run(str(page_num))
                
                # Add hyperlink to bookmark
                if heading.block_id in heading_page_map:
                    bookmark_name = f"_Toc_Bookmark_{heading.block_id}"
                    self._add_hyperlink(p, bookmark_name, str(page_num))

        # Add spacing after TOC
        doc.add_paragraph()

    def _add_hyperlink(self, paragraph, bookmark_name, text):
        """Add a hyperlink to a bookmark in the document."""
        # Create the w:hyperlink element
        hyperlink = OxmlElement('w:hyperlink')
        
        # Set the bookmark reference
        hyperlink.set(qn('w:anchor'), bookmark_name)
        
        # Create the w:r element
        new_run = OxmlElement('w:r')
        
        # Create the w:rPr element
        rPr = OxmlElement('w:rPr')
        
        # Create the w:rStyle element
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        
        # Add the elements to the paragraph
        rPr.append(rStyle)
        new_run.append(rPr)
        
        # Create the w:t element
        t = OxmlElement('w:t')
        t.text = text
        
        # Add the text to the run
        new_run.append(t)
        
        # Add the run to the hyperlink
        hyperlink.append(new_run)
        
        # Add the hyperlink to the paragraph
        paragraph._p.append(hyperlink)

    def _add_content_block(self, doc, block_item, image_folder_path):
        """Add various content block types to the document. Handles both model objects and dicts."""
        try:
            block_handled = False
            if STRUCTURED_MODEL_AVAILABLE:
                if isinstance(block_item, Heading):
                    self._add_heading_block(doc, block_item)
                    block_handled = True
                elif isinstance(block_item, Paragraph):
                    self._add_paragraph_block(doc, block_item)
                    block_handled = True
                elif isinstance(block_item, ImagePlaceholder):
                    self._add_image_placeholder_block(doc, block_item, image_folder_path)
                    block_handled = True
                elif isinstance(block_item, ListItem):
                    self._add_list_item_block(doc, block_item)
                    block_handled = True
                elif isinstance(block_item, Table):
                    self._add_table_block(doc, block_item)
                    block_handled = True
                elif isinstance(block_item, CodeBlock):
                    self._add_code_block(doc, block_item)
                    block_handled = True
                # Add elif for other specific structured_document_model types here
                # e.g., Footnote, Equation, Caption, Metadata if they need special handling

            if not block_handled and isinstance(block_item, dict):
                dict_type = block_item.get('type', '').lower()
                if dict_type.startswith('h') and dict_type[1:].isdigit(): # h1, h2, etc.
                    self._add_heading_block(doc, block_item)
                elif dict_type in ['paragraph', 'text', 'p', 'normal']: # 'normal' for some legacy dicts
                    self._add_paragraph_block(doc, block_item)
                elif dict_type in ['image', 'img', 'imageplaceholder']:
                    self._add_image_placeholder_block(doc, block_item, image_folder_path)
                elif dict_type in ['list_item', 'listitem', 'li']:
                    self._add_list_item_block(doc, block_item)
                elif dict_type == 'table':
                    self._add_table_block(doc, block_item)
                elif dict_type in ['code_block', 'codeblock', 'code']:
                    self._add_code_block(doc, block_item)
                # Add other dict types as needed
                else:
                    # Fallback for unknown dict types: try to treat as paragraph if it has text
                    text_content = block_item.get('content') or block_item.get('text')
                    if text_content:
                        logger.warning(f"Unknown dictionary block type: '{dict_type}', but found text. Treating as paragraph.")
                        self._add_paragraph_block(doc, block_item)
                    else:
                        logger.warning(f"Unknown or empty dictionary block type: '{dict_type}'. Content: {str(block_item)[:200]}")
                block_handled = True # Assume dict was handled or logged

            if not block_handled: # Not a known model object and not a dict
                 logger.warning(f"Unknown block type: {type(block_item)}. Attempting to add as raw string.")
                 try:
                     # Attempt to add as a simple paragraph if it's some other basic type
                     doc.add_paragraph(str(block_item))
                 except Exception as e_raw:
                     logger.error(f"Could not add unknown block type {type(block_item)} as string: {e_raw}")

        except Exception as e:
            logger.error(f"Error processing content block (type: {type(block_item)}): {e}", exc_info=True)

    def _add_cover_page(self, doc, cover_page_data, image_folder_path):
        """Add cover page to document"""
        try:
            cover_paragraph = doc.add_paragraph()
            cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if cover_page_data and image_folder_path:
                cover_image_path = cover_page_data.get('filepath')
                if cover_image_path and os.path.exists(cover_image_path):
                    run = cover_paragraph.add_run()
                    run.add_picture(cover_image_path, width=Inches(6))
                    logger.info("Cover page image added")
                else:
                    cover_paragraph.add_run("Cover Page").font.size = Pt(24)
            else:
                cover_paragraph.add_run("Cover Page").font.size = Pt(24)
                
        except Exception as e:
            logger.warning(f"Could not add cover page: {e}")
    
    def _add_table_of_contents(self, doc, structured_content_list):
        """Add enhanced table of contents with professional formatting"""
        try:
            # Add TOC title with enhanced styling and sanitization
            safe_toc_title = sanitize_for_xml(self.word_settings['toc_title'])
            toc_title = doc.add_heading(safe_toc_title, level=1)
            toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add decorative line under title
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run("─" * 50)
            title_run.font.color.rgb = RGBColor(128, 128, 128)

            # Extract and organize headings for TOC
            headings = self._extract_toc_headings(structured_content_list)

            if not headings:
                logger.warning("No headings found for table of contents")
                return

            # Add TOC entries with enhanced formatting
            for i, heading in enumerate(headings):
                self._add_enhanced_toc_entry(doc, heading, i + 1)

            # Add spacing after TOC
            doc.add_paragraph()

            logger.info(f"Enhanced table of contents added with {len(headings)} entries")

        except Exception as e:
            logger.warning(f"Could not add table of contents: {e}")

    def _extract_toc_headings(self, structured_content_list):
        """Extract and organize headings for TOC with proper bookmark handling"""
        raw_headings = []
        seen_headings = set()  # To prevent duplicates

        # Enhanced page estimation based on content analysis
        page_estimator = self._create_page_estimator()

        # First pass: collect all heading items with accurate page estimation
        for item in structured_content_list:
            item_type = item.get('type', '')

            # Update page estimation based on content type and length
            page_estimator.process_item(item)

            if item_type in ['h1', 'h2', 'h3']:
                text = item.get('text', '').strip()
                if text:
                    # Extract TOC bookmark if present
                    toc_bookmark = None
                    bookmark_match = re.search(r'_Toc_Bookmark_(\d+)', text)
                    if bookmark_match:
                        toc_bookmark = bookmark_match.group(0)
                        # Remove bookmark from display text but keep it for structure
                        text = re.sub(r'_Toc_Bookmark_\d+\s*', '', text).strip()

                    # Check for duplicates using cleaned text
                    clean_text = text.lower()
                    if clean_text not in seen_headings:
                        seen_headings.add(clean_text)
                        level = int(item_type[1])  # Extract level from h1, h2, h3
                        estimated_page = page_estimator.get_current_page()

                        raw_headings.append({
                            'text': text,
                            'level': level,
                            'estimated_page': estimated_page,
                            'original_page': item.get('page_num', estimated_page),
                            'bbox': item.get('bbox', [0, 0, 0, 0]),
                            'content_position': page_estimator.get_position_info(),
                            'toc_bookmark': toc_bookmark  # Preserve TOC bookmark
                        })

        # Second pass: merge consecutive headings that appear to be split
        merged_headings = self._merge_split_headings(raw_headings)

        # Third pass: clean and format the merged headings with refined page numbers
        final_headings = []
        for heading in merged_headings:
            refined_page = self._refine_page_estimation(heading, merged_headings)

            final_headings.append({
                'text': heading['text'],
                'level': heading['level'],
                'estimated_page': refined_page,
                'original_page': heading['original_page'],
                'toc_bookmark': heading['toc_bookmark']  # Preserve TOC bookmark
            })

        return final_headings

    def _configure_document_fonts_for_unicode(self, doc):
        """Configure document fonts to properly support Greek and other Unicode characters"""
        try:
            # Set the default font for the entire document to support Greek characters
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'  # Professional font for normal text
            font.size = Pt(11)
            
            # Set the complex script font for Greek and other complex scripts
            rpr = style.element.xpath('.//w:rPr')[0] if style.element.xpath('.//w:rPr') else None
            if rpr is not None:
                rFonts = rpr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rpr.insert(0, rFonts)
                
                # Set fonts for different script types
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:cs'), 'Arial Unicode MS')  # Complex scripts (Greek, Arabic, etc.) - keep Unicode support
                rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')  # East Asian scripts - keep Unicode support
                rFonts.set(qn('w:asciiTheme'), 'Times New Roman')
                rFonts.set(qn('w:hAnsiTheme'), 'Times New Roman')
                rFonts.set(qn('w:eastAsiaTheme'), 'Arial Unicode MS')
                rFonts.set(qn('w:cstheme'), 'Arial Unicode MS')
            
            # Also configure heading styles
            for i in range(1, 7):  # Heading 1 through Heading 6
                try:
                    heading_style = doc.styles[f'Heading {i}']
                    heading_font = heading_style.font
                    heading_font.name = 'Arial Unicode MS'
                    
                    # Set complex script font for headings too
                    heading_rpr = heading_style.element.xpath('.//w:rPr')[0] if heading_style.element.xpath('.//w:rPr') else None
                    if heading_rpr is not None:
                        heading_rFonts = heading_rpr.find(qn('w:rPr'))
                        if heading_rFonts is None:
                            heading_rFonts = OxmlElement("w:rPr")
                            heading_rpr.insert(0, heading_rFonts)
                        
                        heading_rFonts.set(qn('w:ascii'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:cs'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:asciiTheme'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:hAnsiTheme'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:eastAsiaTheme'), 'Arial Unicode MS')
                        heading_rFonts.set(qn('w:cstheme'), 'Arial Unicode MS')
                        
                except KeyError:
                    # Style doesn't exist, skip
                    continue
            
            # Configure TOC styles
            try:
                toc_style = doc.styles['TOC 1']
                toc_font = toc_style.font
                toc_font.name = 'Arial Unicode MS'
                
                # Set complex script font for TOC
                toc_rpr = toc_style.element.xpath('.//w:rPr')[0] if toc_style.element.xpath('.//w:rPr') else None
                if toc_rpr is not None:
                    toc_rFonts = toc_rpr.find(qn('w:rPr'))
                    if toc_rFonts is None:
                        toc_rFonts = OxmlElement("w:rPr")
                        toc_rpr.insert(0, toc_rFonts)
                    
                    toc_rFonts.set(qn('w:ascii'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:hAnsi'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:cs'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:asciiTheme'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:hAnsiTheme'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:eastAsiaTheme'), 'Arial Unicode MS')
                    toc_rFonts.set(qn('w:cstheme'), 'Arial Unicode MS')
            except KeyError:
                logger.warning("TOC style not found, skipping TOC font configuration")
            
            logger.info("Document fonts configured for Unicode support (Greek, etc.)")
            
        except Exception as e:
            logger.warning(f"Could not configure Unicode fonts: {e}")

    def _insert_visual_content(self, doc, visual_content, page_num):
        """Insert visual content into the document without translation"""
        try:
            for area in visual_content:
                bbox = area['bbox']
                content_type = area['content_type']
                confidence = area.get('confidence', 0.5)

                # Only insert if confidence is high enough
                if confidence < 0.5:
                    logger.debug(f"Skipping low confidence visual content on page {page_num + 1}")
                    continue

                # Add a section break before visual content
                doc.add_section_break()

                # Add a caption indicating this is visual content
                p = doc.add_paragraph()
                p.add_run(f"[Original Visual Content - {content_type}]").italic = True

                # Insert the visual content without translation
                if content_type == 'page_with_drawings':
                    self._insert_drawing_area(doc, bbox, page_num)
                elif content_type == 'image_area':
                    self._insert_image_area(doc, bbox, page_num)
                elif content_type == 'sparse_text_page':
                    self._insert_sparse_text_area(doc, bbox, page_num)

                # Add a section break after visual content
                doc.add_section_break()

        except Exception as e:
            logger.error(f"Error inserting visual content: {e}")
            raise

    def _insert_drawing_area(self, doc, bbox, page_num):
        """Insert a drawing area without translation"""
        try:
            # Insert the drawing as is
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(f"page_{page_num + 1}_drawing.png")
            logger.info(f"Inserted original drawing from page {page_num + 1}")

        except Exception as e:
            logger.error(f"Error inserting drawing area: {e}")
            raise

    def _insert_image_area(self, doc, bbox, page_num):
        """Insert an image area without translation"""
        try:
            # Insert the image as is
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(f"page_{page_num + 1}_image.png")
            logger.info(f"Inserted original image from page {page_num + 1}")

        except Exception as e:
            logger.error(f"Error inserting image area: {e}")
            raise

    def _insert_sparse_text_area(self, doc, bbox, page_num):
        """Insert a sparse text area without translation"""
        try:
            # Insert the content as is
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(f"page_{page_num + 1}_sparse.png")
            logger.info(f"Inserted original sparse content from page {page_num + 1}")

        except Exception as e:
            logger.error(f"Error inserting sparse text area: {e}")
            raise

    def _is_area_translated(self, bbox, page_num):
        """Check if a visual content area has been properly translated"""
        try:
            # Get the text content in this area
            text = self.pdf_parser.get_text_in_area(page_num, bbox)
            
            # If there's no text, consider it translated
            if not text.strip():
                return True

            # Check if the text has been translated
            # This assumes you have a way to track translated content
            # You might need to implement this based on your translation tracking system
            return self._is_text_translated(text, page_num)

        except Exception as e:
            logger.error(f"Error checking area translation status: {e}")
            return False

    def _is_text_translated(self, text, page_num):
        """Check if text has been translated by comparing with translation cache"""
        try:
            # Get the translation cache from the translation service
            translation_cache = self.translation_service.get_translation_cache()
            
            # Check if this text exists in the cache
            if text in translation_cache:
                return True
                
            # If not in cache, check if it's part of a larger translated segment
            for original, translated in translation_cache.items():
                if text in original:
                    return True
                    
            # If we get here, the text hasn't been translated
            logger.warning(f"Text on page {page_num + 1} has not been translated")
            return False
            
        except Exception as e:
            logger.error(f"Error checking translation status: {e}")
            return False

    def create_word_document_with_structure(self, structured_content: List[Dict[str, Any]], output_path: str, images_dir: str = None):
        """
        DEPRECATED: This method is deprecated as per Directive I.
        Use create_word_document_from_structured_document instead.
        
        This wrapper ensures backward compatibility while routing all calls through the unified method.
        """
        import warnings
        warnings.warn(
            "create_word_document_with_structure is deprecated. Use create_word_document_from_structured_document instead.",
            DeprecationWarning,
            stacklevel=2
        )
        
        self.logger.warning("⚠️ Using deprecated create_word_document_with_structure method. Please update to use create_word_document_from_structured_document.")
        
        # Route through the unified method
        return self.create_word_document_from_structured_document(structured_content, output_path, images_dir)

    def create_word_document_from_digital_twin(self, digital_twin_doc: DocumentModel, 
                                             output_path: str, target_language: str = 'el') -> bool:
        """
        Generate a Word document from Digital Twin model with structure preservation.
        
        ENHANCED: Now includes proper output directory tracking for image paths.
        """
        try:
            # Store output directory for image path resolution
            self.output_dir = os.path.dirname(output_path)
            
            # Create Word document
            doc = Document()
            self._configure_document_fonts_for_unicode(doc)
            
            # Build heading-to-page mapping for TOC generation
            heading_page_map = self._build_heading_page_map(digital_twin_doc)
            
            # Generate enhanced TOC with proper navigation
            self._generate_enhanced_toc(doc, digital_twin_doc, heading_page_map, target_language)
            
            # Process all content with bookmarks and structure preservation
            self._process_digital_twin_content_with_bookmarks(doc, digital_twin_doc, heading_page_map)
            
            # Apply document quality enhancements
            if hasattr(self, 'quality_enhancer'):
                self.quality_enhancer.enhance_document_quality(doc, digital_twin_doc)
            
            # Save document
            doc.save(output_path)
            
            self.logger.info(f"✅ Digital Twin Word document saved: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Failed to create Word document from Digital Twin: {e}")
            return False
    
    def _build_heading_page_map(self, digital_twin_doc: DocumentModel) -> Dict[str, int]:
        """
        Build a mapping of heading content to estimated page numbers in the final document.
        
        ENHANCED: Now accurately tracks Word document page structure rather than source PDF pages.
        """
        heading_page_map = {}
        current_word_page = 1
        content_on_current_page = 0
        
        # Estimate characters per page for Word document (different from PDF)
        chars_per_page = 3000  # More conservative for Word documents with larger fonts
        
        # Start with TOC pages (typically 1-2 pages)
        toc_pages = max(1, len(digital_twin_doc.toc_entries) // 15)  # ~15 entries per page
        current_word_page += toc_pages
        
        try:
            self.logger.debug(f"🔍 Building Word document page mapping (starting at page {current_word_page} after TOC)")
            
            for page in digital_twin_doc.pages:
                for block in page.get_all_blocks():
                    # Get display text (translated if available, original otherwise)
                    if hasattr(block, 'translated_text') and block.translated_text:
                        display_text = block.translated_text
                    else:
                        display_text = block.get_display_text()
                    
                    if not display_text or not display_text.strip():
                        continue
                    
                    # Check if this is a heading block
                    if hasattr(block, 'block_type') and ('heading' in str(block.block_type).lower() or 'title' in str(block.block_type).lower()):
                        # Map this heading to the current Word document page
                        heading_key = self._normalize_heading_text(display_text)
                        heading_page_map[heading_key] = current_word_page
                        
                        self.logger.debug(f"📍 Mapped heading '{display_text[:50]}...' to Word page {current_word_page}")
                    
                    # Update page estimation based on translated content length
                    content_length = len(display_text)
                    content_on_current_page += content_length
                    
                    # Check if we should move to next page
                    if content_on_current_page >= chars_per_page:
                        current_word_page += 1
                        content_on_current_page = 0
                        self.logger.debug(f"📄 Moving to Word page {current_word_page}")
            
            self.logger.info(f"📊 Built heading-to-page mapping: {len(heading_page_map)} headings across ~{current_word_page} pages")
            return heading_page_map
            
        except Exception as e:
            self.logger.error(f"Failed to build heading page map: {e}")
            return {}
    
    def _generate_enhanced_toc(self, doc: Document, digital_twin_doc: DocumentModel, 
                             heading_page_map: Dict[str, int], target_language: str) -> None:
        """
        ENHANCED: Comprehensive two-way TOC reconstruction with intelligent mapping and validation.
        
        This implements the complete two-way TOC reconstruction approach:
        1. Uses enhanced TOC entries with content mapping
        2. Validates translated titles against document content  
        3. Generates accurate page numbers through content scanning
        4. Creates functional navigation with proper bookmarks
        5. Maintains hierarchical structure and formatting
        
        This solves the user's TOC corruption problem by generating structurally accurate TOC.
        """
        try:
            self.logger.info(f"📋 Starting comprehensive two-way TOC reconstruction...")
            
            # Phase 1: Validate and prepare TOC entries
            validated_entries = self._validate_toc_entries_for_reconstruction(digital_twin_doc.toc_entries)
            
            if not validated_entries:
                self.logger.warning("⚠️ No valid TOC entries found for reconstruction")
                return
            
            # Phase 2: Generate TOC title with proper localization
            toc_title = self._get_localized_toc_title(target_language)
            title_para = doc.add_heading(toc_title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Phase 3: Add decorative separator
            self._add_toc_separator(doc)
            
            # Phase 4: Build comprehensive heading-to-page mapping
            comprehensive_page_map = self._build_comprehensive_page_mapping(digital_twin_doc, heading_page_map)
            
            # Phase 5: Generate TOC entries with intelligent reconstruction
            for entry in validated_entries:
                self._add_comprehensive_toc_entry(doc, entry, comprehensive_page_map, target_language)
            
            # Phase 6: Add TOC conclusion and spacing
            self._finalize_toc_section(doc)
            
            self.logger.info(f"✅ Comprehensive TOC reconstruction completed: {len(validated_entries)} entries generated")
            
        except Exception as e:
            self.logger.error(f"❌ Failed to generate comprehensive TOC: {e}")
    
    def _validate_toc_entries_for_reconstruction(self, toc_entries: List) -> List:
        """
        Validate and filter TOC entries for reliable reconstruction.
        """
        validated_entries = []
        
        for entry in toc_entries:
            # Check if entry has minimum required data
            if not hasattr(entry, 'title') or not entry.title:
                self.logger.warning(f"⚠️ Skipping TOC entry with missing title")
                continue
            
            # Check if entry has reasonable level
            if not hasattr(entry, 'level') or entry.level < 1 or entry.level > 6:
                self.logger.warning(f"⚠️ Fixing invalid level for TOC entry: {entry.title[:30]}...")
                entry.level = 1  # Default to level 1
            
            # Ensure entry has necessary attributes
            if not hasattr(entry, 'translated_title'):
                entry.translated_title = entry.title
            
            if not hasattr(entry, 'processing_notes'):
                entry.processing_notes = []
            
            if not hasattr(entry, 'confidence_score'):
                entry.confidence_score = 1.0
            
            # Add validation note
            entry.processing_notes.append("Validated for comprehensive reconstruction")
            validated_entries.append(entry)
        
        self.logger.info(f"📋 Validated {len(validated_entries)} TOC entries for reconstruction")
        return validated_entries
    
    def _get_localized_toc_title(self, target_language: str) -> str:
        """
        Get properly localized TOC title based on target language.
        """
        toc_titles = {
            'el': 'Πίνακας Περιεχομένων',
            'en': 'Table of Contents',
            'es': 'Índice de Contenidos',
            'fr': 'Table des Matières',
            'de': 'Inhaltsverzeichnis',
            'it': 'Indice dei Contenuti',
            'pt': 'Índice de Conteúdos',
            'ru': 'Содержание',
            'zh': '目录',
            'ja': '目次'
        }
        return toc_titles.get(target_language, 'Table of Contents')
    
    def _add_toc_separator(self, doc: Document) -> None:
        """
        Add decorative separator below TOC title.
        """
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("─" * 60)
        line_run.font.color.rgb = RGBColor(100, 100, 100)
        line_run.font.size = Pt(10)
        
        # Add spacing
        line_para.paragraph_format.space_after = Pt(12)
    
    def _build_comprehensive_page_mapping(self, digital_twin_doc: DocumentModel, 
                                        existing_page_map: Dict[str, int]) -> Dict[str, int]:
        """
        Build comprehensive page mapping using multiple strategies.
        """
        comprehensive_map = existing_page_map.copy()
        
        # Strategy 1: Use TOC entries with reliable content mapping
        for toc_entry in digital_twin_doc.toc_entries:
            if hasattr(toc_entry, 'has_reliable_mapping') and toc_entry.has_reliable_mapping():
                # Use the mapped page if available
                if hasattr(toc_entry, 'translated_page_in_document') and toc_entry.translated_page_in_document > 0:
                    title_key = self._normalize_heading_text(toc_entry.get_display_title())
                    comprehensive_map[title_key] = toc_entry.translated_page_in_document
                elif hasattr(toc_entry, 'original_page_in_document') and toc_entry.original_page_in_document > 0:
                    title_key = self._normalize_heading_text(toc_entry.get_display_title())
                    comprehensive_map[title_key] = toc_entry.original_page_in_document
        
        # Strategy 2: Scan document content for additional heading matches
        additional_mappings = self._scan_document_for_heading_pages(digital_twin_doc)
        comprehensive_map.update(additional_mappings)
        
        self.logger.info(f"📊 Built comprehensive page mapping: {len(comprehensive_map)} entries")
        return comprehensive_map
    
    def _scan_document_for_heading_pages(self, digital_twin_doc: DocumentModel) -> Dict[str, int]:
        """
        Scan document pages to find additional heading-to-page mappings.
        """
        heading_page_map = {}
        
        try:
            for page in digital_twin_doc.pages:
                page_num = page.page_number
                
                # Look for heading-style text blocks
                for text_block in page.text_blocks:
                    # Check if this looks like a heading
                    if self._is_heading_style_block(text_block):
                        text = text_block.get_display_text()
                        normalized_text = self._normalize_heading_text(text)
                        
                        # Only add if not already mapped
                        if normalized_text not in heading_page_map:
                            heading_page_map[normalized_text] = page_num
                            
                            self.logger.debug(f"📄 Mapped heading '{text[:30]}...' to page {page_num}")
            
            self.logger.info(f"🔍 Content scan found {len(heading_page_map)} additional heading mappings")
            return heading_page_map
            
        except Exception as e:
            self.logger.error(f"❌ Document content scan failed: {e}")
            return {}
    
    def _is_heading_style_block(self, text_block) -> bool:
        """
        Determine if a text block appears to be a heading.
        """
        try:
            # Check block type
            if hasattr(text_block, 'block_type'):
                block_type_str = str(text_block.block_type).lower()
                if 'heading' in block_type_str or 'title' in block_type_str:
                    return True
            
            # Check structural role
            if hasattr(text_block, 'structural_role'):
                role_str = str(text_block.structural_role).lower()
                if 'navigation' in role_str:
                    return True
            
            # Check formatting characteristics
            if hasattr(text_block, 'font_size') and text_block.font_size > 12:
                if hasattr(text_block, 'is_bold') and text_block.is_bold:
                    return True
            
            # Check content characteristics
            text = text_block.get_display_text()
            if len(text) < 200 and not text.endswith('.'):  # Short, not sentence-like
                return True
            
            return False
            
        except Exception:
            return False
    
    def _add_comprehensive_toc_entry(self, doc: Document, toc_entry, 
                                   comprehensive_page_map: Dict[str, int], target_language: str) -> None:
        """
        Add a single TOC entry with comprehensive page mapping and enhanced formatting.
        """
        try:
            # Get the display title (translated if available)
            display_title = toc_entry.get_display_title()
            if not display_title:
                self.logger.warning(f"TOC entry has no display title, skipping")
                return
            
            # Normalize the title for mapping lookup
            normalized_title = self._normalize_heading_text(display_title)
            
            # Try to get accurate page number using multiple strategies
            page_number = None
            
            # Strategy 1: Use comprehensive page mapping (most accurate)
            if normalized_title in comprehensive_page_map:
                page_number = comprehensive_page_map[normalized_title]
                self.logger.debug(f"📍 Found page {page_number} for '{display_title[:30]}...' via comprehensive mapping")
            
            # Strategy 2: Look for partial matches in the mapping
            if page_number is None:
                for mapped_title, mapped_page in comprehensive_page_map.items():
                    # Check if titles are similar (handle translation variations)
                    if self._titles_are_similar(normalized_title, mapped_title):
                        page_number = mapped_page
                        self.logger.debug(f"📍 Found page {page_number} for '{display_title[:30]}...' via similar title match")
                        break
            
            # Strategy 3: Use TOC entry's own page information (adjusted for Word document)
            if page_number is None:
                # Convert PDF page to estimated Word page
                if hasattr(toc_entry, 'page_number') and toc_entry.page_number > 0:
                    # Adjust for TOC pages at beginning and different page layouts
                    toc_pages = max(1, len(doc.paragraphs) // 100)  # Rough estimate
                    estimated_page = toc_entry.page_number + toc_pages
                    page_number = max(1, estimated_page)
                    self.logger.debug(f"📍 Estimated page {page_number} for '{display_title[:30]}...' from PDF page {toc_entry.page_number}")
            
            # Strategy 4: Fallback to sequential numbering
            if page_number is None:
                # Count existing TOC entries to estimate position
                existing_entries = len([p for p in doc.paragraphs if p.style.name.startswith('TOC')])
                page_number = max(3, existing_entries + 2)  # Start after TOC itself
                self.logger.debug(f"📍 Fallback page {page_number} for '{display_title[:30]}...' via sequential numbering")
            
            # Get indentation level
            level = getattr(toc_entry, 'level', 1)
            indent_level = max(0, level - 1)  # Convert to 0-based indentation
            
            # Create TOC entry paragraph
            toc_para = doc.add_paragraph()
            
            # Add indentation based on level
            if indent_level > 0:
                indent_spaces = "  " * indent_level  # 2 spaces per level
                toc_para.add_run(indent_spaces)
            
            # Add the title
            title_run = toc_para.add_run(display_title)
            
            # Add dots and page number
            dots_run = toc_para.add_run("." * max(3, 50 - len(display_title) - len(str(page_number))))
            page_run = toc_para.add_run(f" {page_number}")
            
            # Style the TOC entry based on level
            if level == 1:
                title_run.bold = True
                title_run.font.size = Pt(12)
            elif level == 2:
                title_run.font.size = Pt(11)
            else:
                title_run.font.size = Pt(10)
            
            # Style the dots and page number
            dots_run.font.size = Pt(8)
            page_run.bold = True
            page_run.font.size = Pt(10)
            
            self.logger.debug(f"✅ Added TOC entry: {display_title[:30]}... → page {page_number}")
            
        except Exception as e:
            self.logger.error(f"❌ Failed to add TOC entry '{getattr(toc_entry, 'title', 'unknown')}': {e}")

    def _titles_are_similar(self, title1: str, title2: str, threshold: float = 0.7) -> bool:
        """
        Check if two titles are similar enough to be considered the same heading.
        Handles translation variations and minor differences.
        """
        try:
            # Normalize both titles
            norm1 = title1.lower().strip()
            norm2 = title2.lower().strip()
            
            # Direct match
            if norm1 == norm2:
                return True
            
            # Check if one is a substring of the other
            if norm1 in norm2 or norm2 in norm1:
                return True
            
            # Simple similarity check based on common words
            words1 = set(norm1.split())
            words2 = set(norm2.split())
            
            if not words1 or not words2:
                return False
            
            # Calculate Jaccard similarity
            intersection = len(words1.intersection(words2))
            union = len(words1.union(words2))
            
            similarity = intersection / union if union > 0 else 0
            return similarity >= threshold
            
        except Exception:
            return False
    
    def _deduce_accurate_page_number(self, toc_entry, page_map: Dict[str, int], display_title: str) -> int:
        """
        Deduce the most accurate page number using multiple strategies.
        """
        # Strategy 1: Direct mapping from display title
        normalized_title = self._normalize_heading_text(display_title)
        if normalized_title in page_map:
            return page_map[normalized_title]
        
        # Strategy 2: Use TOC entry's mapped page information
        if hasattr(toc_entry, 'translated_page_in_document') and toc_entry.translated_page_in_document > 0:
            return toc_entry.translated_page_in_document
        
        if hasattr(toc_entry, 'original_page_in_document') and toc_entry.original_page_in_document > 0:
            return toc_entry.original_page_in_document
        
        # Strategy 3: Fuzzy matching with existing page map
        best_match_page = self._find_fuzzy_page_match(display_title, page_map)
        if best_match_page > 0:
            return best_match_page
        
        # Strategy 4: Use original TOC page number as fallback
        if hasattr(toc_entry, 'page_number') and toc_entry.page_number > 0:
            return toc_entry.page_number
        
        # Strategy 5: Estimate based on position in TOC
        estimated_page = self._estimate_page_by_toc_position(toc_entry)
        
        return max(1, estimated_page)  # Ensure page number is at least 1
    
    def _find_fuzzy_page_match(self, display_title: str, page_map: Dict[str, int]) -> int:
        """
        Find page number using fuzzy matching of titles.
        """
        normalized_display = self._normalize_heading_text(display_title)
        best_score = 0
        best_page = 0
        
        for mapped_title, page_num in page_map.items():
            # Calculate similarity score
            score = self._calculate_title_similarity(normalized_display, mapped_title)
            
            if score > best_score and score >= 0.7:  # Minimum threshold
                best_score = score
                best_page = page_num
        
        if best_page > 0:
            self.logger.debug(f"🔍 Fuzzy match: '{display_title[:30]}...' → Page {best_page} (score: {best_score:.2f})")
        
        return best_page
    
    def _calculate_title_similarity(self, title1: str, title2: str) -> float:
        """
        Calculate similarity score between two titles.
        """
        if not title1 or not title2:
            return 0.0
        
        # Exact match
        if title1 == title2:
            return 1.0
        
        # Word-based similarity
        words1 = set(title1.lower().split())
        words2 = set(title2.lower().split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = len(words1.intersection(words2))
        union = len(words1.union(words2))
        
        jaccard_score = intersection / union if union > 0 else 0.0
        
        # Boost score for substring matches
        if title1 in title2 or title2 in title1:
            jaccard_score = min(1.0, jaccard_score + 0.2)
        
        return jaccard_score
    
    def _estimate_page_by_toc_position(self, toc_entry) -> int:
        """
        Estimate page number based on TOC entry position and level.
        """
        # Simple estimation: use level and entry ID to estimate
        base_page = 1
        
        if hasattr(toc_entry, 'entry_id'):
            # Extract number from entry ID if possible
            import re
            numbers = re.findall(r'\d+', toc_entry.entry_id)
            if numbers:
                entry_num = int(numbers[-1])  # Use last number found
                # Estimate: assume ~2-5 pages per major section
                pages_per_section = 3
                base_page = 1 + (entry_num - 1) * pages_per_section
        
        # Adjust based on level (deeper levels are typically later in document)
        level_adjustment = max(0, toc_entry.level - 1) * 2
        estimated_page = base_page + level_adjustment
        
        return max(1, estimated_page)
    
    def _calculate_dots_count(self, title: str, page_number: int, level: int) -> int:
        """
        Calculate appropriate number of dots for TOC entry.
        """
        # Base calculation on title length and page number length
        title_length = len(title)
        page_str_length = len(str(page_number))
        
        # Target total line length (adjusted for level)
        target_length = 65 - (level - 1) * 5  # Shorter lines for deeper levels
        
        # Calculate available space for dots
        used_space = title_length + page_str_length + 2  # +2 for spaces
        available_space = max(10, target_length - used_space)
        
        return min(50, max(10, available_space))  # Between 10 and 50 dots
    
    def _create_toc_bookmark_name(self, display_title: str, toc_entry) -> str:
        """
        Create a unique bookmark name for TOC entry.
        """
        # Sanitize title for bookmark name
        sanitized = re.sub(r'[^\w\s]', '', display_title)
        sanitized = re.sub(r'\s+', '_', sanitized.strip())
        
        # Add level and entry ID for uniqueness
        level = toc_entry.level
        entry_id = getattr(toc_entry, 'entry_id', 'unknown')
        
        bookmark_name = f"TOC_{level}_{sanitized}_{entry_id}"
        
        # Ensure reasonable length
        if len(bookmark_name) > 40:
            bookmark_name = bookmark_name[:37] + "..."
        
        return bookmark_name
    
    def _normalize_heading_text(self, text: str) -> str:
        """
        Normalize heading text for consistent matching between TOC and content.
        """
        if not text:
            return ""
        
        # Remove special characters, normalize whitespace, convert to lowercase
        import re
        normalized = re.sub(r'[^\w\s]', '', text.lower())
        normalized = re.sub(r'\s+', ' ', normalized).strip()
        return normalized
    
    def _finalize_toc_section(self, doc: Document) -> None:
        """
        Add final spacing and formatting to complete TOC section.
        """
        # Add spacing after TOC
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_after = Pt(24)
        
        # Optional: Add page break after TOC
        # doc.add_page_break()
        
        self.logger.debug("📄 Finalized TOC section formatting")
    
    def _add_enhanced_toc_entry(self, doc: Document, toc_entry, 
                              heading_page_map: Dict[str, int], target_language: str) -> None:
        """
        LEGACY METHOD: Redirected to comprehensive implementation.
        """
        # Redirect to the new comprehensive method
        self._add_comprehensive_toc_entry(doc, toc_entry, heading_page_map, target_language)
    
    def _process_digital_twin_content_with_bookmarks(self, doc: Document, 
                                                   digital_twin_doc: DocumentModel,
                                                   heading_page_map: Dict[str, int]) -> None:
        """
        Process Digital Twin content while creating proper bookmarks for TOC linking.
        ENHANCED: Now ensures ALL content is processed without skipping blocks or cutting off pages.
        """
        try:
            total_blocks_processed = 0
            skipped_blocks = 0
            final_output_blocks = []  # For debug export
            for page_index, page in enumerate(digital_twin_doc.pages):
                self.logger.info(f"📄 Processing page {page.page_number} with {len(page.get_all_blocks())} blocks")
                all_blocks = page.get_all_blocks()
                if not all_blocks:
                    self.logger.warning(f"⚠️ Page {page.page_number} has no content blocks, skipping")
                    continue
                sorted_blocks = sorted(all_blocks, key=lambda block: block.bbox[1] if block.bbox else 0)
                if page_index > 0:
                    doc.add_page_break()
                    self.logger.debug(f"📄 Added page break before page {page.page_number}")
                blocks_on_page = 0
                for block_index, block in enumerate(sorted_blocks):
                    try:
                        block_type = str(getattr(block, 'block_type', 'text')).lower()
                        block_id = getattr(block, 'block_id', f'block_{page_index}_{block_index}')
                        if hasattr(block, 'translated_text') and block.translated_text:
                            display_text = block.translated_text
                        else:
                            display_text = block.get_display_text()
                        # --- DEBUG: Record output block order ---
                        final_output_blocks.append({
                            'block_id': block_id,
                            'block_type': block_type,
                            'text': display_text
                        })
                        # --- END DEBUG ---
                        # ... existing code for processing block ...
                        processed = False
                        if 'heading' in block_type or 'title' in block_type:
                            self._process_digital_twin_heading_with_bookmark(doc, block)
                            processed = True
                        elif 'text' in block_type or 'paragraph' in block_type:
                            self._process_digital_twin_text_block_with_bookmark(doc, block)
                            processed = True
                        elif 'image' in block_type:
                            try:
                                if hasattr(block, 'image_path') and block.image_path:
                                    self._process_digital_twin_image_block(doc, block)
                                    processed = True
                                else:
                                    placeholder_para = doc.add_paragraph(f"[Image placeholder: {block_id}]")
                                    placeholder_para.italic = True
                                    processed = True
                            except Exception as img_error:
                                self.logger.warning(f"⚠️ Image processing failed for {block_id}: {img_error}")
                                error_para = doc.add_paragraph(f"[Image error: {block_id}]")
                                error_para.italic = True
                                processed = True
                        elif 'table' in block_type:
                            self._process_digital_twin_table_block(doc, block)
                            processed = True
                        else:
                            self._process_digital_twin_text_block_with_bookmark(doc, block)
                            processed = True
                        if processed:
                            blocks_on_page += 1
                            self.logger.debug(f"✅ Processed block {block_id}")
                        else:
                            self.logger.error(f"❌ Failed to process block {block_id}")
                            skipped_blocks += 1
                    except Exception as block_error:
                        self.logger.error(f"❌ Exception processing block {getattr(block, 'block_id', 'unknown')} on page {page.page_number}: {block_error}")
                        try:
                            error_para = doc.add_paragraph(f"[Processing error for block {getattr(block, 'block_id', 'unknown')}]")
                            error_para.italic = True
                            blocks_on_page += 1
                        except:
                            pass
                        skipped_blocks += 1
                        continue
                total_blocks_processed += blocks_on_page
                self.logger.info(f"✅ Page {page.page_number}: {blocks_on_page}/{len(sorted_blocks)} blocks processed successfully")
            # --- DEBUG: Write final output order ---
            output_dir = getattr(self, 'output_dir', '.')
            output_name = getattr(self, 'output_name', 'digital_twin_output')
            debug_path = os.path.join(output_dir, f"{output_name}_final_output_order.json")
            import json
            with open(debug_path, 'w', encoding='utf-8') as f:
                json.dump(final_output_blocks, f, ensure_ascii=False, indent=2)
            self.logger.info(f"🪪 Final output block order exported to: {debug_path}")
            # --- END DEBUG ---
            self.logger.info(f"📄 Document generation completed:")
            self.logger.info(f"   • Total pages: {len(digital_twin_doc.pages)}")  
            self.logger.info(f"   • Total blocks processed: {total_blocks_processed}")
            self.logger.info(f"   • Blocks skipped: {skipped_blocks}")
            if skipped_blocks > 0:
                self.logger.warning(f"⚠️ {skipped_blocks} blocks were skipped due to processing errors")
        except Exception as e:
            self.logger.error(f"❌ Critical failure in document processing: {e}")
            self.logger.warning("⚠️ Continuing with partial document generation")
    
    def _process_digital_twin_heading_with_bookmark(self, doc: Document, heading_block) -> None:
        """
        Process heading block and create proper bookmark for TOC linking.
        """
        try:
            # Get heading text (translated if available)
            if hasattr(heading_block, 'translated_text') and heading_block.translated_text:
                heading_text = heading_block.translated_text
            else:
                heading_text = heading_block.get_display_text()
            
            if not heading_text:
                return
            
            # Get heading level
            level = getattr(heading_block, 'level', 1)
            
            # Create heading paragraph
            heading_para = doc.add_heading(sanitize_for_xml(heading_text), level=level)
            
            # Create bookmark for TOC linking
            normalized_text = self._normalize_heading_text(heading_text)
            if hasattr(self, 'toc_bookmark_map') and normalized_text in self.toc_bookmark_map:
                bookmark_info = self.toc_bookmark_map[normalized_text]
                bookmark_name = bookmark_info['bookmark_name']
                
                # Add bookmark to heading
                self._add_bookmark_to_paragraph(heading_para, bookmark_name)
                
                self.logger.debug(f"🔗 Created bookmark '{bookmark_name}' for heading: {heading_text[:50]}...")
            
        except Exception as e:
            self.logger.error(f"Failed to process heading with bookmark: {e}")
    
    def _add_bookmark_to_paragraph(self, paragraph, bookmark_name: str) -> None:
        """
        Add a bookmark to a paragraph for TOC linking.
        """
        try:
            # Create bookmark start and end elements
            bookmark_start = OxmlElement('w:bookmarkStart')
            bookmark_start.set(qn('w:id'), str(self.bookmark_id))
            bookmark_start.set(qn('w:name'), bookmark_name)
            
            bookmark_end = OxmlElement('w:bookmarkEnd')
            bookmark_end.set(qn('w:id'), str(self.bookmark_id))
            
            # Insert bookmark elements
            paragraph._element.insert(0, bookmark_start)
            paragraph._element.append(bookmark_end)
            
            self.bookmark_id += 1
            
        except Exception as e:
            self.logger.error(f"Failed to add bookmark: {e}")
    
    def _process_digital_twin_text_block_with_bookmark(self, doc: Document, text_block) -> None:
        """
        Process Digital Twin text blocks with proper formatting and translation awareness.
        
        ENHANCED: Skip footnotes here as they are processed separately in _process_footnotes_section.
        ENHANCED: Improved bookmark creation for headings to match TOC entries.
        """
        try:
            # CRITICAL: Skip footnotes - they are processed separately
            if text_block.block_type == BlockType.FOOTNOTE:
                self.logger.debug(f"📝 Skipping footnote in main text processing: {text_block.get_display_text()[:50]}...")
                return
            
            # Get display text (prefer translation if available)
            display_text = text_block.get_display_text(prefer_translation=True)
            
            if not display_text.strip():
                return  # Skip empty blocks
            
            safe_text = sanitize_for_xml(display_text)
            
            # Handle different block types with proper formatting
            if text_block.block_type == BlockType.HEADING:
                level = text_block.heading_level or 1
                p = doc.add_heading(safe_text, level=level)
                
                # Create consistent bookmark for TOC linking
                self._create_heading_bookmark(p, safe_text, text_block)
                
            elif text_block.block_type == BlockType.TITLE:
                p = doc.add_heading(safe_text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Create bookmark for title as well
                self._create_heading_bookmark(p, safe_text, text_block)
                
            elif text_block.block_type == BlockType.PARAGRAPH:
                p = doc.add_paragraph(safe_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justify paragraphs for better appearance
                
            elif text_block.block_type == BlockType.LIST_ITEM:
                p = doc.add_paragraph(safe_text)
                p.style = 'List Bullet'
                if hasattr(text_block, 'list_level') and text_block.list_level and text_block.list_level > 1:
                    p.paragraph_format.left_indent = Inches(0.25 * text_block.list_level)
                
            elif text_block.block_type == BlockType.QUOTE:
                p = doc.add_paragraph(safe_text)
                p.style = 'Quote'
                
            else:  # Default to paragraph
                p = doc.add_paragraph(safe_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
            # Apply font formatting if available
            if hasattr(text_block, 'font_family') and text_block.font_family:
                for run in p.runs:
                    run.font.name = text_block.font_family
            
            if hasattr(text_block, 'font_size') and text_block.font_size > 0:
                for run in p.runs:
                    run.font.size = Pt(text_block.font_size)
                    
        except Exception as e:
            self.logger.error(f"❌ Failed to process text block: {e}")
    
    def _create_heading_bookmark(self, heading_paragraph, heading_text: str, text_block) -> None:
        """
        Create a bookmark for a heading that can be referenced by TOC entries.
        
        This ensures TOC links work properly and prevents "Error! Bookmark not defined" issues.
        """
        try:
            # Generate a consistent bookmark name
            # Use text content to match with TOC entries
            clean_text = re.sub(r'[^\w\s]', '', heading_text).strip()
            bookmark_name = f"_Heading_{clean_text.replace(' ', '_')}"[:40]  # Limit length
            
            # Ensure bookmark name is unique
            if not hasattr(self, '_bookmark_counter'):
                self._bookmark_counter = 0
            self._bookmark_counter += 1
            bookmark_name = f"{bookmark_name}_{self._bookmark_counter}"
            
            # Create bookmark elements
            tag_start = OxmlElement('w:bookmarkStart')
            tag_start.set(qn('w:id'), str(self._bookmark_counter))
            tag_start.set(qn('w:name'), bookmark_name)
            
            tag_end = OxmlElement('w:bookmarkEnd')
            tag_end.set(qn('w:id'), str(self._bookmark_counter))
            
            # Insert bookmark around the heading
            heading_paragraph._p.insert(0, tag_start)
            heading_paragraph._p.append(tag_end)
            
            # Try to update corresponding TOC entry
            self._update_toc_entry_bookmark(heading_text, bookmark_name)
            
            self.logger.debug(f"📑 Created bookmark '{bookmark_name}' for heading: {heading_text[:30]}...")
            
        except Exception as e:
            self.logger.warning(f"Failed to create bookmark for heading '{heading_text}': {e}")
    
    def _update_toc_entry_bookmark(self, heading_text: str, bookmark_name: str) -> None:
        """
        Update TOC entry to use the actual bookmark created for the heading.
        
        This creates the link between TOC entries and their corresponding headings.
        """
        try:
            if not hasattr(self, 'pending_toc_entries'):
                return
            
            # Find matching TOC entry
            for toc_info in self.pending_toc_entries:
                # Match by title (allowing for translation differences)
                if (self._titles_match(toc_info['title'], heading_text) or 
                    self._titles_match(toc_info['original_title'], heading_text)):
                    
                    # Update the TOC paragraph to include hyperlink
                    self._convert_toc_to_hyperlink(toc_info, bookmark_name)
                    break
            
        except Exception as e:
            self.logger.warning(f"Failed to update TOC entry bookmark: {e}")
    
    def _titles_match(self, title1: str, title2: str) -> bool:
        """
        Check if two titles match, allowing for minor differences.
        """
        if not title1 or not title2:
            return False
        
        # Normalize titles for comparison
        t1 = re.sub(r'[^\w\s]', '', title1.lower()).strip()
        t2 = re.sub(r'[^\w\s]', '', title2.lower()).strip()
        
        # Exact match
        if t1 == t2:
            return True
        
        # Partial match (useful for translations)
        if len(t1) > 10 and len(t2) > 10:
            # Check if one contains the other (allowing for translation differences)
            return t1 in t2 or t2 in t1
        
        return False
    
    def _convert_toc_to_hyperlink(self, toc_info: dict, bookmark_name: str) -> None:
        """
        Convert a TOC entry to include a hyperlink to the bookmark.
        
        This replaces the static text with a functional hyperlink.
        """
        try:
            paragraph = toc_info['paragraph']
            
            # Clear existing content
            paragraph.clear()
            
            # Recreate with hyperlink
            title_run = self._add_hyperlink(paragraph, toc_info['title'], bookmark_name)
            if toc_info['level'] == 1:
                title_run.bold = True
                title_run.font.size = Pt(12)
            elif toc_info['level'] == 2:
                title_run.font.size = Pt(11)
            else:
                title_run.font.size = Pt(10)
                title_run.italic = True
            
            # Add dot leaders
            dots_length = max(3, 60 - len(toc_info['title']) - len(str(toc_info['page_number'])))
            dots_run = paragraph.add_run(" " + "." * dots_length + " ")
            dots_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add page number
            page_run = paragraph.add_run(str(toc_info['page_number']))
            page_run.bold = (toc_info['level'] == 1)
            
            self.logger.debug(f"🔗 Converted TOC entry to hyperlink: {toc_info['title']} -> {bookmark_name}")
            
        except Exception as e:
            self.logger.warning(f"Failed to convert TOC entry to hyperlink: {e}")

    def _process_digital_twin_image_block(self, doc: Document, image_block: ImageBlock) -> None:
        """
        Process Digital Twin image block with enhanced path resolution and insertion.
        
        ENHANCED: Improved image path resolution and error handling for reliable image insertion.
        """
        try:
            self.logger.info(f"🖼️ Processing image block: {getattr(image_block, 'block_id', 'unknown')}")
            
            if not image_block.image_path:
                self.logger.warning(f"❌ Image block has no image path specified: {image_block}")
                return
            
            self.logger.debug(f"📍 Original image path: {image_block.image_path}")
            
            # Enhanced path resolution logic
            resolved_path = self._resolve_image_path(image_block.image_path)
            self.logger.debug(f"📍 Resolved image path: {resolved_path}")
            
            if not resolved_path or not os.path.exists(resolved_path):
                self.logger.warning(f"❌ Image not found: {image_block.image_path}")
                self.logger.debug(f"   Tried to resolve to: {resolved_path}")
                
                # List files in common directories for debugging
                common_dirs = ['images', 'output', 'fenixoutput', '.']
                for dir_name in common_dirs:
                    if os.path.exists(dir_name):
                        files = os.listdir(dir_name)
                        image_files = [f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
                        if image_files:
                            self.logger.debug(f"   Available images in {dir_name}: {image_files[:5]}")  # Show first 5
                
                # Add placeholder instead of failing silently
                self._add_image_placeholder(doc, image_block)
                return
            
            try:
                # Validate image file
                if not self._is_valid_image_file(resolved_path):
                    self.logger.warning(f"❌ Invalid image file: {resolved_path}")
                    self._add_image_placeholder(doc, image_block)
                    return
                
                # Create paragraph for image
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Calculate image size based on bbox if available
                width_inches = self._calculate_image_width(image_block)
                self.logger.debug(f"📏 Calculated image width: {width_inches} inches")
                
                # Add image with size constraints
                run = p.add_run()
                run.add_picture(resolved_path, width=Inches(width_inches))
                
                # Add caption if available
                if hasattr(image_block, 'caption_text') and image_block.caption_text:
                    self._add_image_caption(doc, image_block.caption_text)
                
                self.logger.info(f"📸 ✅ Successfully inserted image: {os.path.basename(resolved_path)}")
                
            except Exception as e:
                self.logger.error(f"❌ Failed to insert image {resolved_path}: {e}")
                # More detailed error information
                self.logger.error(f"   Image path exists: {os.path.exists(resolved_path)}")
                if os.path.exists(resolved_path):
                    self.logger.error(f"   Image file size: {os.path.getsize(resolved_path)} bytes")
                self._add_image_placeholder(doc, image_block)
                
        except Exception as e:
            self.logger.error(f"❌ Failed to process image block: {e}")
            self._add_image_placeholder(doc, image_block)
    
    def _resolve_image_path(self, image_path: str) -> Optional[str]:
        """
        Resolve image path with multiple fallback strategies.
        
        This method tries various path resolution strategies to find the actual image file.
        """
        if not image_path:
            return None
        
        # Strategy 1: Use path as-is if it's absolute and exists
        if os.path.isabs(image_path) and os.path.exists(image_path):
            return image_path
        
        # Strategy 2: Try relative to current working directory
        if os.path.exists(image_path):
            return os.path.abspath(image_path)
        
        # Strategy 3: Try relative to output directory (if available)
        if hasattr(self, 'output_dir') and self.output_dir:
            output_relative = os.path.join(self.output_dir, image_path)
            if os.path.exists(output_relative):
                return output_relative
        
        # Strategy 4: Try relative to common image directories
        common_dirs = ['images', 'output', 'fenixoutput', '.']
        for base_dir in common_dirs:
            if os.path.exists(base_dir):
                candidate_path = os.path.join(base_dir, os.path.basename(image_path))
                if os.path.exists(candidate_path):
                    return candidate_path
        
        # Strategy 5: Search for image by filename in current directory tree
        filename = os.path.basename(image_path)
        for root, dirs, files in os.walk('.'):
            if filename in files:
                candidate_path = os.path.join(root, filename)
                if os.path.exists(candidate_path):
                    return candidate_path
        
        return None
    
    def _is_valid_image_file(self, file_path: str) -> bool:
        """
        Validate that the file is a valid image that can be inserted into Word document.
        """
        try:
            # Check file extension
            valid_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif'}
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext not in valid_extensions:
                return False
            
            # Check file size (should be > 0)
            if os.path.getsize(file_path) == 0:
                return False
            
            # Try to open with PIL to validate image format
            try:
                from PIL import Image
                with Image.open(file_path) as img:
                    img.verify()
                return True
            except ImportError:
                # PIL not available, rely on extension check
                return True
            except Exception:
                return False
                
        except Exception:
            return False
    
    def _calculate_image_width(self, image_block: ImageBlock) -> float:
        """
        Calculate appropriate image width in inches based on bbox and constraints.
        """
        default_width = 4.0  # Default 4 inches
        max_width = 6.0      # Maximum 6 inches
        
        if image_block.bbox and len(image_block.bbox) >= 4:
            # Convert points to inches (72 points = 1 inch)
            bbox_width = image_block.bbox[2] - image_block.bbox[0]
            if bbox_width > 0:
                width_inches = bbox_width / 72
                return min(width_inches, max_width)
        
        return default_width
    
    def _add_image_placeholder(self, doc: Document, image_block: ImageBlock) -> None:
        """
        Add a placeholder when image cannot be inserted.
        """
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        placeholder_text = f"[Image: {os.path.basename(image_block.image_path or 'unknown')}]"
        run = p.add_run(placeholder_text)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add caption if available
        if image_block.caption_text:
            self._add_image_caption(doc, image_block.caption_text)
    
    def _add_image_caption(self, doc: Document, caption_text: str) -> None:
        """
        Add image caption with proper formatting.
        """
        caption_p = doc.add_paragraph()
        caption_run = caption_p.add_run(f"Figure: {sanitize_for_xml(caption_text)}")
        caption_run.font.italic = True
        caption_run.font.size = Pt(10)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _process_digital_twin_table_block(self, doc: Document, table_block: TableBlock) -> None:
        """Process Digital Twin table blocks with proper structure"""
        try:
            if not table_block.rows:
                self.logger.warning("Empty table block, skipping")
                return
            
            # Create Word table
            word_table = doc.add_table(rows=table_block.num_rows, cols=table_block.num_cols)
            word_table.style = 'Table Grid'
            
            # Populate table cells
            for row_idx, row_data in enumerate(table_block.rows):
                table_row = word_table.rows[row_idx]
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(table_row.cells):
                        cell = table_row.cells[col_idx]
                        cell.text = sanitize_for_xml(str(cell_data))
                        
                        # Make header row bold
                        if row_idx == 0 and table_block.has_header:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
            
            self.logger.info(f"📊 Successfully inserted table: {table_block.num_rows}x{table_block.num_cols}")
            
        except Exception as e:
            self.logger.error(f"❌ Failed to process table block: {e}")
    
    def _process_metadata_blocks(self, doc: Document, metadata_blocks: List) -> None:
        """
        Process metadata blocks (headers/footers) by placing them in correct document sections.
        
        This implements the user's requirement for proper header/footer placement
        rather than mixing them with body content.
        """
        try:
            if not metadata_blocks:
                return
            
            # TODO: Implement proper header/footer section handling
            # For now, add as regular paragraphs with special formatting
            for block in metadata_blocks:
                if isinstance(block, TextBlock):
                    display_text = block.get_display_text(prefer_translation=True)
                    if display_text.strip():
                        p = doc.add_paragraph(sanitize_for_xml(display_text))
                        run = p.runs[0] if p.runs else p.add_run()
                        run.font.size = Pt(9)
                        run.font.italic = True
                        
                        if block.block_type == BlockType.HEADER:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif block.block_type == BlockType.FOOTER:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            self.logger.error(f"❌ Failed to process metadata blocks: {e}")

    def _process_digital_twin_block(self, doc: Document, block) -> None:
        """
        Process individual Digital Twin blocks with type-aware handling.
        
        This implements the user's block-aware processing logic.
        """
        try:
            # Handle Text Blocks
            if isinstance(block, TextBlock):
                self._process_digital_twin_text_block(doc, block)
            
            # Handle Image Blocks with proper file linking
            elif isinstance(block, ImageBlock):
                self._process_digital_twin_image_block(doc, block)
            
            # Handle Table Blocks
            elif isinstance(block, TableBlock):
                self._process_digital_twin_table_block(doc, block)
            
            else:
                self.logger.warning(f"Unknown Digital Twin block type: {type(block)}")
                
        except Exception as e:
            self.logger.error(f"❌ Failed to process Digital Twin block {getattr(block, 'block_id', 'unknown')}: {e}")

class DocumentQualityEnhancer:
    """
    Comprehensive document quality enhancement system that addresses all formatting issues
    identified in the user's analysis of Greek text reconstruction problems.
    """
    
    def __init__(self, logger=None):
        self.logger = logger or logging.getLogger(__name__)
        self.bookmark_registry = {}  # Track all bookmarks for consistency
        self.placeholder_registry = {}  # Track placeholder codes for restoration
        self.image_registry = {}  # Track image paths for proper insertion
        self.formatting_issues = []  # Collect formatting issues for reporting
    
    def enhance_document_quality(self, doc, digital_twin_doc=None):
        """Apply comprehensive quality enhancements to document"""
        self.logger.info("🔧 Applying comprehensive document quality enhancements...")
        
        # Phase 1: Fix TOC and bookmark issues
        self._fix_toc_bookmark_consistency(doc)
        
        # Phase 2: Clean up placeholder codes
        self._restore_placeholder_codes(doc)
        
        # Phase 3: Fix paragraph fragmentation
        self._consolidate_fragmented_paragraphs(doc)
        
        # Phase 4: Validate and fix image references
        if digital_twin_doc:
            self._validate_image_references(doc, digital_twin_doc)
        
        # Phase 5: Improve mathematical formula rendering
        self._enhance_mathematical_formulas(doc)
        
        # Phase 6: Fix header/footer placement
        self._separate_metadata_content(doc)
        
        # Phase 7: Remove empty pages
        self._remove_empty_pages(doc)
        
        self.logger.info(f"✅ Document quality enhancement completed. Issues found and fixed: {len(self.formatting_issues)}")
        return self.formatting_issues
    
    def _fix_toc_bookmark_consistency(self, doc):
        """Fix TOC bookmark naming consistency to eliminate 'Error! Bookmark not defined.'"""
        self.logger.info("🔗 Fixing TOC bookmark consistency...")
        
        # Scan all bookmarks in the document
        bookmarks_found = []
        for element in doc.element.iter():
            if element.tag.endswith('bookmarkStart'):
                bookmark_name = element.get(qn('w:name'))
                if bookmark_name:
                    bookmarks_found.append(bookmark_name)
        
        # Scan all hyperlinks in TOC
        hyperlinks_found = []
        for element in doc.element.iter():
            if element.tag.endswith('hyperlink'):
                anchor = element.get(qn('w:anchor'))
                if anchor:
                    hyperlinks_found.append(anchor)
        
        # Report mismatches
        missing_bookmarks = set(hyperlinks_found) - set(bookmarks_found)
        if missing_bookmarks:
            self.logger.warning(f"Found {len(missing_bookmarks)} missing bookmarks: {missing_bookmarks}")
            self.formatting_issues.extend([f"Missing bookmark: {bm}" for bm in missing_bookmarks])
        
        # Fix hyperlinks to point to existing bookmarks
        self._fix_hyperlink_references(doc, bookmarks_found, hyperlinks_found)
    
    def _fix_hyperlink_references(self, doc, available_bookmarks, requested_bookmarks):
        """Fix hyperlink references to point to actually existing bookmarks"""
        # Create a mapping from requested to available bookmarks
        bookmark_mapping = {}
        
        for requested in requested_bookmarks:
            if requested not in available_bookmarks:
                # Try to find a similar bookmark
                best_match = self._find_best_bookmark_match(requested, available_bookmarks)
                if best_match:
                    bookmark_mapping[requested] = best_match
                    self.logger.info(f"Mapping broken link '{requested}' to '{best_match}'")
        
        # Apply the mapping to hyperlinks
        for element in doc.element.iter():
            if element.tag.endswith('hyperlink'):
                anchor = element.get(qn('w:anchor'))
                if anchor in bookmark_mapping:
                    element.set(qn('w:anchor'), bookmark_mapping[anchor])
                    self.formatting_issues.append(f"Fixed hyperlink: {anchor} → {bookmark_mapping[anchor]}")
    
    def _find_best_bookmark_match(self, requested_bookmark, available_bookmarks):
        """Find the best matching bookmark for a broken reference"""
        import difflib
        
        # Try exact numeric matching first (common pattern: _Toc_123 vs _Toc_Bookmark_123)
        if '_Toc_' in requested_bookmark:
            number_part = requested_bookmark.split('_Toc_')[-1]
            for available in available_bookmarks:
                if number_part in available:
                    return available
        
        # Use fuzzy matching for similar names
        matches = difflib.get_close_matches(requested_bookmark, available_bookmarks, n=1, cutoff=0.6)
        return matches[0] if matches else None
    
    def _restore_placeholder_codes(self, doc):
        """Restore placeholder codes like PRESERVE0007 to their original characters"""
        self.logger.info("🔤 Restoring placeholder codes...")
        
        # Common placeholder patterns and their likely meanings
        placeholder_mappings = {
            r'PRESERVE(\d{4})': self._restore_preserve_code,
            r'\[PARAGRAPH_BREAK\]': '\n\n',
            r'\[IMAGE_INSERTION_FAILED[^\]]*\]': '[Image not available]',
            r'\[MISSING_PAGE[^\]]*\]': '',
            r'\[ERROR[^\]]*\]': '',
        }
        
        text_elements_fixed = 0
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            fixed_text = original_text
            
            for pattern, replacement in placeholder_mappings.items():
                if callable(replacement):
                    import re
                    fixed_text = re.sub(pattern, replacement, fixed_text)
                else:
                    import re
                    fixed_text = re.sub(pattern, replacement, fixed_text)
            
            if fixed_text != original_text:
                # Update paragraph text
                self._update_paragraph_text(paragraph, fixed_text)
                text_elements_fixed += 1
                self.formatting_issues.append(f"Restored placeholder codes in paragraph")
        
        self.logger.info(f"Fixed placeholder codes in {text_elements_fixed} text elements")
    
    def _restore_preserve_code(self, match):
        """Restore PRESERVE codes to original characters"""
        code_num = int(match.group(1))
        
        # Common mappings for mathematical and special characters
        preserve_mappings = {
            7: 'Φ',  # Greek letter Phi
            8: '≤',  # Less than or equal
            9: '≥',  # Greater than or equal
            10: '∑', # Summation
            11: '∫', # Integral
            12: '∂', # Partial derivative
            13: '±', # Plus-minus
            14: '×', # Multiplication
            15: '÷', # Division
            16: '≠', # Not equal
            17: '≈', # Approximately equal
            18: '∞', # Infinity
            19: '→', # Right arrow
            20: '←', # Left arrow
        }
        
        return preserve_mappings.get(code_num, f'[UNKNOWN_PRESERVE_{code_num}]')
    
    def _update_paragraph_text(self, paragraph, new_text):
        """Safely update paragraph text while preserving formatting"""
        if paragraph.runs:
            # Update first run and clear others
            paragraph.runs[0].text = new_text
            for i in range(len(paragraph.runs) - 1, 0, -1):
                paragraph.runs[i].clear()
        else:
            # Add new run if no runs exist
            paragraph.add_run(new_text)
    
    def _consolidate_fragmented_paragraphs(self, doc):
        """Fix paragraph fragmentation by consolidating related content"""
        self.logger.info("📄 Consolidating fragmented paragraphs...")
        
        paragraphs_to_remove = []
        consolidation_count = 0
        
        for i in range(len(doc.paragraphs) - 1):
            current_p = doc.paragraphs[i]
            next_p = doc.paragraphs[i + 1]
            
            # Check if paragraphs should be consolidated
            if self._should_consolidate_paragraphs(current_p, next_p):
                # Consolidate next paragraph into current
                consolidated_text = current_p.text + " " + next_p.text
                self._update_paragraph_text(current_p, consolidated_text)
                
                # Mark next paragraph for removal
                paragraphs_to_remove.append(next_p)
                consolidation_count += 1
        
        # Remove marked paragraphs
        for p in paragraphs_to_remove:
            self._remove_paragraph(p)
        
        self.logger.info(f"Consolidated {consolidation_count} fragmented paragraphs")
        self.formatting_issues.extend([f"Consolidated fragmented paragraph" for _ in range(consolidation_count)])
    
    def _should_consolidate_paragraphs(self, current_p, next_p):
        """Determine if two paragraphs should be consolidated"""
        current_text = current_p.text.strip()
        next_text = next_p.text.strip()
        
        # Don't consolidate if either is empty
        if not current_text or not next_text:
            return False
        
        # Don't consolidate headings
        if any(style in current_p.style.name.lower() for style in ['heading', 'title']):
            return False
        
        # Consolidate if current paragraph doesn't end with sentence punctuation
        if not current_text.endswith(('.', '!', '?', ':', ';')):
            return True
        
        # Consolidate if next paragraph starts with lowercase (likely continuation)
        if next_text[0].islower():
            return True
        
        return False
    
    def _remove_paragraph(self, paragraph):
        """Safely remove a paragraph from the document"""
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    def _validate_image_references(self, doc, digital_twin_doc):
        """Validate and fix image references"""
        self.logger.info("🖼️ Validating image references...")
        
        image_issues_fixed = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text
            
            # Look for image insertion failure messages
            if '[Image insertion failed:' in text or '[Image not found:' in text:
                # Try to find the actual image file
                import re
                filename_match = re.search(r'\[Image[^\]]*:\s*([^\]]+)\]', text)
                if filename_match:
                    filename = filename_match.group(1).strip()
                    actual_path = self._find_image_file(filename, digital_twin_doc)
                    
                    if actual_path and os.path.exists(actual_path):
                        # Replace error message with actual image
                        self._insert_image_in_paragraph(paragraph, actual_path)
                        image_issues_fixed += 1
                        self.formatting_issues.append(f"Fixed image: {filename}")
        
        self.logger.info(f"Fixed {image_issues_fixed} image references")
    
    def _find_image_file(self, filename, digital_twin_doc):
        """Find the actual path of an image file"""
        if not digital_twin_doc:
            return None
        
        # Search through all image blocks in digital twin
        for page in digital_twin_doc.pages:
            for block in page.get_all_blocks():
                if hasattr(block, 'image_path') and block.image_path:
                    if filename in block.image_path or os.path.basename(block.image_path) == filename:
                        return block.image_path
        
        return None
    
    def _insert_image_in_paragraph(self, paragraph, image_path):
        """Insert an actual image into a paragraph, replacing error text"""
        try:
            # Clear paragraph content
            for run in paragraph.runs:
                run.clear()
            
            # Add image
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(4.0))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        except Exception as e:
            self.logger.warning(f"Failed to insert image {image_path}: {e}")
    
    def _enhance_mathematical_formulas(self, doc):
        """Improve mathematical formula rendering"""
        self.logger.info("🔢 Enhancing mathematical formulas...")
        
        formula_patterns = [
            (r'Y\s+a1X1\s+anXn\s*\(([^)]+)\)', r'Y = a₁X₁ + ... + aₙXₙ (\1)'),
            (r'F\s*\(', r'Φ('),  # Replace F with Phi symbol
            (r'\bphi\b', r'φ'),  # Replace phi with symbol
            (r'\balpha\b', r'α'),
            (r'\bbeta\b', r'β'),
            (r'\bgamma\b', r'γ'),
            (r'\bdelta\b', r'δ'),
            (r'<=', r'≤'),
            (r'>=', r'≥'),
            (r'!=', r'≠'),
            (r'+-', r'±'),
        ]
        
        formulas_enhanced = 0
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            enhanced_text = original_text
            
            for pattern, replacement in formula_patterns:
                import re
                enhanced_text = re.sub(pattern, replacement, enhanced_text)
            
            if enhanced_text != original_text:
                self._update_paragraph_text(paragraph, enhanced_text)
                formulas_enhanced += 1
                self.formatting_issues.append(f"Enhanced mathematical formula")
        
        self.logger.info(f"Enhanced {formulas_enhanced} mathematical formulas")
    
    def _separate_metadata_content(self, doc):
        """Separate headers, footers, and citations from main content"""
        self.logger.info("📋 Separating metadata content...")
        
        metadata_separated = 0
        paragraphs_to_remove = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if self._is_metadata_content(text):
                # Convert to footer-style formatting
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)
                
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                metadata_separated += 1
                self.formatting_issues.append(f"Separated metadata: {text[:50]}...")
        
        self.logger.info(f"Separated {metadata_separated} metadata elements")
    
    def _is_metadata_content(self, text):
        """Determine if text is metadata (header/footer/citation)"""
        metadata_indicators = [
            'doi:', 'http://', 'https://', 'www.',
            'copyright', '©', 'page ', 'volume ', 'journal',
            'received:', 'accepted:', 'published:',
            r'^\d+$',  # Just page numbers
            r'^\d{4}$',  # Just years
        ]
        
        text_lower = text.lower()
        for indicator in metadata_indicators:
            import re
            if isinstance(indicator, str):
                if indicator in text_lower:
                    return True
            else:
                if re.match(indicator, text):
                    return True
        
        return False
    
    def _remove_empty_pages(self, doc):
        """Remove unnecessary empty pages"""
        self.logger.info("📄 Removing empty pages...")
        
        empty_paragraphs_removed = 0
        paragraphs_to_remove = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Remove completely empty paragraphs
            if not text:
                paragraphs_to_remove.append(paragraph)
                empty_paragraphs_removed += 1
            
            # Remove paragraphs with only whitespace or formatting artifacts
            elif text in ['', ' ', '\n', '\t', '""', "''"]:
                paragraphs_to_remove.append(paragraph)
                empty_paragraphs_removed += 1
        
        # Remove marked paragraphs
        for p in paragraphs_to_remove:
            self._remove_paragraph(p)
        
        self.logger.info(f"Removed {empty_paragraphs_removed} empty content elements")
        self.formatting_issues.extend([f"Removed empty content" for _ in range(empty_paragraphs_removed)])


# Enhanced PDF conversion function with proper font embedding
def convert_word_to_pdf(docx_filepath, pdf_filepath):
    """Enhanced PDF conversion with proper Greek font support"""
    logger.info(f"Converting {docx_filepath} to PDF with Unicode font support...")
    
    try:
        # Try using docx2pdf first
        from docx2pdf import convert
        convert(docx_filepath, pdf_filepath)
        logger.info(f"Successfully converted {docx_filepath} to {pdf_filepath}")
        return True
        
    except ImportError as e:
        logger.error("❌ MISSING DEPENDENCY: docx2pdf library not available")
        logger.error("📋 To fix this issue, install missing dependencies by running:")
        logger.error("   pip install -r requirements.txt")
        logger.error("   OR manually: pip install docx2pdf docx2txt")
        logger.warning("🔄 Trying alternative method with reportlab...")
        
        try:
            # Alternative: Use python-docx2txt + reportlab for better font control
            import docx2txt
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.fonts import addMapping
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Register Unicode font
            try:
                # Try to register Arial Unicode MS or fallback to DejaVu
                pdfmetrics.registerFont(TTFont('ArialUnicode', 'arial.ttf'))
                addMapping('ArialUnicode', 0, 0, 'ArialUnicode')
            except:
                # Fallback to DejaVu Sans which supports Greek
                try:
                    pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
                    addMapping('DejaVuSans', 0, 0, 'DejaVuSans')
                    font_name = 'DejaVuSans'
                except:
                    font_name = 'Helvetica'  # Last resort
                    logger.warning("Could not register Unicode font, using Helvetica")
            
            # Extract text from DOCX
            text = docx2txt.process(docx_filepath)
            
            # Create PDF with proper font
            doc = SimpleDocTemplate(pdf_filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Create custom style with Unicode font
            unicode_style = ParagraphStyle(
                'UnicodeNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=11,
                encoding='utf-8'
            )
            
            # Build PDF content
            story = []
            for line in text.split('\n'):
                if line.strip():
                    story.append(Paragraph(line, unicode_style))
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            logger.info(f"Successfully converted {docx_filepath} to {pdf_filepath} with Unicode support")
            return True
            
        except ImportError as alt_import_error:
            logger.error("❌ MISSING DEPENDENCIES: Alternative PDF conversion libraries not available")
            logger.error("📋 To fix this issue, install all required dependencies:")
            logger.error("   pip install -r requirements.txt")
            logger.error("📋 Required packages: docx2pdf, docx2txt, reportlab")
            logger.error(f"💡 Import error details: {alt_import_error}")
            return False
        except Exception as e:
            logger.error(f"❌ Alternative PDF conversion failed: {e}")
            return False
            
    except Exception as e:
        logger.error(f"❌ Error converting Word to PDF: {e}")
        logger.error("📋 If this is a dependency issue, try: pip install -r requirements.txt")
        return False

# Create a class to hold the PDF conversion function
class PDFConverter:
    def __init__(self):
        self.convert_word_to_pdf = convert_word_to_pdf

# Create an instance of the PDF converter
pdf_converter = PDFConverter()

# Instantiate the generator for use by other modules
document_generator = WordDocumentGenerator()