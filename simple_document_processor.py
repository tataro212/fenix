"""
Simple Document Processor for Fenix PDF Translation Pipeline

This module provides fast-track processing for simple documents that don't require
the full Digital Twin architecture. It performs direct text extraction and translation
with minimal overhead.
"""

import os
import logging
import fitz  # PyMuPDF
import asyncio
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

@dataclass
class SimpleProcessingResult:
    """Result of simple document processing"""
    success: bool
    output_path: str
    page_count: int
    text_blocks_processed: int
    translation_time: float
    total_time: float
    error_message: Optional[str] = None

class SimpleDocumentProcessor:
    """
    Fast-track processor for simple documents.
    
    This processor handles documents that meet the simple criteria:
    - No images/figures
    - No tables
    - No complex formatting
    - No footnotes/endnotes
    - No mathematical equations
    - No multi-column layouts
    - Single font family
    """
    
    def __init__(self, gemini_service=None):
        self.logger = logging.getLogger(__name__)
        self.gemini_service = gemini_service
        
        # Processing settings
        self.CHUNK_SIZE = 2000  # Characters per translation chunk
        self.MIN_CHUNK_SIZE = 500  # Minimum chunk size
        self.MAX_CONCURRENT_TRANSLATIONS = 5
        
    async def process_simple_document(
        self, 
        pdf_path: str, 
        output_path: str, 
        target_language: str = "el"
    ) -> SimpleProcessingResult:
        """
        Process a simple document with fast-track method.
        
        Args:
            pdf_path: Path to input PDF
            output_path: Path for output Word document
            target_language: Target language code
            
        Returns:
            SimpleProcessingResult with processing details
        """
        start_time = asyncio.get_event_loop().time()
        
        try:
            self.logger.info(f"🚀 Starting fast-track simple document processing")
            self.logger.info(f"   Input: {os.path.basename(pdf_path)}")
            self.logger.info(f"   Output: {os.path.basename(output_path)}")
            self.logger.info(f"   Target Language: {target_language}")
            
            # Step 1: Extract text content
            text_blocks = self._extract_simple_text(pdf_path)
            self.logger.info(f"📄 Extracted {len(text_blocks)} text blocks")
            
            if not text_blocks:
                return SimpleProcessingResult(
                    success=False,
                    output_path=output_path,
                    page_count=0,
                    text_blocks_processed=0,
                    translation_time=0.0,
                    total_time=0.0,
                    error_message="No text content found in document"
                )
            
            # Step 2: Translate content
            translation_start = asyncio.get_event_loop().time()
            translated_blocks = await self._translate_text_blocks(text_blocks, target_language)
            translation_time = asyncio.get_event_loop().time() - translation_start
            
            self.logger.info(f"🔄 Translated {len(translated_blocks)} blocks in {translation_time:.2f}s")
            
            # Step 3: Generate Word document
            doc_generated = self._generate_simple_word_document(
                translated_blocks, 
                output_path, 
                target_language
            )
            
            if not doc_generated:
                return SimpleProcessingResult(
                    success=False,
                    output_path=output_path,
                    page_count=len(text_blocks),
                    text_blocks_processed=len(text_blocks),
                    translation_time=translation_time,
                    total_time=asyncio.get_event_loop().time() - start_time,
                    error_message="Failed to generate Word document"
                )
            
            total_time = asyncio.get_event_loop().time() - start_time
            
            self.logger.info(f"✅ Fast-track processing completed successfully")
            self.logger.info(f"   Total time: {total_time:.2f}s")
            self.logger.info(f"   Translation time: {translation_time:.2f}s")
            self.logger.info(f"   Document generation time: {total_time - translation_time:.2f}s")
            
            return SimpleProcessingResult(
                success=True,
                output_path=output_path,
                page_count=len(text_blocks),
                text_blocks_processed=len(translated_blocks),
                translation_time=translation_time,
                total_time=total_time
            )
            
        except Exception as e:
            self.logger.error(f"❌ Fast-track processing failed: {e}")
            return SimpleProcessingResult(
                success=False,
                output_path=output_path,
                page_count=0,
                text_blocks_processed=0,
                translation_time=0.0,
                total_time=asyncio.get_event_loop().time() - start_time,
                error_message=str(e)
            )
    
    def _extract_simple_text(self, pdf_path: str) -> List[Dict[str, Any]]:
        """
        Extract text from PDF using simple method.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            List of text blocks with metadata
        """
        try:
            doc = fitz.open(pdf_path)
            text_blocks = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract text with basic structure preservation
                text = page.get_text()
                
                if not text.strip():
                    continue
                
                # Split into paragraphs (simple approach)
                paragraphs = self._split_into_paragraphs(text)
                
                for para_idx, paragraph in enumerate(paragraphs):
                    if paragraph.strip():
                        text_blocks.append({
                            'page_num': page_num + 1,
                            'block_id': f"page_{page_num + 1}_para_{para_idx}",
                            'text': paragraph.strip(),
                            'type': 'paragraph',
                            'length': len(paragraph.strip())
                        })
            
            doc.close()
            return text_blocks
            
        except Exception as e:
            self.logger.error(f"Error extracting simple text: {e}")
            return []
    
    def _split_into_paragraphs(self, text: str) -> List[str]:
        """
        Split text into paragraphs using simple heuristics.
        
        Args:
            text: Raw text from PDF
            
        Returns:
            List of paragraph strings
        """
        # Split on double newlines first
        paragraphs = text.split('\n\n')
        
        # Further split long paragraphs
        refined_paragraphs = []
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            # If paragraph is too long, split on single newlines
            if len(para) > 1000:
                sub_paragraphs = para.split('\n')
                for sub_para in sub_paragraphs:
                    sub_para = sub_para.strip()
                    if sub_para and len(sub_para) > 10:  # Minimum paragraph length
                        refined_paragraphs.append(sub_para)
            else:
                # Replace single newlines with spaces within paragraph
                para = para.replace('\n', ' ')
                refined_paragraphs.append(para)
        
        return refined_paragraphs
    
    async def _translate_text_blocks(
        self, 
        text_blocks: List[Dict[str, Any]], 
        target_language: str
    ) -> List[Dict[str, Any]]:
        """
        Translate text blocks using concurrent processing.
        
        Args:
            text_blocks: List of text blocks to translate
            target_language: Target language code
            
        Returns:
            List of translated text blocks
        """
        if not self.gemini_service:
            self.logger.warning("No Gemini service available, returning original text")
            return text_blocks
        
        # Create translation chunks
        translation_chunks = self._create_translation_chunks(text_blocks)
        
        # Process chunks concurrently
        semaphore = asyncio.Semaphore(self.MAX_CONCURRENT_TRANSLATIONS)
        
        async def translate_chunk(chunk):
            async with semaphore:
                try:
                    combined_text = "\n\n".join(block['text'] for block in chunk['blocks'])
                    
                    # Create translation prompt
                    prompt = (
                        f"Translate the following text to {target_language}. "
                        f"Maintain the original paragraph structure and formatting. "
                        f"Do not add any explanations or comments.\n\n"
                        f"TEXT TO TRANSLATE:\n{combined_text}"
                    )
                    
                    # Translate
                    translated_text = await self.gemini_service.translate_text(
                        prompt, 
                        target_language
                    )
                    
                    # Split translated text back into blocks
                    translated_paragraphs = translated_text.split('\n\n')
                    
                    # Update blocks with translations
                    for i, block in enumerate(chunk['blocks']):
                        if i < len(translated_paragraphs):
                            block['translated_text'] = translated_paragraphs[i].strip()
                        else:
                            block['translated_text'] = block['text']  # Fallback
                    
                    return chunk['blocks']
                    
                except Exception as e:
                    self.logger.error(f"Error translating chunk: {e}")
                    # Return original text as fallback
                    for block in chunk['blocks']:
                        block['translated_text'] = block['text']
                    return chunk['blocks']
        
        # Execute translations concurrently
        translation_tasks = [translate_chunk(chunk) for chunk in translation_chunks]
        translated_chunks = await asyncio.gather(*translation_tasks)
        
        # Flatten results
        translated_blocks = []
        for chunk_blocks in translated_chunks:
            translated_blocks.extend(chunk_blocks)
        
        return translated_blocks
    
    def _create_translation_chunks(self, text_blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Create translation chunks from text blocks.
        
        Args:
            text_blocks: List of text blocks
            
        Returns:
            List of translation chunks
        """
        chunks = []
        current_chunk = []
        current_length = 0
        
        for block in text_blocks:
            block_length = len(block['text'])
            
            # If adding this block would exceed chunk size, start new chunk
            if current_length + block_length > self.CHUNK_SIZE and current_chunk:
                chunks.append({
                    'blocks': current_chunk,
                    'total_length': current_length
                })
                current_chunk = []
                current_length = 0
            
            current_chunk.append(block)
            current_length += block_length
        
        # Add final chunk
        if current_chunk:
            chunks.append({
                'blocks': current_chunk,
                'total_length': current_length
            })
        
        return chunks
    
    def _generate_simple_word_document(
        self, 
        translated_blocks: List[Dict[str, Any]], 
        output_path: str,
        target_language: str
    ) -> bool:
        """
        Generate a simple Word document from translated blocks.
        
        Args:
            translated_blocks: List of translated text blocks
            output_path: Path for output document
            target_language: Target language code
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Create Word document
            doc = Document()
            
            # Configure document fonts - Times New Roman 11pt for normal text
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(11)
            
            # Add document title if we can infer one
            title = self._infer_document_title(translated_blocks)
            if title:
                title_para = doc.add_heading(title, level=1)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Add spacing
            
            # Add translated content
            for block in translated_blocks:
                translated_text = block.get('translated_text', block['text'])
                
                if translated_text.strip():
                    # Determine if this looks like a heading
                    if self._is_likely_heading(translated_text):
                        doc.add_heading(translated_text, level=2)
                    else:
                        # Add as regular paragraph
                        para = doc.add_paragraph(translated_text)
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Save document
            doc.save(output_path)
            self.logger.info(f"💾 Simple Word document saved: {output_path}")
            
            return True
            
        except Exception as e:
            self.logger.error(f"Error generating simple Word document: {e}")
            return False
    
    def _infer_document_title(self, text_blocks: List[Dict[str, Any]]) -> Optional[str]:
        """
        Try to infer document title from first few blocks.
        
        Args:
            text_blocks: List of text blocks
            
        Returns:
            Inferred title or None
        """
        if not text_blocks:
            return None
        
        # Check first few blocks for title-like content
        for block in text_blocks[:3]:
            text = block.get('translated_text', block['text'])
            
            # Title heuristics
            if (len(text) < 100 and  # Not too long
                len(text.split()) < 15 and  # Not too many words
                text.isupper() or  # All uppercase
                (text[0].isupper() and not text.endswith('.'))):  # Starts with capital, no period
                return text
        
        return None
    
    def _is_likely_heading(self, text: str) -> bool:
        """
        Determine if text is likely a heading.
        
        Args:
            text: Text to analyze
            
        Returns:
            True if likely a heading
        """
        # Heading heuristics
        return (
            len(text) < 80 and  # Short enough
            len(text.split()) < 12 and  # Not too many words
            not text.endswith('.') and  # Doesn't end with period
            (text[0].isupper() or  # Starts with capital
             text.isupper() or  # All uppercase
             any(char.isdigit() for char in text[:5]))  # Starts with number
        )


# Global instance for easy access
simple_document_processor = SimpleDocumentProcessor() 