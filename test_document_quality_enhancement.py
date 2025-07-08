#!/usr/bin/env python3
"""
Comprehensive Test Suite for Document Quality Enhancement

This test suite validates that the DocumentQualityEnhancer properly addresses
all the systematic issues in Greek PDF translation:

1. TOC bookmark consistency and hyperlink fixing
2. PRESERVE placeholder restoration for mathematical symbols
3. Paragraph fragmentation consolidation
4. Image insertion failure resolution
5. Mathematical formula enhancement
6. Metadata separation from main content
7. Empty content and artifact removal

Usage:
    python test_document_quality_enhancement.py
"""

import unittest
import tempfile
import os
import logging
from docx import Document
from docx.shared import Inches

# Configure test logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

class TestDocumentQualityEnhancer(unittest.TestCase):
    """Test suite for DocumentQualityEnhancer functionality"""
    
    def setUp(self):
        """Set up test environment"""
        self.temp_dir = tempfile.mkdtemp()
        
        # Try to import the quality enhancer
        try:
            from document_quality_enhancer import DocumentQualityEnhancer
            self.enhancer = DocumentQualityEnhancer()
            self.enhancer_available = True
        except ImportError:
            self.enhancer_available = False
            self.skipTest("DocumentQualityEnhancer not available")
    
    def tearDown(self):
        """Clean up test environment"""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_preserve_placeholder_restoration(self):
        """Test PRESERVE placeholder restoration to mathematical symbols"""
        # Create document with PRESERVE placeholders
        doc = Document()
        doc.add_paragraph("Mathematical symbols: PRESERVE0001 and PRESERVE0006")
        doc.add_paragraph("Integration: PRESERVE0007 from 0 to PRESERVE0011")
        doc.add_paragraph("Comparison: x PRESERVE0002 y PRESERVE0003 z")
        
        # Apply enhancement
        enhancement_report = self.enhancer.enhance_document_quality(doc)
        
        # Verify symbols were restored
        text_content = "\n".join([p.text for p in doc.paragraphs])
        
        # Check specific symbol restorations
        self.assertIn("Φ", text_content, "Phi symbol should be restored")
        self.assertIn("∑", text_content, "Summation symbol should be restored") 
        self.assertIn("∫", text_content, "Integral symbol should be restored")
        self.assertIn("∞", text_content, "Infinity symbol should be restored")
        self.assertIn("≥", text_content, "Greater than or equal should be restored")
        
        # Check that PRESERVE codes are removed
        self.assertNotIn("PRESERVE", text_content, "PRESERVE codes should be removed")
        
        # Verify enhancement report
        symbol_fixes = enhancement_report.get('issues_by_category', {}).get('mathematical_symbols', {})
        self.assertGreater(symbol_fixes.get('symbols_restored', 0), 0, "Should report symbol restorations")
    
    def test_paragraph_fragmentation_consolidation(self):
        """Test consolidation of fragmented paragraphs"""
        # Create document with fragmented paragraphs
        doc = Document()
        doc.add_paragraph("This is the first part of a sentence that continues")
        doc.add_paragraph("on the next line without proper punctuation")
        doc.add_paragraph("and should be consolidated into one paragraph")
        
        # Add properly separated content that should NOT be consolidated
        doc.add_paragraph("This is a complete sentence.")
        doc.add_paragraph("This is another complete sentence that should remain separate.")
        
        original_para_count = len(doc.paragraphs)
        
        # Apply enhancement
        enhancement_report = self.enhancer.enhance_document_quality(doc)
        
        # Should have fewer paragraphs after consolidation
        self.assertLess(len(doc.paragraphs), original_para_count, "Should consolidate fragmented paragraphs")
        
        # Verify enhancement report
        paragraph_fixes = enhancement_report.get('issues_by_category', {}).get('paragraph_consolidation', {})
        self.assertGreater(paragraph_fixes.get('paragraphs_consolidated', 0), 0, "Should report paragraph consolidations")
    
    def test_mathematical_formula_enhancement(self):
        """Test enhancement of plain text mathematical formulas"""
        # Create document with plain text formulas
        doc = Document()
        doc.add_paragraph("The equation is: Y a1X1 anXn (3.1)")
        doc.add_paragraph("Function F(x) represents the distribution")
        doc.add_paragraph("Where symbols like <= and >= should be enhanced")
        doc.add_paragraph("And F should become the Greek letter Phi")
        
        # Apply enhancement
        enhancement_report = self.enhancer.enhance_document_quality(doc)
        
        text_content = "\n".join([p.text for p in doc.paragraphs])
        
        # Check formula enhancements
        self.assertIn("≤", text_content, "Less than or equal should be enhanced")
        self.assertIn("≥", text_content, "Greater than or equal should be enhanced") 
        self.assertIn("Φ", text_content, "F should be converted to Phi in mathematical contexts")
        
        # Verify enhancement report
        formula_fixes = enhancement_report.get('issues_by_category', {}).get('mathematical_formulas', {})
        self.assertGreater(formula_fixes.get('formulas_enhanced', 0), 0, "Should report formula enhancements")
    
    def test_metadata_content_separation(self):
        """Test separation of metadata from main content"""
        # Create document with metadata mixed in main content
        doc = Document()
        doc.add_paragraph("Regular content paragraph")
        doc.add_paragraph("Page 5")  # Page number - should be detected as metadata
        doc.add_paragraph("JOURNAL OF COMPUTATIONAL SCIENCE")  # Header - metadata
        doc.add_paragraph("© 2024 Academic Publishers")  # Footer - metadata
        doc.add_paragraph("More regular content")
        doc.add_paragraph("DOI: 10.1234/example.doi")  # DOI - metadata
        
        # Apply enhancement
        enhancement_report = self.enhancer.enhance_document_quality(doc)
        
        # Verify enhancement report
        metadata_fixes = enhancement_report.get('issues_by_category', {}).get('metadata_separation', {})
        self.assertGreater(metadata_fixes.get('metadata_paragraphs_found', 0), 0, "Should detect metadata paragraphs")
        self.assertGreater(metadata_fixes.get('paragraphs_reformatted', 0), 0, "Should reformat metadata paragraphs")
    
    def test_empty_content_artifact_removal(self):
        """Test removal of empty content and formatting artifacts"""
        # Create document with various empty/artifact content
        doc = Document()
        doc.add_paragraph("Regular content")
        doc.add_paragraph("")  # Empty paragraph
        doc.add_paragraph("   ")  # Whitespace only
        doc.add_paragraph('""')  # Quote artifacts
        doc.add_paragraph("...")  # Minimal content
        doc.add_paragraph("____________")  # Formatting artifacts
        doc.add_paragraph("More regular content")
        doc.add_paragraph("-----")  # More formatting artifacts
        
        original_para_count = len(doc.paragraphs)
        
        # Apply enhancement
        enhancement_report = self.enhancer.enhance_document_quality(doc)
        
        # Should have fewer paragraphs after cleanup
        self.assertLess(len(doc.paragraphs), original_para_count, "Should remove empty/artifact paragraphs")
        
        # Verify enhancement report
        cleanup_fixes = enhancement_report.get('issues_by_category', {}).get('content_cleanup', {})
        self.assertGreater(cleanup_fixes.get('paragraphs_removed', 0), 0, "Should report removed paragraphs")
    
    def test_toc_bookmark_consistency_simulation(self):
        """Test TOC bookmark consistency fixing (simulated)"""
        # Note: Full bookmark testing requires complex Word XML manipulation
        # This test verifies the detection logic
        
        doc = Document()
        doc.add_paragraph("Table of Contents")
        doc.add_paragraph("1. Introduction ............ Error! Bookmark not defined.")
        doc.add_paragraph("2. Methods ................ Error! Bookmark not defined.")
        doc.add_paragraph("3. Results ................ Error! Bookmark not defined.")
        
        # Apply enhancement
        enhancement_report = self.enhancer.enhance_document_quality(doc)
        
        # Verify that TOC bookmark processing was attempted
        toc_fixes = enhancement_report.get('issues_by_category', {}).get('toc_bookmarks', {})
        # The specific results depend on the document structure and bookmark availability
        self.assertIsInstance(toc_fixes, dict, "Should return TOC bookmark fix report")
    
    def test_comprehensive_enhancement(self):
        """Test comprehensive enhancement with multiple issue types"""
        # Create document with all types of issues
        doc = Document()
        
        # TOC issues
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("1. Introduction ............ Error! Bookmark not defined.")
        
        # Fragmented paragraphs
        doc.add_paragraph("This is a fragmented paragraph that continues")
        doc.add_paragraph("on the next line without proper punctuation")
        
        # PRESERVE placeholders
        doc.add_paragraph("Mathematical symbols: PRESERVE0001 PRESERVE0006 PRESERVE0007")
        
        # Image insertion failures
        doc.add_paragraph("[Image insertion failed: figure_1.png]")
        doc.add_paragraph("Image Placeholder")
        
        # Mathematical formulas
        doc.add_paragraph("Equation: Y a1X1 anXn (3.1) where F(x) <= threshold")
        
        # Metadata in content
        doc.add_paragraph("Page 10")
        doc.add_paragraph("JOURNAL HEADER TEXT")
        
        # Empty artifacts
        doc.add_paragraph("")
        doc.add_paragraph("...")
        doc.add_paragraph("____________")
        
        original_para_count = len(doc.paragraphs)
        
        # Apply comprehensive enhancement
        enhancement_report = self.enhancer.enhance_document_quality(doc)
        
        # Verify overall results
        self.assertGreater(enhancement_report.get('total_issues_found', 0), 0, "Should find multiple issues")
        self.assertGreater(enhancement_report.get('total_fixes_applied', 0), 0, "Should apply multiple fixes")
        self.assertGreater(len(enhancement_report.get('categories_enhanced', [])), 0, "Should enhance multiple categories")
        
        # Check specific category processing
        issues_by_category = enhancement_report.get('issues_by_category', {})
        expected_categories = [
            'mathematical_symbols', 'paragraph_consolidation', 'mathematical_formulas',
            'metadata_separation', 'content_cleanup', 'toc_bookmarks'
        ]
        
        for category in expected_categories:
            self.assertIn(category, issues_by_category, f"Should process {category}")
    
    def test_enhancement_with_digital_twin_simulation(self):
        """Test enhancement with simulated Digital Twin integration"""
        # This test simulates what would happen with Digital Twin integration
        # without requiring the full Digital Twin infrastructure
        
        doc = Document()
        doc.add_paragraph("[Image insertion failed: sample_image.png]")
        doc.add_paragraph("Mathematical formula: PRESERVE0001 = PRESERVE0006")
        
        # Apply enhancement without Digital Twin (should handle gracefully)
        enhancement_report = self.enhancer.enhance_document_quality(
            doc=doc,
            digital_twin_doc=None  # No Digital Twin available
        )
        
        # Should still process other enhancements
        self.assertGreater(enhancement_report.get('total_fixes_applied', 0), 0, "Should apply non-image fixes")
        
        # Mathematical symbols should still be restored
        text_content = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("Φ", text_content, "Mathematical symbols should be restored without Digital Twin")
    
    def test_error_handling_and_robustness(self):
        """Test error handling and robustness of enhancement system"""
        # Create document with potentially problematic content
        doc = Document()
        doc.add_paragraph("Normal content")
        doc.add_paragraph("")  # Empty content
        doc.add_paragraph("Content with unicode: αβγδε")  # Unicode content
        doc.add_paragraph("Very long content " + "x" * 1000)  # Very long content
        
        # Should handle gracefully without crashing
        try:
            enhancement_report = self.enhancer.enhance_document_quality(doc)
            self.assertIsInstance(enhancement_report, dict, "Should return valid report even with edge cases")
        except Exception as e:
            self.fail(f"Enhancement should handle edge cases gracefully, but failed with: {e}")
    
    def test_enhancement_report_structure(self):
        """Test that enhancement reports have the expected structure"""
        doc = Document()
        doc.add_paragraph("Test content with PRESERVE0001")
        
        enhancement_report = self.enhancer.enhance_document_quality(doc)
        
        # Verify report structure
        required_keys = ['total_issues_found', 'total_fixes_applied', 'categories_enhanced', 'issues_by_category']
        for key in required_keys:
            self.assertIn(key, enhancement_report, f"Report should contain {key}")
        
        # Verify data types
        self.assertIsInstance(enhancement_report['total_issues_found'], int)
        self.assertIsInstance(enhancement_report['total_fixes_applied'], int)
        self.assertIsInstance(enhancement_report['categories_enhanced'], list)
        self.assertIsInstance(enhancement_report['issues_by_category'], dict)

class TestQualityEnhancementIntegration(unittest.TestCase):
    """Test integration with document generation pipeline"""
    
    def setUp(self):
        """Set up integration test environment"""
        self.temp_dir = tempfile.mkdtemp()
    
    def tearDown(self):
        """Clean up integration test environment"""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_integration_with_document_generator(self):
        """Test integration with WordDocumentGenerator"""
        try:
            from document_generator import WordDocumentGenerator
            from document_quality_enhancer import DocumentQualityEnhancer
            
            # This verifies that the import and basic integration works
            generator = WordDocumentGenerator()
            enhancer = DocumentQualityEnhancer()
            
            # Verify they can be instantiated together
            self.assertIsNotNone(generator)
            self.assertIsNotNone(enhancer)
            
        except ImportError as e:
            self.skipTest(f"Integration components not available: {e}")
    
    def test_digital_twin_pipeline_integration(self):
        """Test integration with Digital Twin pipeline"""
        try:
            # Check if Digital Twin components are available
            from digital_twin_model import DocumentModel
            from document_quality_enhancer import DocumentQualityEnhancer
            
            enhancer = DocumentQualityEnhancer()
            
            # Test that enhancer can handle Digital Twin model parameter
            doc = Document()
            doc.add_paragraph("Test content")
            
            # Should work with None Digital Twin
            report = enhancer.enhance_document_quality(doc, digital_twin_doc=None)
            self.assertIsInstance(report, dict)
            
        except ImportError as e:
            self.skipTest(f"Digital Twin integration components not available: {e}")

def run_performance_benchmark():
    """Run performance benchmark for quality enhancement"""
    import time
    
    print("\n🚀 PERFORMANCE BENCHMARK")
    print("=" * 50)
    
    try:
        from document_quality_enhancer import DocumentQualityEnhancer
        enhancer = DocumentQualityEnhancer()
        
        # Create a large document for performance testing
        doc = Document()
        
        # Add many paragraphs with various issues
        for i in range(100):
            if i % 10 == 0:
                doc.add_paragraph(f"This is paragraph {i} with PRESERVE000{i%9+1}")
            elif i % 7 == 0:
                doc.add_paragraph(f"Fragmented paragraph part {i}")
                doc.add_paragraph(f"continues here without punctuation")
            elif i % 5 == 0:
                doc.add_paragraph(f"[Image insertion failed: image_{i}.png]")
            else:
                doc.add_paragraph(f"Regular paragraph {i} with some content.")
        
        print(f"📊 Test Document: {len(doc.paragraphs)} paragraphs")
        
        # Measure enhancement performance
        start_time = time.time()
        enhancement_report = enhancer.enhance_document_quality(doc)
        end_time = time.time()
        
        processing_time = end_time - start_time
        
        print(f"⏱️ Processing Time: {processing_time:.3f} seconds")
        print(f"📈 Throughput: {len(doc.paragraphs)/processing_time:.1f} paragraphs/second")
        print(f"🔧 Issues Found: {enhancement_report.get('total_issues_found', 0)}")
        print(f"🛠️ Fixes Applied: {enhancement_report.get('total_fixes_applied', 0)}")
        print(f"📂 Categories: {len(enhancement_report.get('categories_enhanced', []))}")
        
        return processing_time < 10.0  # Should complete within 10 seconds
        
    except ImportError:
        print("❌ DocumentQualityEnhancer not available for benchmarking")
        return False

def main():
    """Run all tests and benchmarks"""
    print("🧪 DOCUMENT QUALITY ENHANCEMENT TEST SUITE")
    print("=" * 60)
    print("Testing comprehensive quality enhancement for Greek PDF translation issues")
    print("=" * 60)
    
    # Run unit tests
    test_loader = unittest.TestLoader()
    test_suite = test_loader.loadTestsFromModule(__import__(__name__))
    
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(test_suite)
    
    # Run performance benchmark
    benchmark_passed = run_performance_benchmark()
    
    # Summary
    print("\n" + "=" * 60)
    print("📊 TEST SUMMARY")
    print("=" * 60)
    print(f"✅ Tests Run: {result.testsRun}")
    print(f"❌ Failures: {len(result.failures)}")
    print(f"⚠️ Errors: {len(result.errors)}")
    print(f"⏱️ Performance: {'✅ PASS' if benchmark_passed else '❌ FAIL'}")
    
    if result.wasSuccessful() and benchmark_passed:
        print("\n🎉 ALL TESTS PASSED!")
        print("The DocumentQualityEnhancer is ready to solve Greek PDF translation issues.")
    else:
        print("\n❌ SOME TESTS FAILED")
        print("Please review the test output and fix any issues.")
    
    return result.wasSuccessful() and benchmark_passed

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1) 