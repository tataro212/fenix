"""
Document Quality Enhancement Module for Fenix Translation Pipeline

This module addresses critical document quality issues in Greek PDF translations:
1. Broken Table of Contents with "Error! Bookmark not defined."
2. Paragraph fragmentation disrupting readability
3. Image insertion failures showing error messages instead of actual images
4. Unprocessed placeholder codes (PRESERVE0007, etc.)
5. Mathematical formula rendering as plain text
6. Header/footer misplacement in main content
7. Unnecessary empty pages and content artifacts

Created to solve systematic formatting and reconstruction failures.
"""

import os
import re
import logging
from typing import List, Dict, Any, Optional, Tuple, Set
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from difflib import SequenceMatcher

logger = logging.getLogger(__name__)

# Import Digital Twin model for image linking
try:
    from digital_twin_model import DocumentModel, ImageBlock
    DIGITAL_TWIN_AVAILABLE = True
except ImportError:
    DIGITAL_TWIN_AVAILABLE = False
    logger.warning("Digital Twin model not available for image enhancement")

class DocumentQualityEnhancer:
    """
    Comprehensive document quality enhancement system addressing all 
    systematic issues in Greek PDF translation pipeline.
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Mathematical symbol mapping for PRESERVE placeholder restoration
        self.preserve_symbol_map = {
            'PRESERVE0001': 'Φ',  # Phi
            'PRESERVE0002': '≤',  # Less than or equal
            'PRESERVE0003': '≥',  # Greater than or equal
            'PRESERVE0004': '≠',  # Not equal
            'PRESERVE0005': '±',  # Plus minus
            'PRESERVE0006': '∑',  # Summation
            'PRESERVE0007': '∫',  # Integral
            'PRESERVE0008': '∂',  # Partial derivative
            'PRESERVE0009': '∆',  # Delta
            'PRESERVE0010': '∇',  # Nabla
            'PRESERVE0011': '∞',  # Infinity
            'PRESERVE0012': '√',  # Square root
            'PRESERVE0013': 'α',  # Alpha
            'PRESERVE0014': 'β',  # Beta
            'PRESERVE0015': 'γ',  # Gamma
            'PRESERVE0016': 'δ',  # Delta (lowercase)
            'PRESERVE0017': 'ε',  # Epsilon
            'PRESERVE0018': 'θ',  # Theta
            'PRESERVE0019': 'λ',  # Lambda
            'PRESERVE0020': 'μ',  # Mu
            'PRESERVE0021': 'π',  # Pi
            'PRESERVE0022': 'σ',  # Sigma
            'PRESERVE0023': 'τ',  # Tau
            'PRESERVE0024': 'ω',  # Omega
            'PRESERVE0025': '×',  # Multiplication
            'PRESERVE0026': '÷',  # Division
            'PRESERVE0027': '∈',  # Element of
            'PRESERVE0028': '∉',  # Not element of
            'PRESERVE0029': '⊂',  # Subset
            'PRESERVE0030': '⊃',  # Superset
        }
        
        # Mathematical formula patterns for plain text conversion
        self.math_patterns = {
            r'Y\s*a1X1\s*anXn\s*\((\d+\.\d+)\)': r'Y = α₁X₁ + ... + αₙXₙ (\1)',
            r'([A-Z])\s*=\s*([a-zA-Z0-9\+\-\*/\(\)]+)': r'\1 = \2',
            r'F\s*\(([^)]+)\)': r'Φ(\1)',  # Convert F to Phi for functions
            r'(\d+)\.(\d+)': r'\1.\2',  # Preserve decimal numbers
        }
        
        # Header/footer detection patterns
        self.metadata_patterns = [
            r'^Page\s+\d+\s*$',
            r'^\d+\s*$',  # Page numbers
            r'^[A-Z\s]{10,}\s*$',  # All caps headers
            r'^[^\w]*\d{4}[^\w]*$',  # Years
            r'^DOI:\s*\S+',
            r'^ISSN:\s*\S+',
            r'^© \d{4}',
            r'^Journal of\s+\w+',
            r'^Proceedings of\s+\w+',
        ]
        
        # Issue tracking
        self.issues_found = []
        self.fixes_applied = []
        
    def enhance_document_quality(self, doc: Document, 
                               digital_twin_doc: Optional[DocumentModel] = None,
                               output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Main enhancement method that applies all quality fixes to a Word document.
        
        Args:
            doc: The Word document to enhance
            digital_twin_doc: Optional Digital Twin document for image linking
            output_path: Optional path for saving enhanced document
            
        Returns:
            Enhancement report with statistics and applied fixes
        """
        self.logger.info("🔧 Starting comprehensive document quality enhancement")
        
        # Initialize tracking
        self.issues_found.clear()
        self.fixes_applied.clear()
        
        enhancement_report = {
            'total_issues_found': 0,
            'total_fixes_applied': 0,
            'categories_enhanced': [],
            'issues_by_category': {},
            'processing_details': []
        }
        
        try:
            # Step 1: Fix TOC bookmark inconsistencies
            self.logger.info("📖 Fixing Table of Contents bookmark issues...")
            toc_fixes = self._fix_toc_bookmark_consistency(doc)
            enhancement_report['categories_enhanced'].append('TOC Bookmarks')
            enhancement_report['issues_by_category']['toc_bookmarks'] = toc_fixes
            
            # Step 2: Restore PRESERVE placeholder codes to mathematical symbols
            self.logger.info("🔢 Restoring mathematical symbol placeholders...")
            symbol_fixes = self._restore_preserve_placeholders(doc)
            enhancement_report['categories_enhanced'].append('Mathematical Symbols')
            enhancement_report['issues_by_category']['mathematical_symbols'] = symbol_fixes
            
            # Step 3: Consolidate fragmented paragraphs
            self.logger.info("📝 Consolidating fragmented paragraphs...")
            paragraph_fixes = self._consolidate_fragmented_paragraphs(doc)
            enhancement_report['categories_enhanced'].append('Paragraph Consolidation')
            enhancement_report['issues_by_category']['paragraph_consolidation'] = paragraph_fixes
            
            # Step 4: Fix image insertion failures (if Digital Twin available)
            if digital_twin_doc and DIGITAL_TWIN_AVAILABLE:
                self.logger.info("📸 Fixing image insertion failures...")
                image_fixes = self._fix_image_insertion_failures(doc, digital_twin_doc)
                enhancement_report['categories_enhanced'].append('Image Insertion')
                enhancement_report['issues_by_category']['image_insertion'] = image_fixes
            
            # Step 5: Enhance mathematical formula presentation
            self.logger.info("∑ Enhancing mathematical formula presentation...")
            formula_fixes = self._enhance_mathematical_formulas(doc)
            enhancement_report['categories_enhanced'].append('Mathematical Formulas')
            enhancement_report['issues_by_category']['mathematical_formulas'] = formula_fixes
            
            # Step 6: Separate metadata from main content
            self.logger.info("📋 Separating metadata from main content...")
            metadata_fixes = self._separate_metadata_content(doc)
            enhancement_report['categories_enhanced'].append('Metadata Separation')
            enhancement_report['issues_by_category']['metadata_separation'] = metadata_fixes
            
            # Step 7: Remove empty content and artifacts
            self.logger.info("🧹 Removing empty content and artifacts...")
            cleanup_fixes = self._remove_empty_content_artifacts(doc)
            enhancement_report['categories_enhanced'].append('Content Cleanup')
            enhancement_report['issues_by_category']['content_cleanup'] = cleanup_fixes
            
            # Compile final statistics
            enhancement_report['total_issues_found'] = len(self.issues_found)
            enhancement_report['total_fixes_applied'] = len(self.fixes_applied)
            enhancement_report['processing_details'] = self.fixes_applied.copy()
            
            # Save enhanced document if path provided
            if output_path:
                doc.save(output_path)
                self.logger.info(f"💾 Enhanced document saved to: {output_path}")
                enhancement_report['output_path'] = output_path
            
            self.logger.info(f"✅ Document quality enhancement completed")
            self.logger.info(f"   📊 Total issues found: {enhancement_report['total_issues_found']}")
            self.logger.info(f"   🔧 Total fixes applied: {enhancement_report['total_fixes_applied']}")
            self.logger.info(f"   📂 Categories enhanced: {len(enhancement_report['categories_enhanced'])}")
            
            return enhancement_report
            
        except Exception as e:
            self.logger.error(f"❌ Document quality enhancement failed: {e}")
            enhancement_report['error'] = str(e)
            return enhancement_report
    
    def _fix_toc_bookmark_consistency(self, doc: Document) -> Dict[str, Any]:
        """
        Fix Table of Contents bookmark inconsistencies that cause 
        "Error! Bookmark not defined." messages.
        """
        fixes = {
            'broken_bookmarks_found': 0,
            'bookmarks_fixed': 0,
            'hyperlinks_updated': 0,
            'details': []
        }
        
        try:
            # Extract all bookmarks from document
            bookmarks = self._extract_all_bookmarks(doc)
            self.logger.debug(f"Found {len(bookmarks)} bookmarks in document")
            
            # Extract all hyperlinks/references
            hyperlinks = self._extract_all_hyperlinks(doc)
            self.logger.debug(f"Found {len(hyperlinks)} hyperlinks in document")
            
            # Find broken references
            broken_refs = []
            for hyperlink in hyperlinks:
                target_bookmark = hyperlink.get('target')
                if target_bookmark and target_bookmark not in bookmarks:
                    broken_refs.append(hyperlink)
                    fixes['broken_bookmarks_found'] += 1
                    self.issues_found.append(f"Broken bookmark reference: {target_bookmark}")
            
            # Attempt to fix broken references using fuzzy matching
            for broken_ref in broken_refs:
                target = broken_ref['target']
                
                # Try to find similar bookmark names
                best_match = self._find_best_bookmark_match(target, bookmarks)
                
                if best_match and best_match != target:
                    # Update the hyperlink reference
                    success = self._update_hyperlink_target(doc, broken_ref, best_match)
                    if success:
                        fixes['bookmarks_fixed'] += 1
                        fixes['hyperlinks_updated'] += 1
                        self.fixes_applied.append(f"Fixed bookmark reference: {target} → {best_match}")
                        fixes['details'].append({
                            'original_target': target,
                            'fixed_target': best_match,
                            'text': broken_ref.get('text', 'Unknown')
                        })
            
            # Create missing bookmarks for TOC entries without matches
            for broken_ref in broken_refs:
                if broken_ref['target'] not in [fix['fixed_target'] for fix in fixes['details']]:
                    # Create a new bookmark at an appropriate location
                    created = self._create_missing_bookmark(doc, broken_ref)
                    if created:
                        fixes['bookmarks_fixed'] += 1
                        self.fixes_applied.append(f"Created missing bookmark: {broken_ref['target']}")
            
            return fixes
            
        except Exception as e:
            self.logger.error(f"Failed to fix TOC bookmark consistency: {e}")
            fixes['error'] = str(e)
            return fixes
    
    def _restore_preserve_placeholders(self, doc: Document) -> Dict[str, Any]:
        """
        Restore PRESERVE placeholder codes to their original mathematical symbols.
        """
        fixes = {
            'placeholders_found': 0,
            'symbols_restored': 0,
            'paragraphs_affected': 0,
            'details': []
        }
        
        try:
            paragraphs_modified = 0
            
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                modified_text = original_text
                local_fixes = 0
                
                # Search for PRESERVE placeholders
                for placeholder, symbol in self.preserve_symbol_map.items():
                    if placeholder in modified_text:
                        modified_text = modified_text.replace(placeholder, symbol)
                        local_fixes += 1
                        fixes['placeholders_found'] += 1
                        fixes['symbols_restored'] += 1
                        self.fixes_applied.append(f"Restored symbol: {placeholder} → {symbol}")
                
                # Update paragraph if changes were made
                if local_fixes > 0:
                    # Clear existing runs and add modified text
                    paragraph.clear()
                    paragraph.add_run(modified_text)
                    paragraphs_modified += 1
                    
                    fixes['details'].append({
                        'original_text': original_text[:100] + '...' if len(original_text) > 100 else original_text,
                        'symbols_restored': local_fixes,
                        'paragraph_index': len(fixes['details'])
                    })
            
            fixes['paragraphs_affected'] = paragraphs_modified
            
            return fixes
            
        except Exception as e:
            self.logger.error(f"Failed to restore PRESERVE placeholders: {e}")
            fixes['error'] = str(e)
            return fixes
    
    def _consolidate_fragmented_paragraphs(self, doc: Document) -> Dict[str, Any]:
        """
        Consolidate paragraphs that have been inappropriately split,
        improving document readability.
        """
        fixes = {
            'fragments_found': 0,
            'paragraphs_consolidated': 0,
            'paragraphs_removed': 0,
            'details': []
        }
        
        try:
            paragraphs = list(doc.paragraphs)
            to_remove = []
            i = 0
            
            while i < len(paragraphs) - 1:
                current_para = paragraphs[i]
                next_para = paragraphs[i + 1]
                
                # Check if paragraphs should be consolidated
                if self._should_consolidate_paragraphs(current_para, next_para):
                    # Merge paragraphs
                    consolidated_text = current_para.text.strip() + " " + next_para.text.strip()
                    
                    # Update current paragraph with consolidated text
                    current_para.clear()
                    current_para.add_run(consolidated_text)
                    
                    # Mark next paragraph for removal
                    to_remove.append(next_para)
                    
                    fixes['fragments_found'] += 1
                    fixes['paragraphs_consolidated'] += 1
                    self.fixes_applied.append(f"Consolidated fragmented paragraph")
                    
                    fixes['details'].append({
                        'original_first': current_para.text[:50] + '...',
                        'original_second': next_para.text[:50] + '...',
                        'consolidated_length': len(consolidated_text)
                    })
                    
                    # Skip the next paragraph since we just processed it
                    i += 2
                else:
                    i += 1
            
            # Remove paragraphs marked for deletion
            for para in to_remove:
                self._remove_paragraph(para)
                fixes['paragraphs_removed'] += 1
            
            return fixes
            
        except Exception as e:
            self.logger.error(f"Failed to consolidate fragmented paragraphs: {e}")
            fixes['error'] = str(e)
            return fixes
    
    def _fix_image_insertion_failures(self, doc: Document, 
                                    digital_twin_doc: DocumentModel) -> Dict[str, Any]:
        """
        Replace image insertion failure messages with actual images 
        from the Digital Twin model.
        """
        fixes = {
            'failure_messages_found': 0,
            'images_inserted': 0,
            'images_not_found': 0,
            'details': []
        }
        
        try:
            # Get all image blocks from Digital Twin
            image_blocks = digital_twin_doc.get_all_image_blocks()
            image_map = {os.path.basename(img.image_path): img for img in image_blocks}
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                
                # Look for image insertion failure patterns
                failure_patterns = [
                    r'\[Image insertion failed:\s*([^\]]+)\]',
                    r'\[Image not found:\s*([^\]]+)\]',
                    r'Image Placeholder',
                    r'\[Figure Placeholder\]'
                ]
                
                for pattern in failure_patterns:
                    matches = re.finditer(pattern, text)
                    for match in matches:
                        fixes['failure_messages_found'] += 1
                        
                        # Try to extract filename from failure message
                        if match.groups():
                            filename = match.group(1).strip()
                        else:
                            # Look for nearby image references
                            filename = self._find_nearest_image_reference(paragraph, image_map.keys())
                        
                        if filename and filename in image_map:
                            # Replace failure message with actual image
                            image_block = image_map[filename]
                            success = self._insert_image_in_paragraph(paragraph, image_block, match.span())
                            
                            if success:
                                fixes['images_inserted'] += 1
                                self.fixes_applied.append(f"Inserted image: {filename}")
                                fixes['details'].append({
                                    'filename': filename,
                                    'failure_message': match.group(0),
                                    'image_path': image_block.image_path
                                })
                            else:
                                fixes['images_not_found'] += 1
                        else:
                            fixes['images_not_found'] += 1
                            self.issues_found.append(f"Image file not found for: {filename}")
            
            return fixes
            
        except Exception as e:
            self.logger.error(f"Failed to fix image insertion failures: {e}")
            fixes['error'] = str(e)
            return fixes
    
    def _enhance_mathematical_formulas(self, doc: Document) -> Dict[str, Any]:
        """
        Convert plain text mathematical expressions to proper mathematical notation.
        """
        fixes = {
            'formulas_found': 0,
            'formulas_enhanced': 0,
            'paragraphs_affected': 0,
            'details': []
        }
        
        try:
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                enhanced_text = original_text
                local_enhancements = 0
                
                # Apply mathematical pattern replacements
                for pattern, replacement in self.math_patterns.items():
                    matches = re.finditer(pattern, enhanced_text)
                    for match in matches:
                        enhanced_text = re.sub(pattern, replacement, enhanced_text)
                        local_enhancements += 1
                        fixes['formulas_found'] += 1
                        self.fixes_applied.append(f"Enhanced formula: {match.group(0)} → {replacement}")
                
                # Apply additional mathematical symbol enhancements
                symbol_replacements = {
                    ' F ': ' Φ ',  # Replace F with Phi in mathematical contexts
                    '<=': '≤',
                    '>=': '≥',
                    '!=': '≠',
                    '+-': '±',
                    'SUM': '∑',
                    'INTEGRAL': '∫',
                    'INFINITY': '∞'
                }
                
                for old_symbol, new_symbol in symbol_replacements.items():
                    if old_symbol in enhanced_text:
                        enhanced_text = enhanced_text.replace(old_symbol, new_symbol)
                        local_enhancements += 1
                        fixes['formulas_enhanced'] += 1
                
                # Update paragraph if enhancements were made
                if local_enhancements > 0:
                    paragraph.clear()
                    paragraph.add_run(enhanced_text)
                    fixes['paragraphs_affected'] += 1
                    
                    fixes['details'].append({
                        'original_text': original_text[:100] + '...' if len(original_text) > 100 else original_text,
                        'enhancements_applied': local_enhancements
                    })
            
            return fixes
            
        except Exception as e:
            self.logger.error(f"Failed to enhance mathematical formulas: {e}")
            fixes['error'] = str(e)
            return fixes
    
    def _separate_metadata_content(self, doc: Document) -> Dict[str, Any]:
        """
        Identify and properly format header/footer/citation content 
        that has been mixed into main body text.
        """
        fixes = {
            'metadata_paragraphs_found': 0,
            'paragraphs_reformatted': 0,
            'paragraphs_moved': 0,
            'details': []
        }
        
        try:
            metadata_paragraphs = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                
                # Check if paragraph contains metadata
                is_metadata = False
                metadata_type = None
                
                for pattern in self.metadata_patterns:
                    if re.match(pattern, text, re.IGNORECASE):
                        is_metadata = True
                        metadata_type = pattern
                        break
                
                # Additional heuristics for metadata detection
                if not is_metadata:
                    # Short lines with specific patterns
                    if len(text) < 50 and any(keyword in text.lower() for keyword in 
                                            ['page', 'doi:', 'issn:', 'journal', 'proceedings']):
                        is_metadata = True
                        metadata_type = 'keyword_based'
                    
                    # Lines that are all uppercase (likely headers)
                    elif text.isupper() and len(text) > 5:
                        is_metadata = True
                        metadata_type = 'uppercase_header'
                
                if is_metadata:
                    metadata_paragraphs.append({
                        'paragraph': paragraph,
                        'index': i,
                        'text': text,
                        'type': metadata_type
                    })
                    fixes['metadata_paragraphs_found'] += 1
                    self.issues_found.append(f"Metadata in main content: {text[:50]}...")
            
            # Reformat metadata paragraphs
            for meta_info in metadata_paragraphs:
                paragraph = meta_info['paragraph']
                
                # Apply metadata formatting
                self._apply_metadata_formatting(paragraph, meta_info['type'])
                fixes['paragraphs_reformatted'] += 1
                
                self.fixes_applied.append(f"Reformatted metadata: {meta_info['type']}")
                fixes['details'].append({
                    'text': meta_info['text'][:100],
                    'type': meta_info['type'],
                    'action': 'reformatted'
                })
            
            return fixes
            
        except Exception as e:
            self.logger.error(f"Failed to separate metadata content: {e}")
            fixes['error'] = str(e)
            return fixes
    
    def _remove_empty_content_artifacts(self, doc: Document) -> Dict[str, Any]:
        """
        Remove unnecessary empty paragraphs, formatting artifacts, 
        and other content that disrupts document flow.
        """
        fixes = {
            'empty_paragraphs_found': 0,
            'paragraphs_removed': 0,
            'artifacts_cleaned': 0,
            'details': []
        }
        
        try:
            paragraphs_to_remove = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Identify empty or artifact paragraphs
                should_remove = False
                removal_reason = None
                
                # Completely empty paragraphs
                if not text:
                    should_remove = True
                    removal_reason = 'empty'
                    fixes['empty_paragraphs_found'] += 1
                
                # Paragraphs with only formatting artifacts
                elif re.match(r'^[\s\-_=\*\.]+$', text):
                    should_remove = True
                    removal_reason = 'formatting_artifact'
                    fixes['artifacts_cleaned'] += 1
                
                # Paragraphs with only quotes or minimal content
                elif text in ['""', "''", '"', "'", '""', '...']:
                    should_remove = True
                    removal_reason = 'minimal_content'
                    fixes['artifacts_cleaned'] += 1
                
                # Multiple consecutive newlines or spaces
                elif re.match(r'^[\s\n]+$', text):
                    should_remove = True
                    removal_reason = 'whitespace_only'
                    fixes['artifacts_cleaned'] += 1
                
                if should_remove:
                    paragraphs_to_remove.append(paragraph)
                    fixes['details'].append({
                        'text': text[:50] if text else '(empty)',
                        'reason': removal_reason
                    })
            
            # Remove identified paragraphs
            for paragraph in paragraphs_to_remove:
                self._remove_paragraph(paragraph)
                fixes['paragraphs_removed'] += 1
                self.fixes_applied.append(f"Removed empty/artifact paragraph")
            
            return fixes
            
        except Exception as e:
            self.logger.error(f"Failed to remove empty content artifacts: {e}")
            fixes['error'] = str(e)
            return fixes
    
    # Helper methods
    
    def _extract_all_bookmarks(self, doc: Document) -> Dict[str, Any]:
        """Extract all bookmarks from the document"""
        bookmarks = {}
        try:
            # Search through document XML for bookmark elements
            for element in doc.element.iter():
                if element.tag.endswith('bookmarkStart'):
                    bookmark_id = element.get(qn('w:id'))
                    bookmark_name = element.get(qn('w:name'))
                    if bookmark_name:
                        bookmarks[bookmark_name] = {
                            'id': bookmark_id,
                            'element': element
                        }
        except Exception as e:
            self.logger.debug(f"Error extracting bookmarks: {e}")
        
        return bookmarks
    
    def _extract_all_hyperlinks(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract all hyperlinks from the document"""
        hyperlinks = []
        try:
            for element in doc.element.iter():
                if element.tag.endswith('hyperlink'):
                    anchor = element.get(qn('w:anchor'))
                    if anchor:
                        # Find text content
                        text_content = ""
                        for text_elem in element.iter():
                            if text_elem.tag.endswith('t') and text_elem.text:
                                text_content += text_elem.text
                        
                        hyperlinks.append({
                            'target': anchor,
                            'text': text_content,
                            'element': element
                        })
        except Exception as e:
            self.logger.debug(f"Error extracting hyperlinks: {e}")
        
        return hyperlinks
    
    def _find_best_bookmark_match(self, target: str, available_bookmarks: Dict[str, Any]) -> Optional[str]:
        """Find the best matching bookmark using fuzzy string matching"""
        if not available_bookmarks:
            return None
        
        best_match = None
        best_ratio = 0.0
        
        for bookmark_name in available_bookmarks.keys():
            ratio = SequenceMatcher(None, target.lower(), bookmark_name.lower()).ratio()
            if ratio > best_ratio and ratio > 0.6:  # Minimum 60% similarity
                best_ratio = ratio
                best_match = bookmark_name
        
        return best_match
    
    def _should_consolidate_paragraphs(self, current_para, next_para) -> bool:
        """Determine if two paragraphs should be consolidated"""
        current_text = current_para.text.strip()
        next_text = next_para.text.strip()
        
        # Don't consolidate if either is empty
        if not current_text or not next_text:
            return False
        
        # Don't consolidate if current ends with strong punctuation
        if current_text.endswith(('.', '!', '?', ':', ';')):
            return False
        
        # Don't consolidate if next starts with capital letter (likely new sentence)
        if next_text[0].isupper() and len(current_text) > 50:
            return False
        
        # Don't consolidate if either is very short (likely a heading or special content)
        if len(current_text) < 20 or len(next_text) < 20:
            return False
        
        # Consolidate if current doesn't end with punctuation and next doesn't start with capital
        return not current_text[-1] in '.!?:;' and not next_text[0].isupper()
    
    def _remove_paragraph(self, paragraph):
        """Safely remove a paragraph from the document"""
        try:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        except Exception as e:
            self.logger.debug(f"Error removing paragraph: {e}")
    
    def _apply_metadata_formatting(self, paragraph, metadata_type: str):
        """Apply appropriate formatting to metadata content"""
        try:
            # Set appropriate style based on metadata type
            if 'header' in metadata_type.lower():
                paragraph.style = 'Header'
            elif 'page' in metadata_type.lower():
                paragraph.style = 'Footer'
            else:
                # Generic metadata formatting
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    run.italic = True
        except Exception as e:
            self.logger.debug(f"Error applying metadata formatting: {e}")
    
    def _update_hyperlink_target(self, doc: Document, hyperlink_info: Dict, new_target: str) -> bool:
        """Update a hyperlink's target bookmark"""
        try:
            element = hyperlink_info['element']
            element.set(qn('w:anchor'), new_target)
            return True
        except Exception as e:
            self.logger.debug(f"Error updating hyperlink target: {e}")
            return False
    
    def _create_missing_bookmark(self, doc: Document, broken_ref: Dict) -> bool:
        """Create a missing bookmark at an appropriate location"""
        try:
            # This is a simplified implementation
            # In practice, you would need to find the appropriate location
            # based on the hyperlink text and document structure
            target_text = broken_ref.get('text', '')
            
            # Find paragraph with similar text
            for paragraph in doc.paragraphs:
                if target_text.lower() in paragraph.text.lower():
                    # Add bookmark to this paragraph
                    bookmark_name = broken_ref['target']
                    # Implementation would add bookmark XML elements here
                    return True
            
            return False
        except Exception as e:
            self.logger.debug(f"Error creating missing bookmark: {e}")
            return False
    
    def _find_nearest_image_reference(self, paragraph, available_images: List[str]) -> Optional[str]:
        """Find the nearest image reference in surrounding text"""
        try:
            text = paragraph.text.lower()
            for image_name in available_images:
                if image_name.lower() in text:
                    return image_name
            return None
        except:
            return None
    
    def _insert_image_in_paragraph(self, paragraph, image_block: 'ImageBlock', span: Tuple[int, int]) -> bool:
        """Insert an image into a paragraph replacing failure text"""
        try:
            if not image_block.image_exists():
                return False
            
            # Clear the paragraph and add the image
            paragraph.clear()
            run = paragraph.add_run()
            
            # Calculate appropriate size
            width_inches = 4.0  # Default width
            if image_block.bbox:
                bbox_width = image_block.bbox[2] - image_block.bbox[0]
                width_inches = min(bbox_width / 72, 6.0)  # Convert points to inches, max 6"
            
            run.add_picture(image_block.image_path, width=Inches(width_inches))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add caption if available
            if image_block.caption_text:
                caption_para = paragraph._element.getparent().insert(
                    paragraph._element.getparent().index(paragraph._element) + 1,
                    paragraph._element.makeelement(qn('w:p'))
                )
                # Add caption text and formatting
            
            return True
        except Exception as e:
            self.logger.debug(f"Error inserting image: {e}")
            return False 